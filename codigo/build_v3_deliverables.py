from __future__ import annotations

import html
import json
import shutil
import subprocess
from pathlib import Path

import nbformat as nbf
import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
RESULT_PATH = ROOT / "artefactos" / "v3" / "resultados_v3.json"
R = json.loads(RESULT_PATH.read_text(encoding="utf-8"))
V1, V2, V3 = (R["referencias_historicas"][key] for key in ("V1", "V2", "V3"))
PROMO = R["promocion"]
MODEL = R["modelo_v3"]
URL = "https://github.com/DanielBarillasM/Proyecto-1_Grupo-1_DLYSI_Sec-30"


def f(value, digits=3):
    return f"{float(value):,.{digits}f}"


def pct(value):
    return f"{100 * float(value):.2f}%"


def money(value):
    return f"Q{float(value):,.0f}"


def build_readme() -> None:
    walk_rows = "\n".join(
        f"| {row['modelo']} | {f(row['mean'], 4)} | {f(row['std'], 4)} |"
        for row in R["validacion_walk_forward"]["resumen"]
    )
    logistic_rows = "\n".join(
        f"| {row['modelo']} | {f(row['auc_pr_validacion'], 4)} | {f(row['roc_auc_validacion'], 4)} | {row['variables_transformadas']} | {'Sí' if row.get('convergio') else 'No'} |"
        for row in R["baselines_logisticos"]["resultados"]
    )
    text = f"""<div align="center">

# Proyecto 1 · Monitoreo transaccional — V3

### Detección de fraude con variables causales, LightGBM nativo y validación temporal

![Python](https://img.shields.io/badge/Python-3.13-3776AB?logo=python&logoColor=white)
![LightGBM](https://img.shields.io/badge/LightGBM-4.7.0-184e77)
![Estado](https://img.shields.io/badge/V3-PROMOVIDA-2a9d8f)
![Protocolo](https://img.shields.io/badge/Protocolo-walk--forward-e9c46a)

**Wilson Alejandro Calderón Argueta · 22018** · **Pablo Daniel Barillas Moreno · 22193**

Universidad del Valle de Guatemala · Deep Learning y Sistemas Inteligentes · Sección 30 · 2026

</div>

> [!IMPORTANT]
> V3 cumplió los cuatro criterios predefinidos de promoción. El último 15% cronológico continúa siendo un **benchmark histórico reutilizado**, no una prueba ciega. La promoción se decidió con tres ventanas walk-forward y un holdout de umbral anterior al benchmark.

## Contenido

- [Resumen ejecutivo](#resumen-ejecutivo)
- [Problema y datos](#problema-y-datos)
- [Qué cambió en V3](#qué-cambió-en-v3)
- [Correlación, PCA y regresión logística](#correlación-pca-y-regresión-logística)
- [Protocolo temporal](#protocolo-temporal)
- [Resultados](#resultados)
- [Umbrales y decisión económica](#umbrales-y-decisión-económica)
- [Operación, ética y limitaciones](#operación-ética-y-limitaciones)
- [Reproducción](#reproducción)
- [Referencias APA 7](#referencias-apa-7)

## Resumen ejecutivo

El proyecto estudia detección de fraude sobre **{R['datos']['filas']:,} transacciones** IEEE-CIS, con **{pct(R['datos']['prevalencia'])}** de prevalencia y 434 columnas integradas. La tarea es fuertemente desbalanceada; por ello la métrica de ranking principal es AUC-PR y no exactitud. La decisión operativa se complementa con precisión, recall, F1, calibración, capacidad de revisión y un costo académico de Q4,200 por falso negativo y Q180 por falso positivo.

V3 se promueve porque no depende de una mejora aislada. Ganó las tres ventanas temporales, elevó AUC-PR walk-forward de **{f(PROMO['v2_auc_pr_walk'],4)} a {f(PROMO['v3_auc_pr_walk'],4)}**, redujo el costo del holdout de **{money(PROMO['v2_validacion_holdout']['costo_q'])} a {money(PROMO['v3_validacion_holdout']['costo_q'])}** y aumentó recall de **{pct(PROMO['v2_validacion_holdout']['recall'])} a {pct(PROMO['v3_validacion_holdout']['recall'])}**. Los cuatro criterios de promoción resultaron verdaderos.

En el benchmark histórico reutilizado, el umbral balanceado de V3 logra AUC-PR **{f(V3['auc_pr'],4)}**, precisión **{pct(V3['precision'])}**, recall **{pct(V3['recall'])}**, F1 **{f(V3['f1'],4)}** y costo **{money(V3['costo_q'])}**. Frente a V2 mejora todas esas métricas y reduce el costo {pct(R['comparacion_benchmark_v3_vs_v2']['reduccion_costo_relativa'])}. Esta comparación es descriptiva; la evidencia confirmatoria requerirá una cohorte temporal nueva.

## Problema y datos

Los archivos `train_transaction.csv` y `train_identity.csv` se integran mediante `TransactionID` y se ordenan con `TransactionDT`. Ninguna de estas dos columnas entra al modelo como una magnitud predictiva: la primera es una llave técnica y la segunda define el orden. Las tarjetas, direcciones, dominios y dispositivos son códigos anonimizados; se tratan como categorías o componentes de una identidad proxy.

Para una transacción en tiempo $t$, las variables históricas cumplen:

$$x_t^{{hist}}=f\\left(\\{{x_j:t_j<t\\}}\\right).$$

Se calculan causalmente conteo previo, media y desviación histórica del monto, razón entre monto actual y promedio anterior, tiempo desde la operación previa y conteos en 1, 6, 24 y 72 horas. El evento actual se incorpora al historial solamente después de emitir sus variables. Esto evita que una observación se describa usando su propio futuro.

La clave principal `card1 + card2 + card3 + card5 + addr1` produce {R['identidad_secuencial']['tarjeta_direccion']['entidades']:,} entidades aproximadas. No equivale a una persona real: puede mezclar usuarios o fragmentar uno mismo. La cobertura histórica se reporta para advertir que una secuencia nominalmente larga no implica antecedentes confiables.

## Qué cambió en V3

| Dimensión | V2 | V3 promovida |
|---|---|---|
| Numéricas | 110 | 220 |
| Categóricas | 18 codificadas como enteros | 24 categorías nativas |
| Valores faltantes | mediana | `NaN` nativo + conteo de ausencia |
| Redundancia | $|\\rho_s|≥0.985$ | $|\\rho_s|≥0.995$ |
| Entrenamiento | uniforme | uniforme, recencia y 300k recientes |
| Baseline lineal | solo dentro del ensamble | L2, L1, Elastic Net y PCA independientes |
| Umbral | costo en validación completa | early stopping, calibración y umbral en bloques separados |
| Segmentos | ensamble no recomendado | candidato V3 recomendado |
| Promoción | comparación descriptiva | regla de cuatro criterios congelados |

LightGBM recibe categorías declaradas como tales, no códigos ordinales cuyo orden sería artificial. Los valores numéricos faltantes se conservan para que el árbol pueda aprender rutas específicas de ausencia. Además, la ponderación exponencial con vida media de 75 días permite que observaciones recientes influyan más sin descartar por completo el pasado.

## Correlación, PCA y regresión logística

La selección se aprende únicamente en el 55% inicial. Pearson describe asociación lineal, Spearman asociación monótona e información mutua dependencias no lineales. Ninguna implica causalidad. Una correlación marginal baja tampoco basta para excluir una variable, porque los árboles pueden explotar interacciones. La poda solo retira sustitutos casi monotónicos y excluye IDs, constantes o ausencia extrema.

PCA continúa como ablación, no como supuesto de mejora. En V2 redujo dimensión, pero perdió AUC-PR. En V3, la logística con 64 componentes explicó {pct(next(row['varianza_pca'] for row in R['baselines_logisticos']['resultados'] if row['modelo']=='Logistica_PCA64'))} de la varianza numérica y obtuvo AUC-PR {f(next(row['auc_pr_validacion'] for row in R['baselines_logisticos']['resultados'] if row['modelo']=='Logistica_PCA64'),4)}. Conservar varianza no equivale a conservar señal discriminativa.

| Baseline | AUC-PR validación | ROC-AUC | Dimensión | Convergió |
|---|---:|---:|---:|---|
{logistic_rows}

Las variantes `saga` alcanzaron el límite de 100 iteraciones, por lo que se documentan como controles lineales no convergidos y no como optimizaciones exhaustivas. Aun así, su AUC-PR cercana a 0.331 muestra que gran parte de la ventaja de V3 proviene de interacciones no lineales y particiones categóricas, no solo de incorporar más columnas.

![Baselines logísticos](../evidencia/figuras/v3/03_logistica_v3.png)

## Protocolo temporal

Las tres ventanas walk-forward siempre entrenan con pasado y evalúan el futuro inmediato. Dentro de cada ventana se aprenden categorías, poda, pesos y early stopping sin incorporar observaciones posteriores. El benchmark final no decide el modelo ni sus hiperparámetros.

| Modelo | AUC-PR media | Desviación |
|---|---:|---:|
{walk_rows}

La variante con recencia ganó F1, F2 y F3. La mejora no se explica únicamente por usar menos o más filas: el modelo uniforme también supera claramente a V2, mientras el entrenamiento con 300 mil recientes queda detrás de ponderar suavemente todo el pasado.

![Walk-forward V3](../evidencia/figuras/v3/01_walk_forward_v3.png)

## Resultados

| Versión | Modelo | AUC-PR | Precisión | Recall | F1 | Costo |
|---|---|---:|---:|---:|---:|---:|
| V1 | HistGradientBoosting A | {f(V1['auc_pr'])} | {pct(V1['precision'])} | {pct(V1['recall'])} | {f(V1['f1'])} | {money(V1['costo_q'])} |
| V2 | LightGBM depurado | {f(V2['auc_pr'])} | {pct(V2['precision'])} | {pct(V2['recall'])} | {f(V2['f1'])} | {money(V2['costo_q'])} |
| **V3** | **LightGBM nativo + recencia** | **{f(V3['auc_pr'])}** | **{pct(V3['precision'])}** | **{pct(V3['recall'])}** | **{f(V3['f1'])}** | **{money(V3['costo_q'])}** |

La diferencia pareada de AUC-PR V3−V2 en el benchmark es {f(R['comparacion_pareada_benchmark']['delta_auc_pr'],4)}, con intervalo descriptivo por bloques [{f(R['comparacion_pareada_benchmark']['li95'],4)}, {f(R['comparacion_pareada_benchmark']['ls95'],4)}]. Que el intervalo quede por encima de cero refuerza la consistencia descriptiva, aunque no convierte al benchmark reutilizado en prueba ciega.

![Curvas PR](../evidencia/figuras/v3/02_curvas_pr_v2_v3.png)

## Umbrales y decisión económica

El umbral principal, **{f(MODEL['threshold_recomendado_balanceado'],5)}**, maximiza F1 dentro del holdout sujeto a recall mínimo de 70%. Es el punto recomendado porque mejora simultáneamente precisión, recall, F1 y costo frente a V2. Produce {f(V3['alertas_por_100k'],0)} alertas por 100 mil transacciones.

El umbral económico alternativo, **{f(MODEL['threshold_economico'],5)}**, minimiza:

$$C(\\tau)=Q4,200\\,FN(\\tau)+Q180\\,FP(\\tau).$$

En el benchmark recupera {pct(MODEL['benchmark_economico']['recall'])} del fraude y reduce el costo a {money(MODEL['benchmark_economico']['costo_q'])}, pero genera más alertas y baja precisión. Por ello no existe un único “mejor” umbral independiente de capacidad operativa. La política balanceada es adecuada para la entrega; la económica sirve como escenario cuando omitir fraude domina la carga de revisión.

![Frontera costo-recall](../evidencia/figuras/v3/04_costo_recall_v3.png)

## Operación, ética y limitaciones

`Precision@K` y `Recall@K` traducen el ranking a cupos de revisión. Al revisar el 1% con mayor riesgo, V3 obtiene precisión {pct(next(x['precision_at_k'] for x in R['metricas_top_k'] if x['tasa_revision']==0.01))} y recupera {pct(next(x['recall_at_k'] for x in R['metricas_top_k'] if x['tasa_revision']==0.01))} del fraude. Los segmentos por producto, dispositivo, monto e historial se calcularon con el modelo y umbral recomendados.

El sistema es un prototipo académico para **priorizar revisión humana**. No debe bloquear operaciones ni atribuir culpabilidad. Antes de producción se requieren una cohorte reciente nunca observada, validación de costos reales, privacidad, seguridad, explicabilidad, análisis de sesgo, monitoreo de deriva y capacidad operativa. La identidad proxy, anonimización, horizonte de 182 días y reutilización del benchmark limitan generalización.

V1 y V2 se retiraron del árbol activo solo después de que V3 cumplió la regla de promoción. Su historia permanece recuperable mediante Git y V3 conserva una referencia mínima de métricas/predicciones para auditar la comparación.

## Reproducción

```powershell
py -3.13 -m venv .venv
.\\.venv\\Scripts\\Activate.ps1
python -m pip install -r configuracion/v3/requirements-v3.txt
python -m pip install -r configuracion/v3/requirements-docs-v3.txt
python codigo/download_data.py
python -u codigo/proyecto1_v3_pipeline.py
python codigo/postprocess_v3.py
python codigo/finalize_v3.py
python codigo/build_v3_deliverables.py
python codigo/audit_project1_v3.py
```

```text
.github/README.md                  documentación principal
artefactos/v3/                    resultados, modelos y predicciones
codigo/                            pipeline, constructores y auditoría
configuracion/v3/                 dependencias e instrucciones
datos/processed/v3/               selección y correlación
datos/raw/                        IEEE-CIS local, ignorado por Git
entregables/cuaderno/              reporte ejecutable V3
entregables/informe/               LaTeX y PDF V3
entregables/presentacion/          HTML y PDF de 8 diapositivas
entregables/ficha/                 ficha DOCX/PDF del repositorio
evidencia/figuras/v3/              visualizaciones de la versión promovida
legal/                             licencia
```

La fuente única de verdad es [`artefactos/v3/resultados_v3.json`](../artefactos/v3/resultados_v3.json). Consulte [`configuracion/v3/INSTRUCCIONES_V3.md`](../configuracion/v3/INSTRUCCIONES_V3.md) para detalles.

## Referencias APA 7

IEEE Computational Intelligence Society. (2019). *IEEE-CIS Fraud Detection* [Conjunto de datos]. Kaggle. https://www.kaggle.com/competitions/ieee-fraud-detection

Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T.-Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. *Advances in Neural Information Processing Systems, 30*. https://papers.neurips.cc/paper/6907-lightgbm-a-highly-efficient-gradient-boosting-decision-tree

Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. *PLOS ONE, 10*(3), e0118432. https://doi.org/10.1371/journal.pone.0118432

Scikit-learn developers. (2026). *LogisticRegression*. https://scikit-learn.org/stable/modules/generated/sklearn.linear_model.LogisticRegression.html

Scikit-learn developers. (2026). *Probability calibration*. https://scikit-learn.org/stable/modules/calibration.html

### Declaración de IA

Se utilizó IA para estructurar código, redacción, visualización y controles de consistencia. Los autores ejecutaron los experimentos, verificaron las cifras y asumen responsabilidad por interpretación, seguridad y defensa. La IA no se considera fuente académica.
"""
    (ROOT / ".github" / "README.md").write_text(text, encoding="utf-8", newline="\n")


def notebook_style() -> str:
    return """<style>.hero{padding:40px 44px;border-radius:24px;color:#f8fbff;background:linear-gradient(125deg,#102a43,#184e77 55%,#2a9d8f);box-shadow:0 16px 38px #102a433d;font-family:Inter,'Segoe UI',sans-serif}.hero h1{font-size:38px;color:white;margin:12px 0}.chips span{display:inline-block;padding:6px 13px;margin:3px;border:1px solid #ffffff55;border-radius:999px;background:#ffffff20;font-size:12px;font-weight:800}.grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(200px,1fr));gap:12px}.card{padding:15px;border-radius:12px;background:#ffffff16;border:1px solid #ffffff38}.section{margin:28px 0 14px;padding:14px 20px;border-radius:13px;background:linear-gradient(90deg,#102a43,#184e77);color:white}.call{margin:17px 0;padding:20px 23px;border:1px solid #c9d9e6;border-left:6px solid #2a9d8f;border-radius:13px;background:#edf7f6;color:#172033}.warn{border-left-color:#e9a23b;background:#fff8e8}.kpi{display:inline-block;min-width:155px;margin:5px;padding:13px;border-radius:11px;background:#edf5fb;border:1px solid #c9d9e6;text-align:center}.kpi b{display:block;font-size:22px;color:#184e77}table{width:100%}th{background:#184e77!important;color:white!important;text-align:left!important}td,th{padding:9px!important}</style>"""


def build_notebook() -> None:
    hero = notebook_style() + """<div class="hero"><div class="chips"><span>DEEP LEARNING</span><span>PROYECTO 1</span><span>V3 PROMOVIDA</span><span>GRUPO 1 · SECCIÓN 30</span></div><h1>Monitoreo transaccional y detección de fraude</h1><p style="font-size:20px">LightGBM nativo, recencia, calibración y decisión económica</p><div class="grid"><div class="card"><b>Institución</b><br>Universidad del Valle de Guatemala</div><div class="card"><b>Curso</b><br>Deep Learning y Sistemas Inteligentes</div><div class="card"><b>Docente</b><br>Kevin Recinos</div><div class="card"><b>Integrantes</b><br>Wilson Calderón · 22018<br>Pablo Barillas · 22193</div><div class="card"><b>Dataset</b><br>IEEE-CIS Fraud Detection</div><div class="card"><b>Candidato</b><br>LGB native recency V3</div></div></div>"""
    cells = [nbf.v4.new_markdown_cell(hero)]
    cells.append(nbf.v4.new_code_cell("from pathlib import Path\nimport json, pandas as pd\nfrom IPython.display import Image, display\nROOT=Path.cwd()\nif not (ROOT/'artefactos').exists(): ROOT=ROOT.parents[1]\nR=json.loads((ROOT/'artefactos/v3/resultados_v3.json').read_text(encoding='utf-8'))\nR['recomendacion']"))
    sections = [
        ("1 · Resumen ejecutivo", f"""<div class="call"><b>Decisión:</b> V3 se promueve porque cumplió los cuatro criterios definidos antes de consultar el benchmark. Ganó 3/3 folds, elevó AP walk-forward en {f(PROMO['delta_auc_pr_walk'],4)} y redujo costo holdout {pct(PROMO['reduccion_costo_holdout'])}.</div>

El objetivo no fue aumentar complejidad indiscriminadamente, sino mejorar representación, protocolo y punto operativo. AUC-PR evalúa ranking bajo desbalance; el costo $$C(\\tau)=4200FN(\\tau)+180FP(\\tau)$$ expresa el supuesto económico académico.""", "pd.DataFrame(R['promocion']['criterios'],index=['cumple']).T"),
        ("2 · Datos y causalidad", """Los 590,540 eventos se ordenan por `TransactionDT`. `TransactionID` une tablas, pero ninguno entra como predictor continuo. Para cada evento, $$x_t^{hist}=f(\\{x_j:t_j<t\\}).$$ Los agregados de monto, recencia y frecuencia se emiten antes de incorporar la transacción actual.""", "pd.DataFrame(R['datos']['particiones']).T"),
        ("3 · Variables, correlación y categorías", """La selección se aprende en el 55% inicial. Se retienen 220 numéricas y 24 categóricas; solo se eliminan redundancias casi monotónicas con $|\\rho_s|\\ge0.995$. LightGBM recibe categorías nativas y conserva `NaN`, evitando orden artificial e imputación que borre patrones de ausencia.""", "pd.DataFrame({'tipo':['numéricas','categóricas'],'cantidad':[R['seleccion_variables']['n_numericas'],R['seleccion_variables']['n_categoricas']]})"),
        ("4 · Walk-forward", """Tres futuros simulados comparan entrenamiento uniforme, ponderación por recencia y las 300 mil observaciones más recientes. La recencia gana los tres folds; la evidencia sugiere deriva gradual y utilidad parcial del pasado antiguo.""", "pd.read_csv(ROOT/'artefactos/v3/validacion_walk_forward_v3.csv').pivot(index='fold',columns='modelo',values='auc_pr')"),
        ("5 · Regresión logística y PCA", """L2, L1 y Elastic Net corroboran una línea base lineal. Alcanzaron el máximo de iteraciones y se etiquetan como no convergidos. PCA64 conserva 99.53% de varianza, pero reduce AUC-PR: varianza total y señal de fraude no son equivalentes.""", "pd.read_csv(ROOT/'artefactos/v3/baselines_logisticos_v3.csv')"),
        ("6 · Calibración y umbrales", """Early stopping, calibración y umbral usan bloques sucesivos. El umbral balanceado maximiza F1 con recall mínimo de 70%; el económico minimiza costo. La calibración reduce Brier y ECE sin alterar el orden de los puntajes.""", "pd.DataFrame(R['modelo_v3']['calibracion'],index=['valor']).T"),
        ("7 · Resultados V1/V2/V3", """V3 mejora simultáneamente AUC-PR, precisión, recall, F1 y costo frente a V2 usando el umbral balanceado. El benchmark es histórico reutilizado, de modo que la tabla es evidencia descriptiva y no confirmación externa.""", "pd.DataFrame(R['referencias_historicas']).T[['auc_pr','precision','recall','f1','costo_q']]"),
        ("8 · Capacidad y segmentos", """Top-K conecta el ranking con cupos reales de analistas. Los segmentos fueron recalculados con el candidato y umbral recomendados; no describen el ensamble descartado.""", "pd.DataFrame(R['metricas_top_k'])"),
        ("9 · Conclusión y límites", """La mejora proviene principalmente de más cobertura, categorías nativas, ausencia preservada y ponderación de recencia. La GRU no se amplió porque la falsificación V1 mostró señal de orden débil. Antes de producción se necesita cohorte nueva, identidad confiable, costos reales, privacidad, equidad, seguridad y monitoreo.""", "R['limitaciones']"),
    ]
    for title, body, code in sections:
        cells.append(nbf.v4.new_markdown_cell(f'<div class="section"><h2>{title}</h2></div>\n\n{body}'))
        cells.append(nbf.v4.new_code_cell(code))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>Referencias APA 7</h2></div>

IEEE Computational Intelligence Society. (2019). *IEEE-CIS Fraud Detection* [Conjunto de datos]. Kaggle. https://www.kaggle.com/competitions/ieee-fraud-detection

Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T.-Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. *Advances in Neural Information Processing Systems, 30*.

Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. *PLOS ONE, 10*(3), e0118432. https://doi.org/10.1371/journal.pone.0118432

**Declaración de IA.** Se utilizó asistencia para código, redacción y auditoría. Los autores ejecutaron y verificaron el experimento y asumen responsabilidad por sus conclusiones."""))
    notebook = nbf.v4.new_notebook(cells=cells, metadata={"kernelspec":{"display_name":"Python 3","language":"python","name":"python3"},"language_info":{"name":"python","version":"3.13"}})
    nbf.write(notebook, ROOT / "entregables" / "cuaderno" / "Proyecto_1_Monitoreo_Transaccional_V3.ipynb")


def build_report() -> Path:
    tex = rf"""\documentclass[10pt]{{article}}
\usepackage[utf8]{{inputenc}}\usepackage[T1]{{fontenc}}\usepackage[letterpaper,margin=1.55cm]{{geometry}}
\usepackage{{graphicx,booktabs,xcolor,amsmath,hyperref,array}}\definecolor{{navy}}{{HTML}}{{102A43}}\definecolor{{teal}}{{HTML}}{{2A9D8F}}
\hypersetup{{colorlinks=true,urlcolor=teal}}\setlength{{\parindent}}{{0pt}}\setlength{{\parskip}}{{5pt}}
\begin{{document}}\begin{{titlepage}}\pagecolor{{navy}}\color{{white}}\raggedright\vspace*{{1cm}}{{\Large DEEP LEARNING Y SISTEMAS INTELIGENTES\par}}\vspace{{1cm}}{{\Huge\bfseries Proyecto 1\\Monitoreo transaccional --- V3\par}}\vspace{{.7cm}}{{\LARGE LightGBM nativo, recencia y validación temporal\par}}\vfill{{\Large Universidad del Valle de Guatemala\\Kevin Recinos · Grupo 1 · Sección 30\\Semestre II 2026\par}}\vspace{{1cm}}{{\Large Wilson Alejandro Calderón Argueta · 22018\\Pablo Daniel Barillas Moreno · 22193\par}}\vfill\textbf{{Estado:}} V3 promovida por cuatro criterios. Benchmark final histórico reutilizado.\end{{titlepage}}\nopagecolor

\section*{{Resumen ejecutivo}} Se analizaron {R['datos']['filas']:,} transacciones IEEE-CIS con prevalencia {pct(R['datos']['prevalencia'])}. V3 amplía a 220 variables numéricas y 24 categóricas, conserva ausencias, declara categorías nativas y pondera recencia. Ganó tres de tres ventanas y elevó AP walk-forward de {PROMO['v2_auc_pr_walk']:.4f} a {PROMO['v3_auc_pr_walk']:.4f}. En el holdout de umbral redujo costo {100*PROMO['reduccion_costo_holdout']:.1f}\% y aumentó recall. Se promueve V3; la confirmación requiere una cohorte nueva.

\section{{Datos, identidad y fuga}} Las tablas se unen con TransactionID y ordenan con TransactionDT; ninguna entra como magnitud predictiva. Para cada $t$, $x_t^{{hist}}=f(\{{x_j:t_j<t\}})$. Se crean conteos previos, media/desviación de monto, razón de monto, tiempo desde operación anterior y actividad 1/6/24/72 h. La clave de tarjeta--dirección genera {R['identidad_secuencial']['tarjeta_direccion']['entidades']:,} proxies; puede mezclar o fragmentar usuarios. El protocolo usa 70\% train, 15\% validación y 15\% benchmark reutilizado.

\section{{Selección y modelos}} Correlación e información mutua se aprenden en el 55\% inicial. Se excluyen IDs, constantes y ausencia extrema; la poda $|\rho_s|\ge0.995$ solo elimina sustitutos casi monotónicos. LightGBM usa categorías nativas y NaN; se comparan pesos uniformes, vida media de 75 días y 300k recientes. Como control se ajustan logísticas L2, L1, Elastic Net y PCA64. Las tres variantes SAGA no convergieron en 100 iteraciones y se tratan como baselines diagnósticos.

\begin{{figure}}[h]\centering\includegraphics[width=.78\linewidth]{{../../evidencia/figuras/v3/01_walk_forward_v3.png}}\caption{{Comparación walk-forward de V3.}}\end{{figure}}
\newpage
\section{{Validación temporal}} Las AP medias fueron {R['validacion_walk_forward']['resumen'][0]['mean']:.4f} con recencia, {R['validacion_walk_forward']['resumen'][1]['mean']:.4f} uniforme y {R['validacion_walk_forward']['resumen'][2]['mean']:.4f} con 300k recientes. Recencia gana los tres folds. Frente a V2, el aumento medio es {PROMO['delta_auc_pr_walk']:.4f}; por tanto, supera el mínimo predefinido de 0.015 sin depender de un solo corte.

PCA64 explica 99.53\% de varianza y obtiene AP 0.3032, mientras Elastic Net obtiene 0.3312. Esto confirma que conservar varianza no equivale a conservar señal supervisada y que el problema necesita interacciones no lineales.

\section{{Resultados y promoción}}\begin{{center}}\begin{{tabular}}{{lrrrrr}}\toprule Versión&AUC-PR&Prec.&Recall&F1&Costo\\\midrule V1&{V1['auc_pr']:.3f}&{V1['precision']:.3f}&{V1['recall']:.3f}&{V1['f1']:.3f}&{money(V1['costo_q'])}\\V2&{V2['auc_pr']:.3f}&{V2['precision']:.3f}&{V2['recall']:.3f}&{V2['f1']:.3f}&{money(V2['costo_q'])}\\\textbf{{V3}}&\textbf{{{V3['auc_pr']:.3f}}}&\textbf{{{V3['precision']:.3f}}}&\textbf{{{V3['recall']:.3f}}}&\textbf{{{V3['f1']:.3f}}}&\textbf{{{money(V3['costo_q'])}}}\\\bottomrule\end{{tabular}}\end{{center}}

La promoción no usa el benchmark: exige AP walk-forward +0.015, costo holdout $-3\%$, recall dentro de un punto y ganar al menos dos folds. V3 cumple los cuatro. En el benchmark reutilizado, la diferencia pareada AP es {R['comparacion_pareada_benchmark']['delta_auc_pr']:.4f}, IC descriptivo [{R['comparacion_pareada_benchmark']['li95']:.4f}, {R['comparacion_pareada_benchmark']['ls95']:.4f}].

\begin{{figure}}[h]\centering\includegraphics[width=.68\linewidth]{{../../evidencia/figuras/v3/02_curvas_pr_v2_v3.png}}\caption{{Curvas PR en benchmark histórico reutilizado.}}\end{{figure}}
\newpage
\section{{Calibración y umbrales}} La validación se divide cronológicamente para early stopping, calibración y selección de umbral. La calibración reduce Brier de {MODEL['calibracion']['brier_raw_validacion']:.4f} a {MODEL['calibracion']['brier_calibrado_validacion']:.4f} y ECE de {MODEL['calibracion']['ece_raw_validacion']:.4f} a {MODEL['calibracion']['ece_calibrado_validacion']:.4f}.

El umbral balanceado {MODEL['threshold_recomendado_balanceado']:.5f} maximiza F1 con recall mínimo 70\%; mejora simultáneamente todas las métricas frente a V2. El umbral económico {MODEL['threshold_economico']:.5f} minimiza $4200FN+180FP$: alcanza recall {MODEL['benchmark_economico']['recall']:.3f} y costo {money(MODEL['benchmark_economico']['costo_q'])}, pero genera más alertas y menor precisión.

\begin{{figure}}[h]\centering\includegraphics[width=.75\linewidth]{{../../evidencia/figuras/v3/04_costo_recall_v3.png}}\caption{{Frontera entre recall y costo en el holdout.}}\end{{figure}}

\section{{Operación, límites y ética}} Al revisar el 1\% de mayor riesgo, precisión es {100*next(x['precision_at_k'] for x in R['metricas_top_k'] if x['tasa_revision']==0.01):.2f}\% y recall {100*next(x['recall_at_k'] for x in R['metricas_top_k'] if x['tasa_revision']==0.01):.2f}\%. Se reportan segmentos de producto, dispositivo, monto e historia con el candidato recomendado. El prototipo prioriza revisión humana; no bloquea transacciones ni atribuye culpabilidad.

Limitaciones: benchmark reutilizado, identidad proxy, anonimización, 182 días y costos académicos. Antes de producción se requieren cohorte nueva, privacidad, seguridad, explicabilidad, sesgo, costos reales, latencia y monitoreo de deriva. La GRU no se amplió porque la falsificación V1 no mostró señal de orden suficiente.

\section{{Conclusión}} V3 mejora de manera significativa y consistente. El salto se atribuye a cobertura ampliada, categorías nativas, ausencia preservada y recencia; regresión logística y PCA no sustituyen las interacciones del árbol. V1/V2 pueden retirarse del árbol activo porque Git conserva su historia y V3 mantiene referencias mínimas para auditoría.

\section*{{Referencias}}\small IEEE Computational Intelligence Society. (2019). \textit{{IEEE-CIS Fraud Detection}}. Kaggle.\\Ke, G., et al. (2017). LightGBM. \textit{{NeurIPS, 30}}.\\Saito, T., \& Rehmsmeier, M. (2015). Precision-recall plots for imbalanced data. \textit{{PLOS ONE, 10}}(3), e0118432.\\Scikit-learn developers. (2026). \textit{{LogisticRegression; Probability calibration}}.

\section*{{Declaración de IA}} IA apoyó código, redacción, visualización y auditoría. Los autores ejecutaron, verificaron y asumen responsabilidad.\end{{document}}"""
    out = ROOT / "entregables" / "informe" / "informe_proyecto1_v3.tex"
    out.write_text(tex, encoding="utf-8", newline="\n")
    pdflatex = shutil.which("pdflatex")
    if pdflatex:
        for _ in range(2):
            subprocess.run([pdflatex,"-interaction=nonstopmode","-halt-on-error","-output-directory",str(out.parent),str(out)],cwd=out.parent,check=False,capture_output=True)
    return out


def build_slides() -> Path:
    slides = [
        ("DECISIÓN", "V3 supera la regla de promoción", f'<div class="big">{f(PROMO["v3_auc_pr_walk"],3)}</div><p>AUC-PR walk-forward</p><div class="call">3/3 folds ganados · costo holdout −{100*PROMO["reduccion_costo_holdout"]:.1f}%</div>', "La promoción se decidió sin usar el benchmark final."),
        ("DATOS", "590,540 eventos; futuro fuera", '<div class="cols"><div class="box"><b>434</b><p>columnas integradas</p></div><div class="box"><b>70·15·15</b><p>train, validación, benchmark</p></div></div><p class="formula">xₜʰⁱˢᵗ=f({xⱼ:tⱼ&lt;t})</p>', "TransactionID une; TransactionDT ordena."),
        ("REPRESENTACIÓN", "Más señal, menos supuestos artificiales", '<div class="cols"><div class="box"><b>220</b><p>numéricas</p></div><div class="box"><b>24</b><p>categorías nativas</p></div></div><ul><li>NaN preservado</li><li>poda ρ≥0.995</li><li>vida media 75 días</li></ul>', "Explicar que baja correlación marginal no obliga a excluir."),
        ("WALK-FORWARD", "Recencia gana los tres futuros", '<img src="../../evidencia/figuras/v3/01_walk_forward_v3.png"><div class="call">V2 0.473 → V3 0.581</div>', "Comparar media y dispersión; no solo el mejor fold."),
        ("BASELINES", "La linealidad no basta", '<img src="../../evidencia/figuras/v3/03_logistica_v3.png"><p>Elastic Net AP 0.331 · PCA64 AP 0.303</p>', "Las variantes SAGA no convergieron en 100 iteraciones; son controles."),
        ("RESULTADOS", "V3 mejora todas las métricas operativas", f'<div class="cols"><div class="box accent"><b>AP {f(V3["auc_pr"])}</b><p>Precisión {pct(V3["precision"])}</p><p>Recall {pct(V3["recall"])}</p></div><div class="box"><b>F1 {f(V3["f1"])}</b><p>{money(V3["costo_q"])}</p><p>{f(V3["alertas_por_100k"],0)} alertas/100k</p></div></div>', "Benchmark reutilizado: resultado descriptivo."),
        ("UMBRAL", "Balance operativo, no una cifra mágica", f'<div class="cols"><div class="box"><b>{f(MODEL["threshold_recomendado_balanceado"],4)}</b><p>balanceado</p></div><div class="box"><b>{f(MODEL["threshold_economico"],4)}</b><p>mínimo costo</p></div></div><img src="../../evidencia/figuras/v3/04_costo_recall_v3.png">', "El económico recupera más fraude pero aumenta carga."),
        ("RECOMENDACIÓN", "Promover V3; confirmar con cohorte nueva", '<div class="road"><span>revisión humana</span><span>monitorear deriva</span><span>costos reales</span><span>privacidad</span><span>cohorte nueva</span></div><div class="call">V1/V2 retiradas del árbol activo; historia preservada en Git.</div>', "No prometer producción ni prueba ciega."),
    ]
    cards = []
    for index, (eye, title, body, notes) in enumerate(slides, 1):
        cards.append(f'<section class="slide"><div class="eye">{eye}</div><h1>{title}</h1><div class="content">{body}</div><footer>Grupo 1 · Sección 30 <b>{index}/8</b></footer><aside>{notes}</aside></section>')
    page = f'''<!doctype html><html lang="es"><head><meta charset="utf-8"><title>Proyecto 1 · V3</title><style>*{{box-sizing:border-box}}body{{margin:0;background:#081a2a;color:#eef7ff;font-family:Inter,"Segoe UI",sans-serif;overflow:hidden}}.slide{{display:none;width:100vw;height:100vh;padding:6vh 7vw;background:radial-gradient(circle at 90% 10%,#2a9d8f55,transparent 30%),linear-gradient(135deg,#102a43,#091b2b);position:relative}}.slide.active{{display:block}}.eye{{color:#65d3c3;font-weight:900;letter-spacing:.16em;font-size:1.2vw}}h1{{font-size:4.1vw;line-height:1.06;margin:2vh 0 4vh}}p,li{{font-size:1.55vw;line-height:1.45}}.content{{height:68vh}}.big{{font-size:11vw;font-weight:900;color:#65d3c3;line-height:1}}.cols{{display:grid;grid-template-columns:1fr 1fr;gap:2.5vw}}.box{{padding:1.8vw;border:1px solid #ffffff33;border-radius:1.2vw;background:#ffffff0d}}.box b{{font-size:3vw;color:#65d3c3}}.accent{{border-color:#65d3c3}}.call{{padding:1.2vw;margin-top:2vh;border-left:.5vw solid #2a9d8f;background:#ffffff12;border-radius:.8vw;font-size:1.45vw}}img{{display:block;max-width:82%;max-height:48vh;margin:auto;border-radius:1vw;background:white}}.formula{{text-align:center;font-size:2.8vw;color:#65d3c3}}.road{{display:flex;flex-wrap:wrap;gap:1vw}}.road span{{padding:1.2vw;border:1px solid #65d3c3;border-radius:999px;background:#184e77;font-size:1.35vw}}footer{{position:absolute;bottom:3vh;left:7vw;right:7vw;display:flex;justify-content:space-between;color:#9eb6c8}}aside{{display:none}}body.notes aside{{display:block;position:absolute;right:2vw;bottom:7vh;width:31vw;padding:1vw;background:#fff;color:#172033;border-radius:.7vw}}@page{{size:13.333in 7.5in;margin:0}}@media print{{html,body{{width:13.333in;height:7.5in;overflow:visible}}.slide{{display:block!important;width:13.333in;height:7.5in;page-break-after:always}}aside{{display:none!important}}}}</style></head><body>{''.join(cards)}<script>const s=[...document.querySelectorAll('.slide')];let i=0;function show(n){{i=Math.max(0,Math.min(s.length-1,n));s.forEach((x,j)=>x.classList.toggle('active',j===i))}}document.onkeydown=e=>{{if(['ArrowRight',' ','PageDown'].includes(e.key))show(i+1);if(['ArrowLeft','PageUp'].includes(e.key))show(i-1);if(e.key.toLowerCase()==='n')document.body.classList.toggle('notes')}};show(0)</script></body></html>'''
    out = ROOT / "entregables" / "presentacion" / "presentacion_proyecto1_v3.html"
    out.write_text(page, encoding="utf-8", newline="\n")
    edge_candidates = [shutil.which("msedge"), Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"), Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe")]
    edge = next((str(x) for x in edge_candidates if x and Path(x).exists()), None)
    if edge:
        pdf = out.with_suffix(".pdf")
        subprocess.run([edge,"--headless","--disable-gpu","--no-pdf-header-footer",f"--print-to-pdf={pdf}",out.resolve().as_uri()],check=False,capture_output=True)
    return out


def build_ficha() -> None:
    out_dir = ROOT / "entregables" / "ficha"
    qr_path = ROOT / "evidencia" / "recursos" / "qr_repositorio_v3.png"
    qrcode.make(URL).save(qr_path)
    doc = Document(); section = doc.sections[0]; section.top_margin=Inches(.45); section.bottom_margin=Inches(.45); section.left_margin=Inches(.55); section.right_margin=Inches(.55)
    title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; run=title.add_run("PROYECTO 1 · MONITOREO TRANSACCIONAL · V3"); run.bold=True; run.font.size=Pt(20); run.font.color.rgb=RGBColor(24,78,119)
    sub=doc.add_paragraph("Universidad del Valle de Guatemala · Deep Learning y Sistemas Inteligentes · Sección 30"); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    names=doc.add_paragraph("Wilson Alejandro Calderón Argueta · 22018  |  Pablo Daniel Barillas Moreno · 22193"); names.alignment=WD_ALIGN_PARAGRAPH.CENTER
    table=doc.add_table(rows=2,cols=4); table.style="Table Grid"
    for cell,(label,value) in zip([c for row in table.rows for c in row.cells],[('AUC-PR',f(V3['auc_pr'])),('Precisión',pct(V3['precision'])),('Recall',pct(V3['recall'])),('F1',f(V3['f1'])),('Costo',money(V3['costo_q'])),('AP walk',f(PROMO['v3_auc_pr_walk'])),('Folds','3/3'),('Estado','PROMOVIDA')]):
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run(label.upper()+"\n"); rr.bold=True; rr.font.size=Pt(8); vv=p.add_run(value); vv.bold=True; vv.font.size=Pt(13); vv.font.color.rgb=RGBColor(42,157,143)
    doc.add_heading("Resumen",level=1); doc.add_paragraph("V3 integra 220 variables numéricas y 24 categóricas, conserva valores faltantes, usa LightGBM con categorías nativas y pondera recencia. Cumplió cuatro criterios predefinidos y mejora simultáneamente AUC-PR, precisión, recall, F1 y costo frente a V2 con el umbral balanceado.")
    doc.add_heading("Reproducción y uso",level=1); doc.add_paragraph("Fuente única: artefactos/v3/resultados_v3.json. Instrucciones: configuracion/v3/INSTRUCCIONES_V3.md. Uso académico para priorizar revisión humana; no bloquear transacciones. Benchmark final histórico reutilizado.")
    pic=doc.add_paragraph(); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER; pic.add_run().add_picture(str(qr_path),width=Inches(1.25)); link=doc.add_paragraph(URL); link.alignment=WD_ALIGN_PARAGRAPH.CENTER
    docx_path=out_dir/"Ficha_Repositorio_Proyecto_1_V3.docx"; doc.save(docx_path)

    styles=getSampleStyleSheet(); pdf=SimpleDocTemplate(str(out_dir/"Ficha_Repositorio_Proyecto_1_V3.pdf"),pagesize=letter,rightMargin=.55*inch,leftMargin=.55*inch,topMargin=.45*inch,bottomMargin=.45*inch)
    story=[Paragraph("<b>PROYECTO 1 · MONITOREO TRANSACCIONAL · V3</b>",styles['Title']),Paragraph("Universidad del Valle de Guatemala · Grupo 1 · Sección 30",styles['Heading3']),Spacer(1,10)]
    data=[["AUC-PR",f(V3['auc_pr']),"Precisión",pct(V3['precision'])],["Recall",pct(V3['recall']),"F1",f(V3['f1'])],["Costo",money(V3['costo_q']),"Estado","PROMOVIDA"]]; tab=Table(data,colWidths=[1.1*inch,1.35*inch,1.1*inch,1.35*inch]); tab.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,-1),colors.HexColor('#EDF5FB')),('GRID',(0,0),(-1,-1),.5,colors.HexColor('#9FB7C8')),('FONTNAME',(0,0),(-1,-1),'Helvetica-Bold'),('ALIGN',(0,0),(-1,-1),'CENTER'),('PADDING',(0,0),(-1,-1),8)])); story.extend([tab,Spacer(1,14),Paragraph("V3 usa 220 numéricas, 24 categorías nativas, NaN preservado, recencia y validación walk-forward. Ganó 3/3 folds y cumplió los cuatro criterios de promoción.",styles['BodyText']),Spacer(1,10),Paragraph("Uso académico para priorizar revisión humana. Benchmark histórico reutilizado; requiere cohorte nueva antes de producción.",styles['BodyText']),Spacer(1,14),Paragraph(URL,styles['Heading3'])]); pdf.build(story)


def main() -> None:
    build_readme(); build_notebook(); report=build_report(); slides=build_slides(); build_ficha()
    print("README, notebook, informe, presentación y ficha V3 generados desde", RESULT_PATH)
    print("Informe:", report, "Presentación:", slides)


if __name__ == "__main__":
    main()
