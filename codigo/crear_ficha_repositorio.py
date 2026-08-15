"""Genera una ficha DOCX profesional y reproducible del repositorio."""

from __future__ import annotations

import json
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULTS_PATH = ROOT / "artefactos" / "resultados.json"
OUTPUT_DIR = ROOT / "entregables" / "ficha"
ASSET_DIR = ROOT / "evidencia" / "recursos"
OUTPUT_PATH = OUTPUT_DIR / "Ficha_Repositorio_Proyecto1.docx"
QR_PATH = ASSET_DIR / "qr_repositorio.png"
REPOSITORY_URL = "https://github.com/DanielBarillasM/Proyecto-1_Grupo-1_DLYSI_Sec-30"

NAVY = "102A43"
BLUE = "184E77"
TEAL = "2A9D8F"
PALE = "EDF5FB"
INK = RGBColor(23, 32, 51)
MUTED = RGBColor(89, 108, 124)
WHITE = RGBColor(255, 255, 255)


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, size: float, color: RGBColor, bold=False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = color


def add_hyperlink(paragraph, text: str, url: str, color="184E77") -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((run_color, underline))
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend((properties, text_node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_section_label(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(42, 157, 143)


def add_compact_bullet(cell, label: str, value: str) -> None:
    paragraph = cell.add_paragraph(style=None)
    paragraph.paragraph_format.left_indent = Cm(0.15)
    paragraph.paragraph_format.space_after = Pt(2)
    marker = paragraph.add_run("• ")
    marker.font.color.rgb = RGBColor(42, 157, 143)
    marker.bold = True
    key = paragraph.add_run(label)
    key.bold = True
    key.font.name = "Aptos"
    key.font.size = Pt(8.4)
    value_run = paragraph.add_run(value)
    value_run.font.name = "Aptos"
    value_run.font.size = Pt(8.4)
    value_run.font.color.rgb = INK


def create_repository_sheet() -> Path:
    results = json.loads(RESULTS_PATH.read_text(encoding="utf-8"))
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    qr = qrcode.QRCode(version=4, box_size=8, border=2)
    qr.add_data(REPOSITORY_URL)
    qr.make(fit=True)
    qr.make_image(fill_color=f"#{NAVY}", back_color="white").save(QR_PATH)

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.15)
    section.right_margin = Cm(1.15)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(8.6)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    properties = document.core_properties
    properties.title = "Ficha del repositorio — Proyecto 1: Monitoreo transaccional"
    properties.subject = "Deep Learning y Sistemas Inteligentes"
    properties.author = "Wilson Alejandro Calderón Argueta; Pablo Daniel Barillas Moreno"
    properties.keywords = "IEEE-CIS, fraude, GRU, Deep Learning, monitoreo transaccional"

    header = document.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    header.columns[0].width = Cm(15.4)
    header.columns[1].width = Cm(3.0)
    left, right = header.rows[0].cells
    for cell in (left, right):
        shade(cell, NAVY)
        set_cell_margins(cell, top=175, start=210, bottom=170, end=210)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    eyebrow = left.paragraphs[0]
    eyebrow.paragraph_format.space_after = Pt(3)
    run = eyebrow.add_run("UNIVERSIDAD DEL VALLE DE GUATEMALA · GRUPO 1 · SECCIÓN 30")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(7.8)
    run.font.color.rgb = RGBColor(159, 227, 216)

    title = left.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Ficha del repositorio")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(23)
    run.font.color.rgb = WHITE

    subtitle = left.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(0)
    run = subtitle.add_run("Proyecto 1 · Monitoreo transaccional")
    run.font.name = "Aptos"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(220, 238, 248)

    image_paragraph = right.paragraphs[0]
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(QR_PATH), width=Cm(2.5))
    label = right.add_paragraph("ABRIR REPOSITORIO")
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.space_after = Pt(0)
    label.runs[0].bold = True
    label.runs[0].font.size = Pt(6.5)
    label.runs[0].font.color.rgb = WHITE

    meta = document.add_table(rows=2, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    values = [
        ("CURSO", "Deep Learning y Sistemas Inteligentes"),
        ("DOCENTE", "Kevin Recinos"),
        ("RUTA", "B · Datos reales de Kaggle"),
        ("CANDIDATO", f"Pieza {results['candidate']}"),
        ("INTEGRANTE", "Wilson Alejandro Calderón Argueta · 22018"),
        ("INTEGRANTE", "Pablo Daniel Barillas Moreno · 22193"),
        ("DATOS", f"{results['dataset']['rows']:,} transacciones"),
        ("FRAUDE", f"{results['dataset']['fraud_rate']:.3%}"),
    ]
    for cell, (label_text, value) in zip((c for row in meta.rows for c in row.cells), values):
        shade(cell, PALE)
        set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(label_text)
        r.bold = True
        r.font.size = Pt(6.7)
        r.font.color.rgb = RGBColor(24, 78, 119)
        p = cell.add_paragraph(value)
        p.paragraph_format.space_after = Pt(0)
        p.runs[0].font.size = Pt(7.7)
        p.runs[0].font.color.rgb = INK

    add_section_label(document, "Propósito y diseño experimental")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(
        "El repositorio evalúa si el historial ordenado de transacciones aporta señal incremental "
        "frente a una representación agregada. Compara A: HistGradientBoosting sin orden, B: "
        "embeddings + GRU(32), y C: fusión GRU–agregados. La partición es cronológica 70/15/15; "
        "validación congela arquitectura y umbrales antes de abrir prueba."
    )

    add_section_label(document, "Resultados verificables en prueba cronológica")
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Modelo", "AUC-PR", "Precisión", "Recall", "F1", "Costo")
    for cell, value in zip(table.rows[0].cells, headers):
        shade(cell, BLUE)
        set_cell_text(cell, value, 7.5, WHITE, bold=True)
    for model in ("A", "B", "C"):
        metric = results["test"][model]
        cells = table.add_row().cells
        values = (
            model,
            f"{metric['auc_pr']:.3f}",
            f"{metric['precision']:.3f}",
            f"{metric['recall']:.3f}",
            f"{metric['f1']:.3f}",
            f"Q{metric['cost_q']:,.0f}",
        )
        for cell, value in zip(cells, values):
            shade(cell, "F7FAFC" if model != results["candidate"] else "DDF3EE")
            set_cell_text(cell, value, 7.8, INK, bold=model == results["candidate"])

    callout = document.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    shade(cell, "DDF3EE")
    set_cell_margins(cell, top=110, start=150, bottom=110, end=150)
    set_cell_text(
        cell,
        f"DECISIÓN · Conservar A. La permutación redujo AUC-PR solo "
        f"{results['falsification']['order_auc_pr_drop']:.3f}, por debajo del criterio previo de 0.01; "
        "la evidencia no justifica todavía migrar a una arquitectura secuencial.",
        8.4,
        INK,
        bold=True,
    )

    add_section_label(document, "Contenido y reproducción")
    columns = document.add_table(rows=1, cols=2)
    columns.alignment = WD_TABLE_ALIGNMENT.CENTER
    columns.autofit = False
    left, right = columns.rows[0].cells
    left.width = Cm(9.1)
    right.width = Cm(9.1)
    for cell in (left, right):
        shade(cell, "F7FAFC")
        set_cell_margins(cell, top=80, start=110, bottom=80, end=110)
        cell.text = ""

    heading = left.paragraphs[0]
    heading.add_run("ORGANIZACIÓN").bold = True
    heading.runs[0].font.size = Pt(7.4)
    heading.runs[0].font.color.rgb = RGBColor(24, 78, 119)
    add_compact_bullet(left, "entregables/", "cuaderno, informe, presentación y ficha.")
    add_compact_bullet(left, "codigo/", "pipeline, descarga, construcción y auditoría.")
    add_compact_bullet(left, "artefactos/", "modelos, umbrales, esquema y métricas.")
    add_compact_bullet(left, "evidencia/", "figuras y recursos visuales reproducibles.")

    heading = right.paragraphs[0]
    heading.add_run("CONTROLES").bold = True
    heading.runs[0].font.size = Pt(7.4)
    heading.runs[0].font.color.rgb = RGBColor(24, 78, 119)
    add_compact_bullet(right, "Temporalidad: ", "test posterior y abierto una sola vez.")
    add_compact_bullet(right, "Falsificación: ", "cinco permutaciones + historial truncado.")
    add_compact_bullet(right, "Economía: ", "Q4,200 por FN y Q180 por FP.")
    add_compact_bullet(right, "Seguridad: ", "datos y credenciales no se versionan.")

    add_section_label(document, "Trazabilidad, alcance y límites")
    trace = document.add_table(rows=2, cols=3)
    trace.alignment = WD_TABLE_ALIGNMENT.CENTER
    trace.style = "Table Grid"
    trace_headers = ("DATOS", "EVIDENCIA", "USO RESPONSABLE")
    trace_text = (
        "IEEE-CIS · 182 días\n70/15/15 cronológico\n590,540 transacciones",
        "A/B/C bajo protocolo común\n5 permutaciones + recorte\nUmbral económico validado",
        "Identidad aproximada\nVariables anonimizadas\nRequiere piloto y monitoreo",
    )
    for cell, value in zip(trace.rows[0].cells, trace_headers):
        shade(cell, TEAL)
        set_cell_text(cell, value, 7.2, WHITE, bold=True)
    for cell, value in zip(trace.rows[1].cells, trace_text):
        shade(cell, "F7FAFC")
        set_cell_margins(cell, top=80, start=100, bottom=80, end=100)
        cell.text = ""
        for index, line in enumerate(value.splitlines()):
            paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            marker = paragraph.add_run("• ")
            marker.bold = True
            marker.font.color.rgb = RGBColor(42, 157, 143)
            run = paragraph.add_run(line)
            run.font.name = "Aptos"
            run.font.size = Pt(7.7)
            run.font.color.rgb = INK

    scope = document.add_table(rows=1, cols=1)
    scope.alignment = WD_TABLE_ALIGNMENT.CENTER
    scope_cell = scope.cell(0, 0)
    shade(scope_cell, "FFF4E8")
    set_cell_margins(scope_cell, top=90, start=140, bottom=90, end=140)
    set_cell_text(
        scope_cell,
        "ALCANCE · Repositorio académico reproducible y base para un piloto controlado; "
        "no constituye un motor de autorización listo para producción.",
        8,
        INK,
        bold=True,
    )

    add_section_label(document, "Acceso al repositorio")
    link_paragraph = document.add_paragraph()
    link_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link_paragraph.paragraph_format.space_after = Pt(1)
    add_hyperlink(link_paragraph, REPOSITORY_URL, REPOSITORY_URL)
    note = document.add_paragraph(
        "El README contiene instrucciones completas de reproducción, estructura y limitaciones. "
        "La extrapolación económica es un escenario analítico, no una cifra contable del banco."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(0)
    for run in note.runs:
        run.italic = True
        run.font.size = Pt(7.3)
        run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "PROYECTO 1 · GRUPO 1 · SECCIÓN 30 · SEMESTRE II 2026"
    footer.runs[0].font.size = Pt(6.5)
    footer.runs[0].font.color.rgb = MUTED

    document.save(OUTPUT_PATH)
    print(f"Ficha creada: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == "__main__":
    create_repository_sheet()
