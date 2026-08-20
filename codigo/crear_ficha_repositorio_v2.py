from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
R = json.loads((ROOT / "artefactos/v2/resultados_v2.json").read_text(encoding="utf-8"))
OUT = ROOT / "entregables/ficha/Ficha_Repositorio_Proyecto_1_V2.docx"
QR = ROOT / "evidencia/recursos/qr_repositorio_v2.png"
URL = "https://github.com/DanielBarillasM/Proyecto-1_Grupo-1_DLYSI_Sec-30"


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)


def add_text(cell, text, size=9, bold=False, color="172033"):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def main():
    QR.parent.mkdir(parents=True, exist_ok=True)
    qrcode.make(URL).save(QR)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.38)
    sec.bottom_margin = Inches(0.38)
    sec.left_margin = Inches(0.48)
    sec.right_margin = Inches(0.48)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor(23, 32, 51)

    head = doc.add_table(rows=1, cols=2)
    head.columns[0].width = Inches(5.9)
    head.columns[1].width = Inches(1.25)
    left, right = head.rows[0].cells
    shade(left, "102A43")
    shade(right, "102A43")
    margins(left, 180, 220, 180, 220)
    margins(right, 100, 100, 100, 100)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("PROYECTO 1 · DEEP LEARNING")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(101, 211, 195)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Monitoreo transaccional")
    r.bold = True
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor(248, 251, 255)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Ficha técnica del repositorio · Versión 2")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(220, 238, 248)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    right.paragraphs[0].add_run().add_picture(str(QR), width=Inches(1.05))

    info = doc.add_table(rows=2, cols=3)
    info.style = "Table Grid"
    values = [
        ("Institución", "Universidad del Valle de Guatemala"),
        ("Curso", "Deep Learning y Sistemas Inteligentes · Sec. 30"),
        ("Docente", "Kevin Recinos"),
        ("Integrantes", "Wilson A. Calderón A. · 22018"),
        ("Integrantes", "Pablo D. Barillas M. · 22193"),
        ("Repositorio", "DanielBarillasM / Proyecto-1_Grupo-1_DLYSI_Sec-30"),
    ]
    for cell, (label, value) in zip(
        [c for row in info.rows for c in row.cells], values
    ):
        margins(cell)
        shade(cell, "EDF5FB")
        add_text(cell, label.upper(), 7, True, "376F9E")
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(value)
        rr.font.size = Pt(8.5)
        rr.font.color.rgb = RGBColor(23, 32, 51)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    rr = p.add_run("RESUMEN EJECUTIVO")
    rr.bold = True
    rr.font.size = Pt(10)
    rr.font.color.rgb = RGBColor(24, 78, 119)
    p = doc.add_paragraph(
        f"IEEE-CIS: {R['datos']['filas']:,} transacciones y {100 * R['datos']['prevalencia']:.2f}% de fraude. La V2 audita todas las variables, construye agregados causales y compara LightGBM, CatBoost y PCA mediante tres ventanas temporales. El 15% final es un benchmark histórico reutilizado, no test ciego."
    )
    p.paragraph_format.space_after = Pt(5)

    tab = R["modelo_tabular_v2"]["benchmark_historico"]
    ens = R["ensamble_v2"]["benchmark_historico"]
    metrics = doc.add_table(rows=2, cols=4)
    metrics.style = "Table Grid"
    cards = [
        ("AUC-PR LightGBM", f"{tab['auc_pr']:.3f}"),
        ("Recall LightGBM", f"{100 * tab['recall']:.1f}%"),
        ("AUC-PR ensamble", f"{ens['auc_pr']:.3f}"),
        ("Costo ensamble", f"Q{ens['costo_q']:,.0f}"),
    ]
    for j, (label, value) in enumerate(cards):
        c = metrics.cell(0, j)
        shade(c, "184E77")
        margins(c, 80, 90, 60, 90)
        add_text(c, label.upper(), 7, True, "F8FBFF")
        c = metrics.cell(1, j)
        shade(c, "EDF7F6")
        margins(c, 80, 90, 80, 90)
        p = add_text(c, value, 15, True, "184E77")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    rr = p.add_run("CONTENIDO Y REPRODUCCIÓN")
    rr.bold = True
    rr.font.size = Pt(10)
    rr.font.color.rgb = RGBColor(24, 78, 119)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    bullets = [
        "V1 preservada y V2 separada en artefactos, código y figuras.",
        "Notebook ejecutable, informe LaTeX/PDF y presentación HTML de 8 diapositivas.",
        "Auditoría de correlación, información mutua, redundancia y PCA.",
        "Calibración, umbral económico, top-k, deriva, segmentos e intervalos.",
        "Dependencias exactas e instrucciones PowerShell en configuracion/v2.",
        "Fuente única: artefactos/v2/resultados_v2.json.",
    ]
    for j in range(2):
        c = table.cell(0, j)
        margins(c)
        shade(c, "F8FBFF")
        for item in bullets[j * 3 : (j + 1) * 3]:
            p = c.add_paragraph(item, style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.runs[0].font.size = Pt(8.2)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    rr = p.add_run("USO RESPONSABLE")
    rr.bold = True
    rr.font.size = Pt(9)
    rr.font.color.rgb = RGBColor(218, 118, 35)
    p = doc.add_paragraph(
        "Prototipo académico para priorizar revisión humana. No debe bloquear operaciones ni atribuir culpabilidad. Requiere una cohorte nueva, privacidad, explicabilidad, seguridad, análisis de sesgo, costos reales y monitoreo antes de producción."
    )
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph(URL)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(24, 78, 119)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    libre = shutil.which("libreoffice") or shutil.which("soffice")
    if libre:
        subprocess.run(
            [
                libre,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(OUT.parent),
                str(OUT),
            ],
            check=False,
            capture_output=True,
        )
    print("Ficha V2 creada:", OUT)


if __name__ == "__main__":
    main()
