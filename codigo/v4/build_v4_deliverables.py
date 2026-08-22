from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path

import nbformat as nbf
import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[2]
RESULT_PATH = ROOT / "artefactos" / "v4" / "resultados_v4.json"
R = json.loads(RESULT_PATH.read_text(encoding="utf-8"))
REC = R["modelo_v4_recomendado"]
POLICY = REC["politicas"][REC["politica_recomendada"]]
EVAL = POLICY["evaluacion"]
BENCH = POLICY["benchmark_historico"]
V3E = REC["comparacion_v3_evaluacion"]["v3"]
DELTA = REC["comparacion_v3_evaluacion"]["deltas_v4_menos_v3"]
URL = "https://github.com/DanielBarillasM/Proyecto-1_Grupo-1_DLYSI_Sec-30"


def f(value: float, digits: int = 3) -> str:
    return f"{float(value):,.{digits}f}"


def pct(value: float) -> str:
    return f"{100 * float(value):.2f}%"


def tex_pct(value: float) -> str:
    return f"{100 * float(value):.2f}\\%"


def money(value: float) -> str:
    return f"Q{float(value):,.0f}"


def build_readme() -> None:
    walk_rows = "\n".join(
        f"| {row['modelo']} | {f(row['mean'],4)} | {f(row['std'],4)} |"
        for row in sorted(R["validacion_walk_forward"]["resumen"], key=lambda x: x["mean"], reverse=True)
    )
    candidate_rows = "\n".join(
        f"| {row['modelo']} | {f(row['auc_pr_seleccion'],4)} | {f(row['roc_auc_seleccion'],4)} |"
        for row in REC["seleccion"]
    )
    policy_rows = "\n".join(
        f"| {name} | {f(data['threshold'],5)} | {pct(data['evaluacion']['precision'])} | {pct(data['evaluacion']['recall'])} | {f(data['evaluacion']['f1'])} | {money(data['evaluacion']['costo_q'])} |"
        for name, data in REC["politicas"].items()
    )
    text = rf"""<div align="center">

# Proyecto 1 · Monitoreo transaccional — V4

### Detección de fraude con ingeniería causal, expertos segmentados y evaluación temporal

![Python](https://img.shields.io/badge/Python-3.13-3776AB?logo=python&logoColor=white)
![LightGBM](https://img.shields.io/badge/LightGBM-4.7.0-184e77)
![XGBoost](https://img.shields.io/badge/XGBoost-3.4.1-e76f51)
![Estado](https://img.shields.io/badge/V4-CANDIDATO_CONGELADO-e9c46a)
![ROC--AUC](https://img.shields.io/badge/ROC--AUC-0.9013-2a9d8f)

**Wilson Alejandro Calderón Argueta · 22018** · **Pablo Daniel Barillas Moreno · 22193**

Universidad del Valle de Guatemala · Deep Learning y Sistemas Inteligentes · Sección 30 · 2026

</div>

> [!IMPORTANT]
> V4 mejora todas las métricas operativas frente a V3 bajo la política robusta de recall, pero esa política se definió después de observar la primera evaluación V4. Por integridad experimental se conserva V3 como versión confirmada y se congela V4 para evaluarla, sin cambios, en una cohorte temporal nueva. El último 15 % de IEEE-CIS continúa siendo un benchmark histórico reutilizado, no un test ciego.

En consecuencia, **V4 es un candidato congelado**, no una promoción confirmatoria ni un sistema listo para producción.

## Contenido

- [Resumen ejecutivo](#resumen-ejecutivo)
- [Problema, datos y objetivo correcto](#problema-datos-y-objetivo-correcto)
- [Qué cambia en V4](#qué-cambia-en-v4)
- [Protocolo temporal](#protocolo-temporal)
- [Modelos y selección](#modelos-y-selección)
- [Resultados](#resultados)
- [Políticas de umbral](#políticas-de-umbral)
- [Correlación, PCA y deriva](#correlación-pca-y-deriva)
- [Estructura y reproducción](#estructura-y-reproducción)
- [Limitaciones y uso responsable](#limitaciones-y-uso-responsable)
- [Referencias APA 7](#referencias-apa-7)

## Resumen ejecutivo

El proyecto estudia {R['datos']['filas']:,} transacciones del conjunto IEEE-CIS Fraud Detection, con prevalencia de fraude de {pct(R['datos']['prevalencia'])}. La V3 ya había demostrado que ampliar características y usar categorías nativas era más efectivo que aumentar una GRU cuya falsificación temporal mostraba poca señal de orden. V4 profundiza esa conclusión: integra {R['datos']['columnas_integradas']} columnas después de ingeniería, selecciona {R['seleccion_variables']['n_numericas']} variables numéricas y {R['seleccion_variables']['n_categoricas']} categóricas, optimiza LightGBM con {R['optuna']['trials']} pruebas, compara CatBoost y XGBoost, construye una variante con hard negatives y entrena expertos separados para `ProductCD=W` y el resto.

LightGBM V4 alcanza AUC-PR walk-forward media **{f(next(x['mean'] for x in R['validacion_walk_forward']['resumen'] if x['modelo']=='LightGBM_tuned'),4)}**, frente a **{f(R['promocion']['v3_auc_pr_walk'],4)}** en V3. La diferencia es **+{f(R['promocion']['delta_auc_pr_walk'],4)}** y la ROC-AUC media llega a 0.9220. En el bloque final separado, el candidato seleccionado antes de calibrar y fijar umbral es `LightGBM_expertos_ProductCD`, con AUC-PR **{f(EVAL['auc_pr'],4)}** y ROC-AUC **{f(EVAL['roc_auc'],4)}**.

La política robusta `recall ≥ 0.75`, fijada sobre un bloque anterior, produce precisión **{pct(EVAL['precision'])}**, recall **{pct(EVAL['recall'])}**, F1 **{f(EVAL['f1'],4)}** y costo **{money(EVAL['costo_q'])}**. Frente a V3, mejora AUC-PR {f(DELTA['auc_pr'],4)}, ROC-AUC {f(DELTA['roc_auc'],4)}, precisión {pct(DELTA['precision'])}, recall {pct(DELTA['recall'])}, F1 {f(DELTA['f1'],4)} y reduce el costo {pct(DELTA['reduccion_costo'])}. La mejora pareada de AUC-PR tiene IC 95 % [{f(REC['comparacion_pareada_evaluacion']['li95'],4)}, {f(REC['comparacion_pareada_evaluacion']['ls95'],4)}].

## Problema, datos y objetivo correcto

El fraude es una clasificación binaria extremadamente desbalanceada. Una accuracy alta no demuestra utilidad: predecir todas las transacciones como legítimas superaría 96 % de accuracy y detectaría cero fraudes. Por ello el proyecto prioriza AUC-PR, precisión, recall, F1, costo y métricas `Precision@K/Recall@K`. ROC-AUC se reporta como medida complementaria y V4 supera el objetivo 0.90, pero no se presenta como sustituto de la curva Precision–Recall.

Las tablas `train_transaction.csv` y `train_identity.csv` se unen únicamente por `TransactionID`. Las observaciones se ordenan por `TransactionDT`; ambos campos se excluyen como magnitudes predictivas. Para una operación en tiempo $t$, toda variable histórica cumple:

$$
x_t^{{hist}}=f(\{{x_j:t_j<t\}}).
$$

El falso negativo tiene costo académico Q4,200 y el falso positivo Q180. La relación 23.3:1 explica por qué la política recomendada exige un buffer de recall en vez de maximizar exclusivamente precisión.

## Qué cambia en V4

V4 incorpora 135 variables derivadas adicionales, entre ellas calendario, forma decimal y primer dígito del monto, resúmenes por bloques `V/C/D`, normalizaciones `D_i-día`, logaritmos `C_i`, frecuencias estrictamente previas, cambios de dispositivo/dirección/correo, familias de navegador y dispositivo, y estadísticas históricas para cuatro identidades proxy: tarjeta–dirección, tarjeta–correo, tarjeta–dispositivo y tarjeta–producto.

No se introduce un ID como variable continua. Las claves proxy solo sirven para producir conteos, montos previos, dispersión, razones y recencia. Los valores faltantes se preservan, porque el patrón de ausencia puede tener señal, pero también se documenta que esa señal puede cambiar con el proceso de captura.

La selección utiliza exclusivamente el 55 % cronológico inicial. Se eliminan constantes, columnas presentes en menos de 0.5 % de ese tramo y sustitutos casi idénticos con $|\rho_s|\ge0.999$. La baja correlación marginal con fraude no obliga a excluir una característica: LightGBM puede aprovechar interacciones no lineales. PCA permanece como ablation lineal; en V3 conservó 99.53 % de la varianza y redujo AUC-PR, evidencia de que varianza total no equivale a información discriminativa.

## Protocolo temporal

El orden experimental es 70 % entrenamiento, 15 % validación y 15 % benchmark histórico reutilizado. Dentro de validación, V4 separa bloques para early stopping, entrenamiento del stacking, selección de candidato, calibración, selección de umbral y evaluación. Los tres folds walk-forward nunca entrenan con eventos posteriores al bloque evaluado.

```mermaid
flowchart LR
  A[IEEE-CIS ordenado] --> B[Features causales]
  B --> C[Selección en 55% inicial]
  C --> D[Optuna + walk-forward]
  D --> E[Modelos globales y expertos]
  E --> F[Selección de candidato]
  F --> G[Calibración]
  G --> H[Umbral]
  H --> I[Evaluación final]
  I --> J[Cohorte nueva requerida]
```

El modelo adversarial distingue train de validación con ROC-AUC 1.0, dominado por `day_index` y otras variables temporales acumulativas. No significa que el clasificador de fraude sea perfecto: demuestra una deriva temporal fuerte y justifica ponderación por recencia, walk-forward y monitoreo por períodos.

## Modelos y selección

| Modelo | AUC-PR walk-forward | Desviación |
|---|---:|---:|
{walk_rows}

LightGBM fue optimizado con Optuna sobre learning rate, hojas, profundidad, mínimo por hoja, submuestreo, regularización y suavizado categórico. XGBoost y CatBoost se ejecutaron con presupuestos piloto uniformes de 220 y 120 iteraciones respectivamente, después de comprobar que una configuración CatBoost mayor excedía 20 minutos por fold. Por tanto, son controles de diversidad y no búsquedas exhaustivas.

La selección final ocurre en el bloque 50–60 % de validación, antes de calibración, umbral y evaluación:

| Candidato | AUC-PR selección | ROC-AUC selección |
|---|---:|---:|
{candidate_rows}

El experto por `ProductCD` gana con AUC-PR {f(REC['seleccion'][0]['auc_pr_seleccion'],4)}. El stacking experimental queda segundo y no se recomienda: en evaluación redujo ROC-AUC a 0.8686, señal de sobreajuste del metamodelo. Esta conclusión negativa se conserva porque evita presentar complejidad adicional como mejora automática.

## Resultados

| Modelo/política | AUC-PR | ROC-AUC | Precisión | Recall | F1 | Costo |
|---|---:|---:|---:|---:|---:|---:|
| V3, bloque comparable | {f(V3E['auc_pr'])} | {f(V3E['roc_auc'])} | {pct(V3E['precision'])} | {pct(V3E['recall'])} | {f(V3E['f1'])} | {money(V3E['costo_q'])} |
| **V4 experto, política robusta** | **{f(EVAL['auc_pr'])}** | **{f(EVAL['roc_auc'])}** | **{pct(EVAL['precision'])}** | **{pct(EVAL['recall'])}** | **{f(EVAL['f1'])}** | **{money(EVAL['costo_q'])}** |
| V4, benchmark histórico | {f(BENCH['auc_pr'])} | {f(BENCH['roc_auc'])} | {pct(BENCH['precision'])} | {pct(BENCH['recall'])} | {f(BENCH['f1'])} | {money(BENCH['costo_q'])} |

![Walk-forward V4](../evidencia/figuras/v4/01_walk_forward_v4.png)

![Curvas PR](../evidencia/figuras/v4/07_curvas_pr_candidato_v4.png)

En el benchmark histórico V4 logra AUC-PR {f(BENCH['auc_pr'],4)} y ROC-AUC {f(BENCH['roc_auc'],4)}. La diferencia pareada de AUC-PR V4–V3 es {f(REC['comparacion_pareada_benchmark_historico']['delta_auc_pr'],4)}, IC descriptivo [{f(REC['comparacion_pareada_benchmark_historico']['li95'],4)}, {f(REC['comparacion_pareada_benchmark_historico']['ls95'],4)}. Este resultado apoya consistencia, pero no constituye confirmación ciega.

## Políticas de umbral

| Política | Umbral | Precisión | Recall | F1 | Costo evaluación |
|---|---:|---:|---:|---:|---:|
{policy_rows}

La política balanceada maximiza F1 con recall mínimo 0.70 en el bloque de umbral. La robusta exige 0.75 para absorber deriva y es la recomendada para la siguiente cohorte. La económica minimiza $4200FN+180FP$ y recupera más fraude, pero reduce precisión y F1. Ningún umbral modifica AUC-PR o ROC-AUC; únicamente el punto operativo.

En capacidad limitada, V4 obtiene precisión superior a 0.90 al revisar el 1 % de mayor riesgo y recupera aproximadamente 26 % del fraude. Esto sí es una métrica ≥0.90 válida, pero debe expresarse como `Precision@1%`, no como precisión poblacional.

## Correlación, PCA y deriva

La poda de correlación se usa para redundancia, no para confundir correlación con causalidad o relevancia. Se conserva una sola variable cuando dos candidatas tienen Spearman absoluto ≥0.999. Para árboles, PCA no es la representación principal porque destruye interpretabilidad y puede diluir señal minoritaria. Para modelos lineales continúa siendo un control válido.

La deriva adversarial perfecta exige cautela. Variables temporales explícitas permiten distinguir períodos casi sin error, y las frecuencias acumulativas cambian necesariamente con el tiempo. En una siguiente iteración deben reportarse dos auditorías: deriva total operacional y deriva residual excluyendo marcadores temporales explícitos. También se debe vigilar estabilidad de umbral, prevalencia, PSI, faltantes y métricas por producto/dispositivo.

## Estructura y reproducción

```text
codigo/                  pipelines y auditorías V3/V4
configuracion/v4/        dependencias e instrucciones V4
datos/raw/               IEEE-CIS local, ignorado por Git
datos/processed/v4/      selección y correlación
artefactos/v4/           modelos, predicciones y resultados
evidencia/figuras/v4/    gráficos reproducibles
entregables/cuaderno/    notebook ejecutable
entregables/informe/     LaTeX y PDF
entregables/presentacion/HTML y PDF de 8 diapositivas
entregables/ficha/       ficha DOCX/PDF del repositorio
```

```powershell
py -3.13 -m venv .venv
.\.venv\Scripts\Activate.ps1
python -m pip install -r configuracion/v4/requirements-v4.txt
python -m pip install -r configuracion/v4/requirements-docs-v4.txt
python codigo/compartido/download_data.py
python -u codigo/v4/proyecto1_v4_pipeline.py
python codigo/v4/postprocess_v4.py
python codigo/v4/build_v4_deliverables.py
jupyter nbconvert --to notebook --execute --inplace --ExecutePreprocessor.timeout=300 entregables/cuaderno/v4/Proyecto_1_Monitoreo_Transaccional_V4.ipynb
python codigo/v4/audit_project1_v4.py
```

La fuente única es [`artefactos/v4/resultados_v4.json`](../artefactos/v4/resultados_v4.json). Las instrucciones completas están en [`configuracion/v4/INSTRUCCIONES_V4.md`](../configuracion/v4/INSTRUCCIONES_V4.md).

## Limitaciones y uso responsable

V4 no identifica personas reales: sus identidades son claves proxy que pueden mezclar usuarios o fragmentar uno mismo. IEEE-CIS está anonimizado, cubre un período limitado y no representa automáticamente otro país, comercio o época. Los costos son supuestos académicos. No se estudiaron consecuencias legales, privacidad, seguridad adversarial, explicaciones individuales, latencia productiva ni sesgos sobre atributos no disponibles.

El modelo debe priorizar revisión humana, nunca atribuir culpabilidad ni bloquear transacciones de forma autónoma. Antes de producción se requieren una nueva cohorte, calibración con costos reales, análisis de impacto, trazabilidad de decisiones, controles de acceso, monitoreo y un procedimiento de apelación. La política robusta se considera post-hoc; debe congelarse sin ajustes y validarse una sola vez en datos nuevos.

## Referencias APA 7

Akiba, T., Sano, S., Yanase, T., Ohta, T., & Koyama, M. (2019). Optuna: A next-generation hyperparameter optimization framework. *Proceedings of the 25th ACM SIGKDD International Conference on Knowledge Discovery & Data Mining*, 2623–2631. https://doi.org/10.1145/3292500.3330701

Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. *Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining*, 785–794. https://doi.org/10.1145/2939672.2939785

IEEE Computational Intelligence Society. (2019). *IEEE-CIS Fraud Detection* [Data set]. Kaggle. https://www.kaggle.com/competitions/ieee-fraud-detection

Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T.-Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. *Advances in Neural Information Processing Systems, 30*.

Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A. V., & Gulin, A. (2018). CatBoost: Unbiased boosting with categorical features. *Advances in Neural Information Processing Systems, 31*.

Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. *PLOS ONE, 10*(3), e0118432. https://doi.org/10.1371/journal.pone.0118432

## Declaración de asistencia

Se utilizó asistencia de IA para desarrollo, redacción, visualización y auditoría. Los integrantes ejecutaron los experimentos, revisaron los artefactos y asumen responsabilidad por las decisiones y conclusiones.
"""
    out = ROOT / "entregables" / "ficha" / "v4" / "README_V4.md"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(text, encoding="utf-8", newline="\n")


def notebook_style() -> str:
    return """<style>
    :root{--navy:#102a43;--blue:#184e77;--teal:#2a9d8f;--gold:#e9c46a;--ink:#172033;--pale:#edf5fb}
    .hero{padding:38px 42px;border-radius:24px;color:#f8fbff;background:radial-gradient(circle at 92% 8%,rgba(255,255,255,.16) 0 8%,transparent 9%),linear-gradient(125deg,var(--navy),var(--blue) 55%,var(--teal));box-shadow:0 16px 38px #102a433d;font-family:Inter,'Segoe UI',sans-serif}.chips{display:flex;gap:8px;flex-wrap:wrap}.chips span{padding:6px 13px;border:1px solid #ffffff55;border-radius:999px;background:#ffffff20;font-size:11px;font-weight:800}.hero h1{font-size:38px;margin:18px 0 8px;color:white}.grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(210px,1fr));gap:13px;margin-top:24px}.card{padding:14px 16px;border:1px solid #ffffff44;border-radius:12px;background:#ffffff12}.section{margin:28px 0 14px;padding:16px 22px;border-radius:14px;background:linear-gradient(90deg,var(--navy),var(--blue));color:white}.section h2{margin:0;color:white}.call{padding:18px 22px;margin:14px 0;border:1px solid #c9d9e6;border-left:6px solid var(--teal);border-radius:13px;background:var(--pale);color:var(--ink);line-height:1.7}.warn{border-left-color:#e76f51;background:#fff2ed}.metric{display:inline-block;padding:10px 14px;margin:4px;border-radius:10px;background:#e5f5f2;color:#184e77;font-weight:800}table{font-size:14px}code{background:#102a4312;padding:2px 5px;border-radius:4px}</style>"""


def build_notebook() -> None:
    cells = []
    cells.append(nbf.v4.new_markdown_cell(notebook_style() + f"""<div class="hero"><div class="chips"><span>DEEP LEARNING</span><span>PROYECTO 1</span><span>V4 CANDIDATO</span><span>GRUPO 1 · SECCIÓN 30</span></div><h1>Monitoreo transaccional y detección de fraude</h1><p style="font-size:20px">Ingeniería causal, expertos segmentados y validación temporal</p><div class="grid"><div class="card"><b>Institución</b><br>Universidad del Valle de Guatemala</div><div class="card"><b>Curso</b><br>Deep Learning y Sistemas Inteligentes</div><div class="card"><b>Docente</b><br>Kevin Recinos</div><div class="card"><b>Integrantes</b><br>Wilson Calderón · 22018<br>Pablo Barillas · 22193</div><div class="card"><b>Dataset</b><br>IEEE-CIS Fraud Detection</div><div class="card"><b>Candidato</b><br>LightGBM expertos ProductCD</div></div></div>"""))
    cells.append(nbf.v4.new_code_cell("from pathlib import Path\nimport json, pandas as pd\nfrom IPython.display import Image, display\nROOT=Path.cwd().resolve()\nwhile not (ROOT/'artefactos').exists() and ROOT != ROOT.parent:\n    ROOT=ROOT.parent\nR=json.loads((ROOT/'artefactos/v4/resultados_v4.json').read_text(encoding='utf-8'))\nREC=R['modelo_v4_recomendado']; POLICY=REC['politicas'][REC['politica_recomendada']]\nR['recomendacion']"))
    sections = [
        ("1 · Resumen ejecutivo", f"""<div class="call"><b>Resultado:</b> V4 eleva AUC-PR walk-forward a {f(next(x['mean'] for x in R['validacion_walk_forward']['resumen'] if x['modelo']=='LightGBM_tuned'),4)} y la ROC-AUC media a 0.9220. El candidato experto obtiene ROC-AUC {f(EVAL['roc_auc'],4)} en evaluación y mejora todas las métricas operativas con la política robusta.</div><div class="call warn"><b>Decisión científica:</b> conservar V3 como confirmada y congelar V4 para una cohorte nueva, porque el piso robusto de recall se añadió post-hoc.</div>""", None),
        ("2 · Pregunta y objetivos", """El objetivo no es maximizar accuracy, sino ordenar correctamente el fraude, mantener recall operativo, reducir falsos positivos y minimizar $C=4200FN+180FP$. Se reportan AUC-PR, ROC-AUC, precisión, recall, F1, calibración, top-K y costo.""", None),
        ("3 · Datos y causalidad", r"""Cada característica histórica de una transacción $t$ usa exclusivamente eventos anteriores: $$x_t^{hist}=f(\{x_j:t_j<t\}).$$ `TransactionID` se utiliza para unir tablas y `TransactionDT` para ordenar; ninguno entra como predictor directo.""", "pd.DataFrame(R['datos']['particiones']).T"),
        ("4 · Ingeniería V4", """V4 añade agregados para tarjeta–dirección, tarjeta–correo, tarjeta–dispositivo y tarjeta–producto; ventanas de actividad, forma del monto, normalización temporal, patrones de ausencia, frecuencia previa y cambios de contexto.""", "pd.Series({'columnas integradas':R['datos']['columnas_integradas'],'numéricas':R['seleccion_variables']['n_numericas'],'categóricas':R['seleccion_variables']['n_categoricas'],'variables derivadas':R['identidad_y_features']['variables_ingenieria_v4']}).to_frame('valor')"),
        ("5 · Correlación, IDs y PCA", r"""Se excluyen IDs como magnitudes y se poda solo redundancia casi exacta con $|\rho_s|\ge0.999$. Una correlación marginal baja no descarta interacciones no lineales. PCA permanece como ablation: en V3 conservó 99.53% de varianza y perdió AUC-PR.""", None),
        ("6 · Modelos", """Se comparan LightGBM optimizado, XGBoost y CatBoost piloto, hard-negative mining, expertos `W/NO_W` y stacking. CatBoost/XGBoost usan presupuestos acotados después de auditar su tiempo en CPU.""", None),
        ("7 · Walk-forward", """Los tres folds entrenan solo con el pasado. LightGBM V4 aumenta AUC-PR media en 0.0366 frente a V3 y alcanza ROC-AUC media 0.9220.""", "pd.DataFrame(R['validacion_walk_forward']['detalle']).sort_values(['fold','auc_pr'],ascending=[True,False])"),
        ("8 · Evidencia walk-forward", """La mejora se sostiene especialmente en F3, el corte más cercano al final y más sensible a deriva.""", "display(Image(filename=str(ROOT/'evidencia/figuras/v4/01_walk_forward_v4.png')))"),
        ("9 · Selección del candidato", """La selección usa 50–60% de validación, antes de calibración, umbral y evaluación. El experto por ProductCD gana; el stacking queda segundo y pierde ROC-AUC en evaluación.""", "pd.DataFrame(REC['seleccion'])"),
        ("10 · Calibración", """La calibración sigmoide del candidato reduce Brier y ECE en su bloque reservado. Calibrar no cambia AUC-PR/ROC-AUC; mejora la interpretación probabilística y la estabilidad del umbral.""", "pd.Series(REC['calibracion']).to_frame('valor')"),
        ("11 · Políticas de umbral", """Se conservan tres políticas. La robusta exige recall ≥0.75 en selección; es útil por el costo alto de FN, pero se etiqueta post-hoc.""", "pd.DataFrame({k:v['evaluacion'] for k,v in REC['politicas'].items()}).T[['threshold','precision','recall','f1','costo_q','alertas_por_100k']]"),
        ("12 · Comparación V3–V4", """En el bloque final comparable, la política robusta domina AUC-PR, ROC-AUC, precisión, recall, F1 y costo. El IC pareado de AUC-PR excluye cero.""", "pd.DataFrame({'V3':REC['comparacion_v3_evaluacion']['v3'],'V4':POLICY['evaluacion']}).T[['auc_pr','roc_auc','precision','recall','f1','costo_q']]"),
        ("13 · Benchmark histórico", """El benchmark reutilizado confirma consistencia descriptiva: AUC-PR 0.5592 y ROC-AUC 0.9013. No decide selección ni promoción.""", "pd.Series(POLICY['benchmark_historico']).to_frame('valor')"),
        ("14 · Curvas PR", """V4 mejora el ranking frente a V3 en el benchmark histórico, aunque la confirmación exige datos nuevos.""", "display(Image(filename=str(ROOT/'evidencia/figuras/v4/07_curvas_pr_candidato_v4.png')))"),
        ("15 · Capacidad top-K", """Al revisar 1% de mayor riesgo, precision@K supera 0.90. Esta es la afirmación correcta de una métrica ≥0.90; no equivale a precisión poblacional.""", "pd.DataFrame(R['metricas_top_k_candidato'])"),
        ("16 · Deriva", """La validación adversarial obtiene ROC-AUC 1.0, dominada por tiempo y frecuencias acumulativas. Esto evidencia separación temporal y obliga a monitorear deriva, no perfección del detector.""", "pd.DataFrame(R['deriva_adversarial']['top_variables'])"),
        ("17 · Qué no funcionó", """CatBoost piloto fue inestable; XGBoost no superó LightGBM; el stacking redujo ROC-AUC; PCA perdió señal en V3. Conservar resultados negativos evita justificar complejidad sin beneficio.""", None),
        ("18 · Decisión", """V4 queda congelada como candidato superior post-hoc. V3 permanece confirmada hasta evaluar V4 sin cambios en una cohorte temporal nueva. No se ajustarán variables, modelos o umbrales después de ver esa cohorte.""", "R['decision_v4']"),
        ("19 · Uso responsable", """El sistema prioriza revisión humana. No determina culpabilidad ni bloquea operaciones automáticamente. Antes de producción se requieren privacidad, seguridad, explicabilidad, monitoreo, costos reales y apelación.""", None),
        ("20 · Reproducibilidad", """La fuente única es `artefactos/v4/resultados_v4.json`. El pipeline, postprocesamiento, construcción y auditoría están separados para conservar trazabilidad.""", None),
    ]
    for title, markdown, code in sections:
        cells.append(nbf.v4.new_markdown_cell(f'<div class="section"><h2>{title}</h2></div>\n\n{markdown}'))
        if code:
            cells.append(nbf.v4.new_code_cell(code))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>Referencias APA 7</h2></div>

Akiba, T., Sano, S., Yanase, T., Ohta, T., & Koyama, M. (2019). Optuna: A next-generation hyperparameter optimization framework. *KDD*, 2623–2631.

Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. *KDD*, 785–794.

IEEE Computational Intelligence Society. (2019). *IEEE-CIS Fraud Detection* [Conjunto de datos]. Kaggle.

Ke, G., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. *NeurIPS, 30*.

Prokhorenkova, L., et al. (2018). CatBoost: Unbiased boosting with categorical features. *NeurIPS, 31*.

Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating imbalanced classifiers. *PLOS ONE, 10*(3), e0118432.

**Declaración de IA.** Se utilizó asistencia para código, redacción, visualización y auditoría. Los autores ejecutaron y verificaron los resultados y asumen responsabilidad por sus conclusiones."""))
    notebook = nbf.v4.new_notebook(cells=cells, metadata={"kernelspec":{"display_name":"Python 3","language":"python","name":"python3"},"language_info":{"name":"python","version":"3.13"}})
    nbf.write(notebook, ROOT / "entregables" / "cuaderno" / "v4" / "Proyecto_1_Monitoreo_Transaccional_V4.ipynb")


def build_report() -> Path:
    walk = next(x for x in R["validacion_walk_forward"]["resumen"] if x["modelo"] == "LightGBM_tuned")
    tex = rf"""\documentclass[10pt]{{article}}
\usepackage[utf8]{{inputenc}}\usepackage[T1]{{fontenc}}\usepackage[letterpaper,margin=1.55cm]{{geometry}}
\usepackage{{graphicx,booktabs,xcolor,amsmath,hyperref,array}}\definecolor{{navy}}{{HTML}}{{102A43}}\definecolor{{teal}}{{HTML}}{{2A9D8F}}
\hypersetup{{colorlinks=true,urlcolor=teal}}\setlength{{\parindent}}{{0pt}}\setlength{{\parskip}}{{5pt}}
\begin{{document}}\begin{{titlepage}}\pagecolor{{navy}}\color{{white}}\raggedright\vspace*{{1cm}}{{\Large DEEP LEARNING Y SISTEMAS INTELIGENTES\par}}\vspace{{1cm}}{{\Huge\bfseries Proyecto 1\\Monitoreo transaccional --- V4\par}}\vspace{{.7cm}}{{\LARGE Ingeniería causal, expertos y validación temporal\par}}\vfill{{\Large Universidad del Valle de Guatemala\\Kevin Recinos · Grupo 1 · Sección 30\\Semestre II 2026\par}}\vspace{{1cm}}{{\Large Wilson Alejandro Calderón Argueta · 22018\\Pablo Daniel Barillas Moreno · 22193\par}}\vfill\textbf{{Estado:}} candidato congelado post-hoc; V3 permanece confirmada hasta una cohorte nueva.\end{{titlepage}}\nopagecolor

\section*{{Resumen ejecutivo}} Se analizaron {R['datos']['filas']:,} transacciones IEEE-CIS con prevalencia {tex_pct(R['datos']['prevalencia'])}. V4 integra {R['datos']['columnas_integradas']} columnas tras ingeniería y selecciona 360 numéricas más 38 categóricas. LightGBM optimizado alcanza AP walk-forward {walk['mean']:.4f} y ROC-AUC media 0.9220. El candidato experto por ProductCD, seleccionado antes de calibración y evaluación, obtiene AP {EVAL['auc_pr']:.4f}, ROC-AUC {EVAL['roc_auc']:.4f}, precisión {EVAL['precision']:.4f}, recall {EVAL['recall']:.4f}, F1 {EVAL['f1']:.4f} y costo {money(EVAL['costo_q'])}. Mejora todas las métricas frente a V3, pero la política recall $\ge0.75$ es post-hoc; requiere cohorte nueva.

\section{{Problema y datos}} El fraude representa {tex_pct(R['datos']['prevalencia'])}; accuracy convencional sería engañosa. Se priorizan AUC-PR, precisión, recall, F1, costo y top-K. Las tablas se unen con TransactionID y ordenan con TransactionDT; ninguno entra como magnitud predictiva. Para cada transacción, $x_t^{{hist}}=f(\{{x_j:t_j<t\}})$. Se utiliza 70\% train, 15\% validación y 15\% benchmark histórico reutilizado.

\section{{Ingeniería causal}} V4 añade calendario, forma del monto, resúmenes V/C/D, $D_i-día$, $\log(1+C_i)$, ausencias, frecuencias previas, cambios de contexto y estadísticas para tarjeta--dirección, tarjeta--correo, tarjeta--dispositivo y tarjeta--producto. IDs no se usan como números continuos. Selección y correlación se aprenden en el 55\% inicial; solo se podan sustitutos con $|\rho_s|\ge0.999$. PCA no se aplica al boosting porque en V3 perdió AP aun conservando 99.53\% de varianza.
\newpage
\section{{Modelos y protocolo}} Optuna ejecuta 18 pruebas sobre hojas, profundidad, regularización, submuestreo y suavizado categórico. Se comparan LightGBM, XGBoost y CatBoost piloto, hard-negative mining, expertos W/NO-W y stacking. CatBoost y XGBoost usan 120/220 iteraciones por costo de CPU; son ablations, no búsquedas exhaustivas. La validación se divide para early stopping, meta, selección de candidato, calibración, umbral y evaluación.

\begin{{figure}}[h]\centering\includegraphics[width=.78\linewidth]{{../../../evidencia/figuras/v4/01_walk_forward_v4.png}}\caption{{AUC-PR walk-forward de candidatos V4.}}\end{{figure}}

LightGBM obtiene AP media {walk['mean']:.4f} (DE {walk['std']:.4f}), mejora V3 en {R['promocion']['delta_auc_pr_walk']:.4f} y supera 0.90 de ROC-AUC media. XGBoost obtiene 0.5957 y CatBoost piloto 0.5302. La ventaja de LightGBM aumenta cerca del período final.
\newpage
\section{{Selección y resultados}} El bloque 50--60\% de validación selecciona el candidato antes de calibración y umbral. Expertos ProductCD ganan con AP {REC['seleccion'][0]['auc_pr_seleccion']:.4f}; stacking obtiene {REC['seleccion'][1]['auc_pr_seleccion']:.4f}, pero reduce ROC-AUC de evaluación a 0.8686 y se descarta.

\begin{{center}}\begin{{tabular}}{{lrrrrrr}}\toprule Modelo&AUC-PR&ROC-AUC&Prec.&Recall&F1&Costo\\\midrule V3&{V3E['auc_pr']:.3f}&{V3E['roc_auc']:.3f}&{V3E['precision']:.3f}&{V3E['recall']:.3f}&{V3E['f1']:.3f}&{money(V3E['costo_q'])}\\\textbf{{V4}}&\textbf{{{EVAL['auc_pr']:.3f}}}&\textbf{{{EVAL['roc_auc']:.3f}}}&\textbf{{{EVAL['precision']:.3f}}}&\textbf{{{EVAL['recall']:.3f}}}&\textbf{{{EVAL['f1']:.3f}}}&\textbf{{{money(EVAL['costo_q'])}}}\\\bottomrule\end{{tabular}}\end{{center}}

V4 aumenta AP {DELTA['auc_pr']:.4f}, ROC-AUC {DELTA['roc_auc']:.4f}, precisión {DELTA['precision']:.4f}, recall {DELTA['recall']:.4f} y F1 {DELTA['f1']:.4f}; reduce costo {100*DELTA['reduccion_costo']:.1f}\%. La diferencia pareada AP es {REC['comparacion_pareada_evaluacion']['delta_auc_pr']:.4f}, IC 95\% [{REC['comparacion_pareada_evaluacion']['li95']:.4f}, {REC['comparacion_pareada_evaluacion']['ls95']:.4f}].

\begin{{figure}}[h]\centering\includegraphics[width=.70\linewidth]{{../../../evidencia/figuras/v4/07_curvas_pr_candidato_v4.png}}\caption{{Curvas PR V3/V4 en benchmark histórico reutilizado.}}\end{{figure}}
\newpage
\section{{Calibración, umbral y capacidad}} La calibración del candidato reduce Brier de {REC['calibracion']['brier_raw']:.4f} a {REC['calibracion']['brier_calibrado']:.4f} y ECE de {REC['calibracion']['ece_raw']:.4f} a {REC['calibracion']['ece_calibrado']:.4f}. El umbral robusto {POLICY['threshold']:.5f} exige recall 0.75 en selección para amortiguar deriva y el costo FN/FP de 23.3:1.

La política balanceada prioriza F1; la económica eleva recall a {REC['politicas']['economico']['evaluacion']['recall']:.3f}, pero baja precisión. En el benchmark V4 obtiene AP {BENCH['auc_pr']:.4f}, ROC-AUC {BENCH['roc_auc']:.4f}, precisión {BENCH['precision']:.4f}, recall {BENCH['recall']:.4f} y costo {money(BENCH['costo_q'])}. Al revisar 1\% de mayor riesgo, precision@K supera 0.90 y recall@K ronda 0.26.

\begin{{figure}}[h]\centering\includegraphics[width=.68\linewidth]{{../../../evidencia/figuras/v4/06_seleccion_candidato_v4.png}}\caption{{Selección del candidato en bloque independiente.}}\end{{figure}}
\newpage
\section{{Deriva, limitaciones y decisión}} La validación adversarial alcanza ROC-AUC 1.0, dominada por day\_index y acumulados; demuestra separación temporal y justifica walk-forward, no perfección del fraude. Deben separarse deriva total y residual sin tiempo explícito en una cohorte futura.

Limitaciones: benchmark reutilizado, identidad proxy, datos anonimizados, costos académicos, presupuesto piloto CatBoost/XGBoost y ausencia de cohorte externa. V4 no debe bloquear transacciones ni atribuir culpabilidad; prioriza revisión humana. Antes de producción se requieren privacidad, seguridad, explicabilidad, sesgos, costos reales, latencia, monitoreo y apelación.

\textbf{{Decisión:}} V4 domina V3 bajo la política robusta, pero esa política fue añadida después de observar la primera evaluación. Se congela el pipeline, las 398 variables, los modelos, calibrador y umbral; V3 permanece confirmada hasta evaluar V4 una sola vez y sin ajustes en una cohorte temporal nueva.

\section*{{Referencias}}\small Akiba, T., et al. (2019). Optuna. \textit{{KDD}}, 2623--2631.\\Chen, T., \& Guestrin, C. (2016). XGBoost. \textit{{KDD}}, 785--794.\\IEEE CIS. (2019). \textit{{IEEE-CIS Fraud Detection}}. Kaggle.\\Ke, G., et al. (2017). LightGBM. \textit{{NeurIPS, 30}}.\\Prokhorenkova, L., et al. (2018). CatBoost. \textit{{NeurIPS, 31}}.\\Saito, T., \& Rehmsmeier, M. (2015). Precision-recall plots for imbalanced data. \textit{{PLOS ONE, 10}}(3).

\section*{{Declaración de IA}} IA apoyó código, redacción, visualización y auditoría. Los autores ejecutaron, verificaron y asumen responsabilidad.\end{{document}}"""
    out = ROOT / "entregables" / "informe" / "v4" / "informe_proyecto1_v4.tex"
    out.write_text(tex, encoding="utf-8", newline="\n")
    pdflatex = shutil.which("pdflatex")
    if pdflatex:
        for _ in range(2):
            subprocess.run([pdflatex, "-interaction=nonstopmode", "-halt-on-error", "-output-directory", str(out.parent), str(out)], cwd=out.parent, check=False, capture_output=True)
    return out


def build_slides() -> Path:
    walk = next(x for x in R["validacion_walk_forward"]["resumen"] if x["modelo"] == "LightGBM_tuned")
    slides = [
        ("DECISIÓN", "V4 mejora; promoción requiere cohorte nueva", f'<div class="big">{f(EVAL["roc_auc"],3)}</div><p>ROC-AUC evaluación</p><div class="call">AP +{f(DELTA["auc_pr"],3)} · costo −{100*DELTA["reduccion_costo"]:.1f}%</div>', "La política robusta es post-hoc: congelar antes de confirmar."),
        ("DATOS", "590,540 eventos; el futuro queda fuera", '<div class="cols"><div class="box"><b>569</b><p>columnas tras ingeniería</p></div><div class="box"><b>70·15·15</b><p>train, validación, benchmark</p></div></div><p class="formula">xₜʰⁱˢᵗ=f({xⱼ:tⱼ&lt;t})</p>', "TransactionID une; TransactionDT ordena; no son predictores directos."),
        ("SEÑAL", "360 numéricas + 38 categorías", '<div class="road"><span>4 identidades proxy</span><span>frecuencia previa</span><span>cambios de contexto</span><span>montos históricos</span><span>NaN nativo</span></div><div class="call">PCA queda como ablation; correlación poda redundancia ≥0.999.</div>', "Las variables se seleccionan en el 55% inicial."),
        ("WALK-FORWARD", "LightGBM supera V3", '<img src="../../../evidencia/figuras/v4/01_walk_forward_v4.png"><div class="call">V3 0.581 → V4 0.618 AP media · ROC media 0.922</div>', "CatBoost/XGBoost son pilotos de diversidad."),
        ("CANDIDATO", "Los expertos ganan; stacking no", '<img src="../../../evidencia/figuras/v4/06_seleccion_candidato_v4.png"><div class="call">Expertos ProductCD seleccionados antes de calibración y evaluación.</div>', "El stacking perdió ROC-AUC; no vender complejidad como mejora."),
        ("RESULTADOS", "V4 mejora todas las métricas", f'<div class="cols"><div class="box accent"><b>AP {f(EVAL["auc_pr"])}</b><p>ROC {f(EVAL["roc_auc"])}</p><p>Precisión {pct(EVAL["precision"])}</p></div><div class="box"><b>Recall {pct(EVAL["recall"])}</b><p>F1 {f(EVAL["f1"])}</p><p>{money(EVAL["costo_q"])}</p></div></div>', "Bloque final separado; política robusta recall>=0.75."),
        ("OPERACIÓN", "Un ranking, tres políticas", f'<div class="cols"><div class="box"><b>{f(POLICY["threshold"],4)}</b><p>robusto</p></div><div class="box"><b>{pct(next(x["precision_at_k"] for x in R["metricas_top_k_candidato"] if x["tasa_revision"]==0.01))}</b><p>precision@1%</p></div></div><img src="../../../evidencia/figuras/v4/07_curvas_pr_candidato_v4.png">', "El umbral no cambia AUC-PR; cambia carga, recall y costo."),
        ("SIGUIENTE PASO", "Congelar V4 y probar una cohorte nueva", '<div class="road"><span>sin retuning</span><span>deriva residual</span><span>costos reales</span><span>privacidad</span><span>revisión humana</span></div><div class="call">V3 permanece confirmada; V4 es candidato superior post-hoc.</div>', "No declarar test ciego ni producción."),
    ]
    cards = []
    for index, (eye, title, body, notes) in enumerate(slides, 1):
        cards.append(f'<section class="slide"><div class="eye">{eye}</div><h1>{title}</h1><div class="content">{body}</div><footer>Grupo 1 · Sección 30 <b>{index}/8</b></footer><aside>{notes}</aside></section>')
    page = f'''<!doctype html><html lang="es"><head><meta charset="utf-8"><title>Proyecto 1 · V4</title><style>*{{box-sizing:border-box}}body{{margin:0;background:#081a2a;color:#eef7ff;font-family:Inter,"Segoe UI",sans-serif;overflow:hidden}}.slide{{display:none;width:100vw;height:100vh;padding:6vh 7vw;background:radial-gradient(circle at 90% 10%,#2a9d8f55,transparent 30%),linear-gradient(135deg,#102a43,#091b2b);position:relative}}.slide.active{{display:block}}.eye{{color:#65d3c3;font-weight:900;letter-spacing:.16em;font-size:1.2vw}}h1{{font-size:4vw;line-height:1.06;margin:2vh 0 4vh}}p,li{{font-size:1.5vw;line-height:1.42}}.content{{height:68vh}}.big{{font-size:10vw;font-weight:900;color:#65d3c3;line-height:1}}.cols{{display:grid;grid-template-columns:1fr 1fr;gap:2.5vw}}.box{{padding:1.7vw;border:1px solid #ffffff33;border-radius:1.2vw;background:#ffffff0d}}.box b{{font-size:2.8vw;color:#65d3c3}}.accent{{border-color:#65d3c3}}.call{{padding:1.1vw;margin-top:2vh;border-left:.5vw solid #2a9d8f;background:#ffffff12;border-radius:.8vw;font-size:1.4vw}}img{{display:block;max-width:82%;max-height:46vh;margin:auto;border-radius:1vw;background:white}}.formula{{text-align:center;font-size:2.7vw;color:#65d3c3}}.road{{display:flex;flex-wrap:wrap;gap:1vw}}.road span{{padding:1.1vw;border:1px solid #65d3c3;border-radius:999px;background:#184e77;font-size:1.3vw}}footer{{position:absolute;bottom:3vh;left:7vw;right:7vw;display:flex;justify-content:space-between;color:#9eb6c8}}aside{{display:none}}body.notes aside{{display:block;position:absolute;right:2vw;bottom:7vh;width:31vw;padding:1vw;background:#fff;color:#172033;border-radius:.7vw}}@page{{size:13.333in 7.5in;margin:0}}@media print{{html,body{{width:13.333in;height:7.5in;overflow:visible}}.slide{{display:block!important;width:13.333in;height:7.5in;page-break-after:always}}aside{{display:none!important}}}}</style></head><body>{''.join(cards)}<script>const s=[...document.querySelectorAll('.slide')];let i=0;function show(n){{i=Math.max(0,Math.min(s.length-1,n));s.forEach((x,j)=>x.classList.toggle('active',j===i))}}document.onkeydown=e=>{{if(['ArrowRight',' ','PageDown'].includes(e.key))show(i+1);if(['ArrowLeft','PageUp'].includes(e.key))show(i-1);if(e.key.toLowerCase()==='n')document.body.classList.toggle('notes')}};show(0)</script></body></html>'''
    out = ROOT / "entregables" / "presentacion" / "v4" / "presentacion_proyecto1_v4.html"
    out.write_text(page, encoding="utf-8", newline="\n")
    edge_candidates = [shutil.which("msedge"), Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"), Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe")]
    edge = next((str(x) for x in edge_candidates if x and Path(x).exists()), None)
    if edge:
        subprocess.run([edge, "--headless", "--disable-gpu", "--no-pdf-header-footer", f"--print-to-pdf={out.with_suffix('.pdf')}", out.resolve().as_uri()], check=False, capture_output=True)
    return out


def build_ficha() -> None:
    out_dir = ROOT / "entregables" / "ficha" / "v4"
    qr_path = ROOT / "evidencia" / "recursos" / "qr_repositorio_v4.png"
    qr_path.parent.mkdir(parents=True, exist_ok=True)
    qrcode.make(URL).save(qr_path)
    doc = Document(); section = doc.sections[0]
    section.top_margin=Inches(.42); section.bottom_margin=Inches(.42); section.left_margin=Inches(.55); section.right_margin=Inches(.55)
    title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; run=title.add_run("PROYECTO 1 · MONITOREO TRANSACCIONAL · V4"); run.bold=True; run.font.size=Pt(19); run.font.color.rgb=RGBColor(24,78,119)
    sub=doc.add_paragraph("Universidad del Valle de Guatemala · Deep Learning y Sistemas Inteligentes · Sección 30"); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    names=doc.add_paragraph("Wilson Alejandro Calderón Argueta · 22018  |  Pablo Daniel Barillas Moreno · 22193"); names.alignment=WD_ALIGN_PARAGRAPH.CENTER
    table=doc.add_table(rows=2,cols=4); table.style="Table Grid"
    values=[('AUC-PR',f(EVAL['auc_pr'])),('ROC-AUC',f(EVAL['roc_auc'])),('Precisión',pct(EVAL['precision'])),('Recall',pct(EVAL['recall'])),('F1',f(EVAL['f1'])),('Costo',money(EVAL['costo_q'])),('AP walk',f(next(x['mean'] for x in R['validacion_walk_forward']['resumen'] if x['modelo']=='LightGBM_tuned'))),('Estado','CANDIDATO')]
    for cell,(label,value) in zip([c for row in table.rows for c in row.cells],values):
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run(label.upper()+"\n"); rr.bold=True; rr.font.size=Pt(8); vv=p.add_run(value); vv.bold=True; vv.font.size=Pt(12); vv.font.color.rgb=RGBColor(42,157,143)
    doc.add_heading("Resumen",level=1); doc.add_paragraph("V4 integra 360 variables numéricas y 38 categóricas, optimiza LightGBM y selecciona expertos por ProductCD antes de calibración y evaluación. La política robusta mejora todas las métricas frente a V3 y reduce costo 18.5% en el bloque comparable.")
    doc.add_heading("Decisión",level=1); doc.add_paragraph("V4 es un candidato superior post-hoc. V3 permanece confirmada hasta validar este pipeline y umbral sin cambios en una cohorte temporal nueva. Uso académico para priorizar revisión humana; no bloquear transacciones.")
    pic=doc.add_paragraph(); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER; pic.add_run().add_picture(str(qr_path),width=Inches(1.15)); link=doc.add_paragraph(URL); link.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.save(out_dir / "Ficha_Repositorio_Proyecto_1_V4.docx")

    styles=getSampleStyleSheet(); pdf=SimpleDocTemplate(str(out_dir/"Ficha_Repositorio_Proyecto_1_V4.pdf"),pagesize=letter,rightMargin=.55*inch,leftMargin=.55*inch,topMargin=.45*inch,bottomMargin=.45*inch)
    story=[Paragraph("<b>PROYECTO 1 · MONITOREO TRANSACCIONAL · V4</b>",styles['Title']),Paragraph("Universidad del Valle de Guatemala · Grupo 1 · Sección 30",styles['Heading3']),Paragraph("Wilson Alejandro Calderón Argueta · 22018 | Pablo Daniel Barillas Moreno · 22193",styles['BodyText']),Spacer(1,10)]
    data=[["AUC-PR",f(EVAL['auc_pr']),"ROC-AUC",f(EVAL['roc_auc'])],["Precisión",pct(EVAL['precision']),"Recall",pct(EVAL['recall'])],["F1",f(EVAL['f1']),"Costo",money(EVAL['costo_q'])]]
    tab=Table(data,colWidths=[1.1*inch,1.35*inch,1.1*inch,1.35*inch]); tab.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,-1),colors.HexColor('#EDF5FB')),('GRID',(0,0),(-1,-1),.5,colors.HexColor('#9FB7C8')),('FONTNAME',(0,0),(-1,-1),'Helvetica-Bold'),('ALIGN',(0,0),(-1,-1),'CENTER'),('PADDING',(0,0),(-1,-1),8)]))
    qr = RLImage(str(qr_path), width=1.12*inch, height=1.12*inch); qr.hAlign = "CENTER"
    story.extend([tab,Spacer(1,14),Paragraph("V4 usa expertos ProductCD, 360 numéricas, 38 categorías, ponderación por recencia y validación walk-forward. Mejora todas las métricas en evaluación, pero la política robusta es post-hoc.",styles['BodyText']),Spacer(1,8),Paragraph("Estado: candidato congelado. Requiere cohorte temporal nueva antes de promoción confirmatoria o producción.",styles['BodyText']),Spacer(1,10),qr,Spacer(1,6),Paragraph(URL,styles['Heading3'])]); pdf.build(story)


def main() -> None:
    build_readme(); build_notebook(); report=build_report(); slides=build_slides(); build_ficha()
    print("README, notebook, informe, presentación y ficha V4 generados desde", RESULT_PATH)
    print("Informe:", report, "Presentación:", slides)


if __name__ == "__main__":
    main()
