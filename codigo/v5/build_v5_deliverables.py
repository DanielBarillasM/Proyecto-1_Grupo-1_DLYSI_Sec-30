"""Genera todos los entregables rubricados de Proyecto 1 V5 desde resultados_v5.json."""

from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path

import nbformat as nbf
import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[2]
RESULT_PATH = ROOT / "artefactos" / "v5" / "resultados_v5.json"
R = json.loads(RESULT_PATH.read_text(encoding="utf-8"))
INTERNAL = R["evaluacion_interna"]
BENCH = R["benchmark_historico"]
CANDIDATE = R["candidato"]["modelo"]
FALS = R["falsificaciones"]
URL = "https://github.com/DanielBarillasM/Proyecto-1_Grupo-1_DLYSI_Sec-30"


def f(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def pct(value: float, digits: int = 2) -> str:
    return f"{100 * value:.{digits}f}%"


def money(value: float) -> str:
    return f"Q{value:,.0f}"


def notebook_style() -> str:
    return """<style>
:root{--navy:#102a43;--blue:#184e77;--teal:#2a9d8f;--gold:#e9c46a;--red:#e76f51;--ink:#172033;--pale:#edf5fb}
.hero{padding:38px 42px;border-radius:24px;color:#f8fbff;background:radial-gradient(circle at 92% 8%,rgba(255,255,255,.16) 0 8%,transparent 9%),linear-gradient(125deg,var(--navy),var(--blue) 55%,var(--teal));box-shadow:0 16px 38px #102a433d;font-family:Inter,'Segoe UI',sans-serif}.hero h1{font-size:38px;margin:18px 0 8px;color:white}.chips{display:flex;gap:8px;flex-wrap:wrap}.chips span{padding:6px 13px;border:1px solid #ffffff55;border-radius:999px;background:#ffffff20;font-size:11px;font-weight:800}.grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(205px,1fr));gap:13px;margin-top:24px}.card{padding:14px 16px;border:1px solid #ffffff44;border-radius:12px;background:#ffffff12}.section{margin:28px 0 14px;padding:16px 22px;border-radius:14px;background:linear-gradient(90deg,var(--navy),var(--blue));color:white}.section h2{margin:0;color:white}.call{padding:18px 22px;margin:14px 0;border:1px solid #c9d9e6;border-left:6px solid var(--teal);border-radius:13px;background:var(--pale);color:var(--ink);line-height:1.7}.warn{border-left-color:var(--red);background:#fff2ed}.proof{border-left-color:var(--gold);background:#fff9e9}.metric{display:inline-block;min-width:145px;padding:12px 14px;margin:4px;border-radius:10px;background:#e5f5f2;color:#184e77;text-align:center}.metric b{display:block;font-size:21px}table{font-size:14px}th{background:#184e77!important;color:white!important;text-align:left!important}td,th{padding:9px!important}code{background:#102a4312;padding:2px 5px;border-radius:4px}
</style>"""


def build_notebook() -> Path:
    out_dir = ROOT / "entregables" / "cuaderno" / "v5"
    out_dir.mkdir(parents=True, exist_ok=True)
    hero = notebook_style() + """<div class="hero"><div class="chips"><span>DEEP LEARNING</span><span>PROYECTO 1</span><span>V5 INTEGRADA</span><span>GRUPO 1 · SECCIÓN 30</span></div><h1>Monitoreo transaccional: detectar lo que el orden revela</h1><p style="font-size:20px">Comparación rubricada A/B/C, falsificación temporal y decisión económica</p><div class="grid"><div class="card"><b>Institución</b><br>Universidad del Valle de Guatemala</div><div class="card"><b>Curso</b><br>Deep Learning y Sistemas Inteligentes</div><div class="card"><b>Docente</b><br>Kevin Recinos</div><div class="card"><b>Integrantes</b><br>Wilson Alejandro Calderón Argueta · 22018<br>Pablo Daniel Barillas Moreno · 22193</div><div class="card"><b>Datos</b><br>IEEE-CIS · 590,540 transacciones</div><div class="card"><b>Candidato</b><br>A · LightGBM expertos V4</div></div></div>"""
    cells = [nbf.v4.new_markdown_cell(hero)]
    cells.append(nbf.v4.new_code_cell("""from pathlib import Path
import json, pandas as pd, numpy as np
from sklearn.metrics import average_precision_score, roc_auc_score
from IPython.display import Image, display
ROOT=Path.cwd().resolve()
while not (ROOT/'artefactos').exists() and ROOT != ROOT.parent:
    ROOT=ROOT.parent
R=json.loads((ROOT/'artefactos/v5/resultados_v5.json').read_text(encoding='utf-8'))
VAL=pd.read_csv(ROOT/'artefactos/v5/predicciones_validacion_v5.csv')
BENCH=pd.read_csv(ROOT/'artefactos/v5/predicciones_benchmark_v5.csv')
R['candidato']"""))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>1 · Pregunta y decisión ejecutiva</h2></div>

La investigación responde: **¿el orden de las transacciones aporta información que las variables agregadas no capturan, bajo qué condiciones y cuánto vale en quetzales?**

<div class="call"><b>Decisión:</b> conservar A —LightGBM con expertos por <code>ProductCD</code>— como candidato y usarlo para priorizar revisión humana. B no supera A y la permutación mejora, en vez de degradar, su AUC-PR; por ello no se atribuye valor predictivo material al orden.</div>

<div class="call warn"><b>Alcance:</b> el último 15 % es un benchmark histórico reutilizado. La selección de V5 ocurre dentro de validación; una promoción confirmatoria requiere una cohorte temporal nueva.</div>"""))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>2 · Integridad de datos y protocolo temporal</h2></div>

IEEE-CIS contiene transacciones reales anonimizadas de Vesta. `TransactionID` une las tablas y `TransactionDT` establece el orden; ninguno se introduce como número predictivo. La identidad se aproxima con `card1 + card2 + card3 + card5 + addr1` y puede mezclar o fragmentar clientes.

Para una transacción en tiempo $t$, toda estadística histórica satisface:

$$x_t^{hist}=f(\{x_j:t_j<t\}).$$

La imputación, estandarización y vocabularios categóricos se ajustan únicamente con entrenamiento."""))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(R['datos']['particiones']).T"))
    cells.append(nbf.v4.new_code_cell("pd.Series(R['secuencias']).to_frame('valor')"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>3 · Núcleo comparable A/B</h2></div>

| Pieza | Entrada | Modelo | Uso del orden |
|---|---|---|---|
| **A** | Vector V4 con agregados causales y categorías | LightGBM con expertos W/NO-W | No lee eventos ordenados |
| **B** | Hasta 16 eventos cronológicos, 57 variables numéricas y 12 categóricas | Embeddings + GRU(64) | Sí puede aprender transiciones |

Los dos modelos generan un puntaje continuo y predicen las mismas transacciones con el mismo horizonte. B se entrenó sobre las 413,378 observaciones del período de entrenamiento, con BCE ponderada, AdamW, clipping y early stopping."""))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(R['entrenamiento_B'])"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>4 · Hipótesis previa de la apuesta C</h2></div>

<div class="call proof"><b>Hipótesis congelada antes del ajuste de C.</b><br>Creemos que fusionar el puntaje tabular de LightGBM con el puntaje secuencial de la GRU mejorará el AUC-PR porque ambos modelos representan información complementaria. Lo consideraremos útil si incrementa AUC-PR al menos 0.01 y reduce el costo al menos 5 % frente al mejor modelo individual en la evaluación interna de validación temporal.</div>

C utiliza regresión logística sobre $[\operatorname{logit}(s_A),\operatorname{logit}(s_B),\log(1+monto),L,\mathbb{1}(ProductCD=W)]$. El metamodelo, calibradores, umbrales y evaluación ocupan bloques cronológicos distintos."""))
    cells.append(nbf.v4.new_code_cell("pd.Series(R['hipotesis_C']).to_frame('resultado')"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>5 · Comparación común en evaluación interna</h2></div>

AUC-PR es la métrica de ranking principal. Precisión, recall y F1 se calculan con el umbral seleccionado por costo mínimo sujeto al requisito predeclarado de recall $\ge 0.75$ en el bloque de umbral. La exactitud no se utiliza para seleccionar modelos."""))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(R['evaluacion_interna']).T[['auc_pr','roc_auc','precision','recall','f1','cost_q','alertas_por_100k','threshold']]"))
    cells.append(nbf.v4.new_code_cell("""eval_rows=np.arange(int(len(VAL)*.75),len(VAL))
for name in ['A_V4','B_GRU','C_fusion']:
    score=VAL[f'score_{name}'].to_numpy()[eval_rows]
    assert abs(average_precision_score(VAL.y.to_numpy()[eval_rows],score)-R['evaluacion_interna'][name]['auc_pr'])<1e-10
print('Métricas A/B/C reproducidas desde los puntajes guardados.')"""))
    cells.append(nbf.v4.new_code_cell("display(Image(filename=str(ROOT/'evidencia/figuras/v5/01_comparacion_abc_validacion.png')))"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>6 · Dos intentos de refutar el valor del orden</h2></div>

La permutación conserva los eventos y la transacción objetivo, pero baraja solo los antecedentes. Se repite con cinco semillas. La segunda prueba recorta el historial a 3 y 8 eventos sin reentrenar.

<div class="call warn"><b>Resultado:</b> B original obtiene AP 0.4245 y B permutada 0.4355 ± 0.0021. La diferencia original−permutada es −0.0111. Destruir el orden no perjudica a B; por el contrario, mejora el ranking. No existe evidencia para afirmar que el orden aporte.</div>"""))
    cells.append(nbf.v4.new_code_cell("""pd.DataFrame({
'B original':R['falsificaciones']['original_internal'],
'B permutada':{'auc_pr':R['falsificaciones']['permutation_mean_auc_pr']},
'Historia 3':R['falsificaciones']['historia_3'],
'Historia 8':R['falsificaciones']['historia_8'],
}).T"""))
    cells.append(nbf.v4.new_code_cell("display(Image(filename=str(ROOT/'evidencia/figuras/v5/03_falsificaciones_orden_v5.png')))"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>7 · Veredicto de C</h2></div>

C aumenta recall, pero no satisface la regla previa: pierde 0.0020 de AUC-PR y aumenta 0.43 % el costo frente a A en evaluación interna. La apuesta se declara **no útil**, aunque en el benchmark histórico su ranking sea descriptivamente cercano a A. La decisión no se reescribe después de observar ese bloque."""))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>8 · Umbral y decisión económica</h2></div>

La función de costo suministrada es:

$$C(\tau)=Q4{,}200\,FN(\tau)+Q180\,FP(\tau).$$

El candidato A usa $\tau=0.05783$. En el benchmark histórico logra AP 0.5592, precisión 22.15 %, recall 73.27 %, F1 0.3402 y costo Q4,890,000. La proyección mensual es un escenario, no una cifra contable, porque el PDF no fija transacciones por tarjeta."""))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(R['benchmark_historico']).T[['auc_pr','roc_auc','precision','recall','f1','cost_q','alertas_por_100k']]"))
    cells.append(nbf.v4.new_code_cell("""rows=[]
for model,scenarios in R['economia_mensual'].items():
    for tx,values in scenarios.items(): rows.append({'modelo':model,'tx_tarjeta_mes':int(tx),**values})
pd.DataFrame(rows).set_index(['modelo','tx_tarjeta_mes'])"""))
    cells.append(nbf.v4.new_code_cell("display(Image(filename=str(ROOT/'evidencia/figuras/v5/04_costos_abc_v5.png')))"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>9 · Errores, recomendación y límites</h2></div>

Los falsos negativos de alto monto son el error prioritario por su costo. El archivo `falsos_negativos_alto_monto_v5.csv` permite inspeccionarlos sin confundir una predicción con culpabilidad.

Se recomienda **conservar y complementar** el sistema agregado con A para ordenar revisiones humanas, no bloquear automáticamente. La recomendación cambiaría si una identidad bancaria fiable produjera una caída material al permutar, si una cohorte nueva revirtiera el resultado o si la capacidad operativa no absorbiera las alertas.

Limitaciones: identidad proxy, anonimización, benchmark reutilizado, costos académicos, deriva temporal y ausencia de evaluación de privacidad, equidad, latencia, seguridad y apelación."""))
    cells.append(nbf.v4.new_code_cell("pd.read_csv(ROOT/'artefactos/v5/metricas_segmentos_v5.csv').sort_values(['dimension','auc_pr'],ascending=[True,False])"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>10 · Matriz de evidencias</h2></div>

| Evidencia | Figura o tabla | Conclusión | Limitación |
|---|---|---|---|
| Integridad de datos | Particiones y cobertura | 70/15/15 cronológico; preprocesamiento train-only | Identidad aproximada |
| Comparación A/B | Figura 1 y tabla A/B/C | A supera B en AP, F1 y costo | A posee ingeniería tabular más amplia |
| Valor del orden | Figura 3 | Permutar no reduce AP; no se atribuye valor al orden | Una GRU y una clave proxy |
| Apuesta C | Tabla de hipótesis | No alcanza +0.01 AP ni −5 % costo | Fusión logística específica |
| Decisión económica | Figura 4 | A minimiza costo interno bajo política predefinida | Costos y volumen son escenarios |
| Recomendación y límites | Sección 9 | Complementar con revisión humana | Requiere cohorte nueva y controles productivos |"""))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>11 · Reproducibilidad y declaración de uso de inteligencia artificial</h2></div>

Los pesos, transformaciones, calibradores, umbrales, contrato y puntajes se encuentran en `artefactos/v5/`. El pipeline se ejecuta con `python -u codigo/v5/proyecto1_v5_pipeline.py` y los entregables se reconstruyen con `python codigo/v5/build_v5_deliverables.py`.

Se utilizó asistencia de IA para estructurar código, revisar consistencia, redactar, diseñar HTML/CSS/LaTeX y automatizar auditorías. Los integrantes ejecutaron los experimentos, verificaron alineación temporal, métricas, artefactos y conclusiones; deben poder defender todas las decisiones."""))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>Referencias APA 7</h2></div>

Cho, K., et al. (2014). Learning phrase representations using RNN encoder–decoder for statistical machine translation. *Proceedings of EMNLP*, 1724–1734. https://doi.org/10.3115/v1/D14-1179

IEEE Computational Intelligence Society. (2019). *IEEE-CIS Fraud Detection* [Data set]. Kaggle. https://www.kaggle.com/competitions/ieee-fraud-detection

Ke, G., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. *Advances in Neural Information Processing Systems, 30*.

Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. *PLOS ONE, 10*(3), e0118432. https://doi.org/10.1371/journal.pone.0118432"""))
    notebook = nbf.v4.new_notebook(cells=cells, metadata={"kernelspec":{"display_name":"Python 3","language":"python","name":"python3"},"language_info":{"name":"python","version":"3.13"}})
    out = out_dir / "proyecto1_calderon_barillas.ipynb"
    nbf.write(notebook, out)
    return out


def build_readme() -> str:
    center = R["economia_mensual"][CANDIDATE]["12"]
    text = f"""<div align="center">

# Proyecto 1 · Monitoreo transaccional · V5 integrada

### ¿El orden aporta información más allá de los agregados?

![Python](https://img.shields.io/badge/Python-3.13-3776AB?logo=python&logoColor=white)
![PyTorch](https://img.shields.io/badge/PyTorch-2.13-EE4C2C?logo=pytorch&logoColor=white)
![LightGBM](https://img.shields.io/badge/LightGBM-4.7-184e77)
![Estado](https://img.shields.io/badge/Candidato-A__V4-2a9d8f)

**Wilson Alejandro Calderón Argueta · 22018** · **Pablo Daniel Barillas Moreno · 22193**

Universidad del Valle de Guatemala · Deep Learning y Sistemas Inteligentes · Sección 30 · 2026

</div>

> [!IMPORTANT]
> El último 15 % de IEEE-CIS ya fue observado en iteraciones anteriores y se reporta como benchmark temporal histórico reutilizado. Todas las decisiones V5 se toman dentro de validación. Una promoción confirmatoria exige una cohorte nueva.

## Resumen ejecutivo

El proyecto estudia {R['datos']['filas']:,} transacciones IEEE-CIS, con {R['datos']['fraudes']:,} fraudes y prevalencia {pct(R['datos']['prevalencia'])}. Compara una línea tabular competitiva sin leer eventos ordenados (A), una GRU causal sobre secuencias de hasta 16 eventos (B) y una fusión leakage-safe de sus puntajes (C). A obtiene AP interna {f(INTERNAL['A_V4']['auc_pr'],4)}, mientras B obtiene {f(INTERNAL['B_GRU']['auc_pr'],4)} y C {f(INTERNAL['C_fusion']['auc_pr'],4)}.

La permutación controlada no perjudica a B: su AP cambia de {f(FALS['original_internal']['auc_pr'],4)} a {f(FALS['permutation_mean_auc_pr'],4)} ± {f(FALS['permutation_std_auc_pr'],4)}. La diferencia original−permutada es {f(FALS['order_auc_pr_drop'],4)}. Con esta evidencia no se afirma que el orden aporte. C tampoco supera su criterio previo: cambio AP {f(R['hipotesis_C']['auc_pr_gain'],4)} y reducción de costo {pct(R['hipotesis_C']['cost_reduction'])}. El candidato es A.

## Datos y protocolo temporal

Se utiliza la competencia pública IEEE-CIS Fraud Detection de Kaggle. `train_transaction.csv` y `train_identity.csv` se unen por `TransactionID`, las filas se ordenan por `TransactionDT` y ambos campos se excluyen como magnitudes predictivas. La identidad secuencial es una clave aproximada formada por `card1`, `card2`, `card3`, `card5` y `addr1`.

La separación es 70 % entrenamiento, 15 % validación y 15 % benchmark histórico. Validación se subdivide cronológicamente en early stopping, ajuste de C, calibración, umbral y evaluación interna. Imputación, escalado y vocabularios se aprenden exclusivamente con entrenamiento. Las características históricas utilizan solo eventos anteriores; no se aplican particiones aleatorias.

## Modelos A/B/C

| Pieza | Diseño | Resultado interno |
|---|---|---:|
| A | LightGBM V4 con expertos `ProductCD=W/NO-W` | AP {f(INTERNAL['A_V4']['auc_pr'],4)} · costo {money(INTERNAL['A_V4']['cost_q'])} |
| B | Embeddings + GRU(64), hasta 16 eventos | AP {f(INTERNAL['B_GRU']['auc_pr'],4)} · costo {money(INTERNAL['B_GRU']['cost_q'])} |
| C | Regresión logística sobre A, B, monto, historia y producto | AP {f(INTERNAL['C_fusion']['auc_pr'],4)} · costo {money(INTERNAL['C_fusion']['cost_q'])} |

A incorpora la ingeniería causal y selección V4. B utiliza 57 variables numéricas por evento, 12 categorías, BCE ponderada, AdamW, clipping y early stopping. C se entrena en un bloque independiente con puntajes A/B y no recibe el benchmark para decidir su arquitectura.

## Hipótesis y falsificaciones

**Hipótesis previa de C:** {R['hipotesis_C']['declaracion_previa']}

La hipótesis no se cumple. C pierde {abs(R['hipotesis_C']['auc_pr_gain']):.4f} de AP y aumenta el costo {abs(100*R['hipotesis_C']['cost_reduction']):.2f} % frente a A en evaluación interna.

Las dos pruebas obligatorias son:

1. Permutación de antecedentes con cinco semillas, manteniendo la transacción objetivo al final.
2. Recorte de la historia a 3 y 8 eventos.

La historia de 3 eventos obtiene AP {f(FALS['historia_3']['auc_pr'],4)} y la de 8 obtiene {f(FALS['historia_8']['auc_pr'],4)}. Ninguna evidencia justifica afirmar que el orden mejore el detector.

## Resultados y decisión económica

| Modelo | AUC-PR benchmark | Precisión | Recall | F1 | Costo |
|---|---:|---:|---:|---:|---:|
| **A** | **{f(BENCH['A_V4']['auc_pr'])}** | **{pct(BENCH['A_V4']['precision'])}** | {pct(BENCH['A_V4']['recall'])} | **{f(BENCH['A_V4']['f1'])}** | **{money(BENCH['A_V4']['cost_q'])}** |
| B | {f(BENCH['B_GRU']['auc_pr'])} | {pct(BENCH['B_GRU']['precision'])} | {pct(BENCH['B_GRU']['recall'])} | {f(BENCH['B_GRU']['f1'])} | {money(BENCH['B_GRU']['cost_q'])} |
| C | {f(BENCH['C_fusion']['auc_pr'])} | {pct(BENCH['C_fusion']['precision'])} | **{pct(BENCH['C_fusion']['recall'])}** | {f(BENCH['C_fusion']['f1'])} | {money(BENCH['C_fusion']['cost_q'])} |

La política de umbral minimiza $4200FN+180FP$ sujeta a recall ≥ 0.75 en selección. El umbral de A es {R['umbrales']['A_V4']:.5f}. En el escenario central de 12 transacciones por tarjeta al mes, A representa un costo mensual proyectado de {money(center['costo_mensual_q'])}. Es una extrapolación académica, no una cifra contable.

## Tres decisiones técnicas importantes

1. **A tabular V4 frente al HistGradientBoosting V1.** Se consideró conservar el baseline antiguo. Se eligió LightGBM con expertos porque usa más variables causales y obtiene AP y costo claramente mejores.
2. **GRU frente a LSTM, TCN o Transformer.** Se eligió GRU por eficiencia en CPU, menor cantidad de parámetros y porque la rúbrica evalúa evidencia del orden, no complejidad. La permutación permitió comprobar si realmente aprovechó la secuencia.
3. **Fusión tardía frente al híbrido interno V1.** Se eligió stacking logístico con puntajes A/B porque separa el aporte de cada modelo, permite controles temporales y reduce el riesgo de que C reconstruya toda la función tabular dentro de una red opaca.

## Reproducción

```powershell
py -3.13 -m venv .venv
.\.venv\Scripts\Activate.ps1
python -m pip install -r configuracion/v5/requirements-v5.txt
python -m pip install -r configuracion/v5/requirements-docs-v5.txt
python codigo/compartido/download_data.py
python -u codigo/v5/proyecto1_v5_pipeline.py
python codigo/v5/build_v5_deliverables.py
jupyter nbconvert --to notebook --execute --inplace --ExecutePreprocessor.timeout=300 entregables/cuaderno/v5/proyecto1_calderon_barillas.ipynb
python codigo/v5/audit_project1_v5.py
```

La descarga requiere aceptar las reglas de IEEE-CIS en Kaggle y configurar las credenciales fuera del repositorio. Los CSV de casi 700 MB no se versionan.

## Estructura

```text
codigo/v5/              pipeline, construcción y auditoría de V5
configuracion/v5/       versiones exactas e instrucciones
datos/raw/              CSV locales ignorados por Git
artefactos/v5/          pesos A/B/C, calibradores, contrato y puntajes
evidencia/figuras/v5/   resultados reproducibles
entregables/cuaderno/v5/ notebook ejecutado
entregables/informe/v5/  fuente LaTeX y PDF
entregables/presentacion/v5/ HTML y PDF de ocho diapositivas
entregables/ficha/v5/    ficha del repositorio
```

## Candidato al Proyecto Final

- **Modelo conservado:** A — LightGBM V4 con expertos por `ProductCD` y calibrador V5.
- **Artefactos:** `artefactos/v4/modelo_experto_w_v4.txt`, `modelo_experto_no_w_v4.txt`, `artefactos/v5/calibradores_v5.joblib` y `contrato_entrada_salida_v5.json`.
- **Usuario:** analista de riesgo o equipo de monitoreo transaccional.
- **Decisión:** ordenar alertas y priorizar revisión; el puntaje no prueba fraude ni autoriza bloqueo autónomo.
- **Entrada preliminar:** transacción actual, variables categóricas y estadísticas históricas causales especificadas por el contrato.
- **Salida:** `risk_score` continuo en [0,1], umbral {R['umbrales']['A_V4']:.5f} y política de revisión.
- **Pendientes:** nueva cohorte etiquetada, identidad bancaria fiable, costos reales, latencia, privacidad, equidad, seguridad, explicaciones y monitoreo.

## Limitaciones y uso responsable

IEEE-CIS está anonimizado y cubre aproximadamente 182 días. La clave proxy no equivale a un cliente real. El benchmark ya fue observado y ninguna conclusión se presenta como confirmación externa. Los costos y volúmenes mensuales son escenarios. El sistema debe apoyar revisión humana, no atribuir culpabilidad ni bloquear de manera autónoma.

## Declaración de uso de inteligencia artificial

Se utilizó asistencia de IA para estructurar y revisar código, diseñar documentación HTML/CSS/LaTeX, localizar bibliografía y automatizar auditorías. Los integrantes ejecutaron el pipeline y verificaron particiones, alineación de IDs, métricas, falsificaciones, umbrales y artefactos. La IA no se utilizó como fuente académica ni reemplaza la defensa de las decisiones.

## Referencias APA 7

Cho, K., et al. (2014). Learning phrase representations using RNN encoder–decoder for statistical machine translation. *Proceedings of EMNLP*, 1724–1734. https://doi.org/10.3115/v1/D14-1179

IEEE Computational Intelligence Society. (2019). *IEEE-CIS Fraud Detection* [Data set]. Kaggle. https://www.kaggle.com/competitions/ieee-fraud-detection

Ke, G., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. *Advances in Neural Information Processing Systems, 30*.

Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. *PLOS ONE, 10*(3), e0118432. https://doi.org/10.1371/journal.pone.0118432
"""
    out = ROOT / "entregables" / "ficha" / "v5" / "README_V5_GENERADO.md"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(text, encoding="utf-8", newline="\n")
    return text


def build_report() -> Path:
    center_a = R["economia_mensual"]["A_V4"]["12"]
    center_b = R["economia_mensual"]["B_GRU"]["12"]
    center_c = R["economia_mensual"]["C_fusion"]["12"]
    tex = rf"""\documentclass[9.5pt]{{article}}
\usepackage[utf8]{{inputenc}}\usepackage[T1]{{fontenc}}\usepackage[letterpaper,margin=1.45cm]{{geometry}}
\usepackage{{graphicx,booktabs,xcolor,amsmath,hyperref,array,tabularx}}\definecolor{{navy}}{{HTML}}{{102A43}}\definecolor{{teal}}{{HTML}}{{2A9D8F}}
\hypersetup{{colorlinks=true,urlcolor=teal}}\setlength{{\parindent}}{{0pt}}\setlength{{\parskip}}{{4pt}}
\begin{{document}}\begin{{titlepage}}\pagecolor{{navy}}\color{{white}}\raggedright\vspace*{{1cm}}{{\Large DEEP LEARNING Y SISTEMAS INTELIGENTES\par}}\vspace{{1cm}}{{\Huge\bfseries Proyecto 1\\Monitoreo transaccional\par}}\vspace{{.6cm}}{{\LARGE Detectar lo que el orden revela · V5 integrada\par}}\vfill{{\Large Universidad del Valle de Guatemala\\Kevin Recinos · Grupo 1 · Sección 30\\Semestre II 2026\par}}\vspace{{1cm}}{{\Large Wilson Alejandro Calderón Argueta · 22018\\Pablo Daniel Barillas Moreno · 22193\par}}\vfill\textbf{{Decisión:}} conservar A para priorizar revisión humana; la evidencia no demuestra valor material del orden.\end{{titlepage}}\nopagecolor

\section*{{Resumen ejecutivo}} Se estudiaron {R['datos']['filas']:,} transacciones IEEE-CIS ({pct(R['datos']['prevalencia'])} fraude). La investigación compara A, LightGBM sin lectura secuencial; B, GRU causal de 16 eventos; y C, fusión A+B. En evaluación interna, A obtiene AP {INTERNAL['A_V4']['auc_pr']:.4f}, B {INTERNAL['B_GRU']['auc_pr']:.4f} y C {INTERNAL['C_fusion']['auc_pr']:.4f}. A minimiza costo ({money(INTERNAL['A_V4']['cost_q'])}) y queda como candidato. Permutar B cambia AP de {FALS['original_internal']['auc_pr']:.4f} a {FALS['permutation_mean_auc_pr']:.4f}; no se atribuye valor al orden. El benchmark es histórico reutilizado, no test ciego.

\section{{Integridad de datos}} Las tablas IEEE-CIS se unen por TransactionID y ordenan por TransactionDT; ambos se excluyen como magnitudes. La identidad proxy combina card1, card2, card3, card5 y addr1. Cada secuencia contiene la transacción objetivo y hasta 15 antecedentes; {R['secuencias']['porcentaje_con_8']:.1f}\% tiene longitud al menos 8 y {R['secuencias']['porcentaje_con_16']:.1f}\% llega a 16. La división es 70\% entrenamiento, 15\% validación y 15\% benchmark histórico. Imputación, escalado, categorías y umbrales se aprenden antes del bloque evaluado. Toda variable histórica cumple $x_t^{{hist}}=f(\{{x_j:t_j<t\}})$.

\begin{{figure}}[h]\centering\includegraphics[width=.72\linewidth]{{../../../evidencia/figuras/v5/01_comparacion_abc_validacion.png}}\caption{{Comparación común A/B/C en evaluación interna temporal.}}\end{{figure}}
\newpage
\section{{Núcleo A/B y apuesta C}} \textbf{{A}} reutiliza LightGBM V4 con expertos W/NO-W y agregados causales. \textbf{{B}} emplea 57 numéricas, 12 categóricas con embeddings, proyección de 96, GRU(64), BCE ponderada, AdamW y early stopping sobre las 413,378 filas de entrenamiento. Ambos producen riesgo continuo sobre las mismas filas. \textbf{{C}} combina logits A/B, monto, longitud y producto mediante regresión logística ajustada en un bloque separado.

\textbf{{Hipótesis previa:}} fusionar A+B será útil si aumenta AP al menos 0.01 y reduce costo al menos 5\% frente al mejor modelo individual. C obtiene cambio AP {R['hipotesis_C']['auc_pr_gain']:.4f} y reducción de costo {100*R['hipotesis_C']['cost_reduction']:.2f}\%; el veredicto es \textbf{{no útil}}.

\begin{{center}}\begin{{tabular}}{{lrrrrrr}}\toprule Modelo&AP&ROC&Prec.&Recall&F1&Costo\\\midrule
A&{INTERNAL['A_V4']['auc_pr']:.3f}&{INTERNAL['A_V4']['roc_auc']:.3f}&{INTERNAL['A_V4']['precision']:.3f}&{INTERNAL['A_V4']['recall']:.3f}&{INTERNAL['A_V4']['f1']:.3f}&{money(INTERNAL['A_V4']['cost_q'])}\\
B&{INTERNAL['B_GRU']['auc_pr']:.3f}&{INTERNAL['B_GRU']['roc_auc']:.3f}&{INTERNAL['B_GRU']['precision']:.3f}&{INTERNAL['B_GRU']['recall']:.3f}&{INTERNAL['B_GRU']['f1']:.3f}&{money(INTERNAL['B_GRU']['cost_q'])}\\
C&{INTERNAL['C_fusion']['auc_pr']:.3f}&{INTERNAL['C_fusion']['roc_auc']:.3f}&{INTERNAL['C_fusion']['precision']:.3f}&{INTERNAL['C_fusion']['recall']:.3f}&{INTERNAL['C_fusion']['f1']:.3f}&{money(INTERNAL['C_fusion']['cost_q'])}\\\bottomrule\end{{tabular}}\end{{center}}
\newpage
\section{{Valor del orden}} La permutación baraja solo los antecedentes, conserva eventos, variables agregadas y transacción objetivo, y se repite cinco veces. B original obtiene AP {FALS['original_internal']['auc_pr']:.4f}; permutada, {FALS['permutation_mean_auc_pr']:.4f}$\pm${FALS['permutation_std_auc_pr']:.4f}. La diferencia original--permutada es {FALS['order_auc_pr_drop']:.4f}. Recortar a 3 eventos produce {FALS['historia_3']['auc_pr']:.4f}; usar 8 produce {FALS['historia_8']['auc_pr']:.4f}. Destruir o reducir la historia no ocasiona una caída material: no puede afirmarse que el orden aporte.

\begin{{figure}}[h]\centering\includegraphics[width=.78\linewidth]{{../../../evidencia/figuras/v5/03_falsificaciones_orden_v5.png}}\caption{{Permutación controlada y segunda prueba de falsificación.}}\end{{figure}}

La caída de B entre evaluación interna y benchmark (AP {INTERNAL['B_GRU']['auc_pr']:.3f} a {BENCH['B_GRU']['auc_pr']:.3f}) revela sensibilidad temporal. Una secuencia más rica no garantiza estabilidad si la identidad es aproximada.
\newpage
\section{{Umbral y valor económico}} La política predefinida minimiza $4200FN+180FP$ sujeto a recall $\ge0.75$ en el bloque de umbral. A usa $\tau={R['umbrales']['A_V4']:.5f}$. En benchmark histórico logra AP {BENCH['A_V4']['auc_pr']:.4f}, precisión {BENCH['A_V4']['precision']:.4f}, recall {BENCH['A_V4']['recall']:.4f}, F1 {BENCH['A_V4']['f1']:.4f} y costo {money(BENCH['A_V4']['cost_q'])}.

\begin{{center}}\begin{{tabular}}{{lrrr}}\toprule Modelo&Costo 12 tx/tarjeta/mes&Ahorro vs. A&Alertas/100k\\\midrule
A&{money(center_a['costo_mensual_q'])}&Q0&{BENCH['A_V4']['alertas_por_100k']:,.0f}\\
B&{money(center_b['costo_mensual_q'])}&{money(center_b['ahorro_vs_A_q'])}&{BENCH['B_GRU']['alertas_por_100k']:,.0f}\\
C&{money(center_c['costo_mensual_q'])}&{money(center_c['ahorro_vs_A_q'])}&{BENCH['C_fusion']['alertas_por_100k']:,.0f}\\\bottomrule\end{{tabular}}\end{{center}}

La proyección usa 1.4 millones de tarjetas y escenarios 5/12/20 transacciones por tarjeta; no es una cifra contable. Ahorros negativos representan pérdidas frente a A.

\begin{{figure}}[h]\centering\includegraphics[width=.66\linewidth]{{../../../evidencia/figuras/v5/04_costos_abc_v5.png}}\caption{{Costo de los tres modelos en evaluación interna.}}\end{{figure}}
\newpage
\section{{Recomendación, errores y límites}} Se recomienda \textbf{{conservar y complementar}} el sistema tabular con A para ordenar revisión humana. No se recomienda migrar a B ni bloquear automáticamente. Los falsos negativos de monto alto son el error prioritario; se exportan cien casos para análisis. La recomendación cambiaría si una identidad fiable causara una caída material al permutar, una cohorte nueva favoreciera B/C o la operación no absorbiera las alertas.

Límites: benchmark reutilizado, identidad proxy, anonimización, 182 días, costos académicos, deriva y ausencia de privacidad, equidad, seguridad, explicaciones, latencia y apelación. Antes de producción se requiere cohorte temporal nueva y piloto humano.

\section*{{Tres decisiones defendibles}} (1) LightGBM V4 reemplaza HGB V1 por cobertura y costo; (2) GRU se prefiere a Transformer/TCN por eficiencia y falsabilidad; (3) stacking tardío reemplaza el híbrido interno para aislar aportes. Ninguna decisión usa accuracy ni partición aleatoria.

\section*{{Referencias}}\small Cho, K., et al. (2014). GRU encoder--decoder. \textit{{EMNLP}}, 1724--1734.\\IEEE CIS. (2019). \textit{{IEEE-CIS Fraud Detection}}. Kaggle.\\Ke, G., et al. (2017). LightGBM. \textit{{NeurIPS, 30}}.\\Saito, T., \& Rehmsmeier, M. (2015). Precision--recall plots for imbalanced data. \textit{{PLOS ONE, 10}}(3).
\newpage
\section*{{Matriz de evidencias}}\small
\begin{{tabularx}}{{\linewidth}}{{p{{2.5cm}}p{{2.7cm}}X X}}\toprule Evidencia&Ubicación&Conclusión&Limitación\\\midrule
Integridad&Sección 1, Fig. 1&70/15/15 temporal y train-only&Identidad aproximada\\
Comparación A/B&Tabla A/B/C&A domina AP, F1 y costo&A posee más ingeniería\\
Valor del orden&Fig. 2&Permutar no reduce AP&Una GRU y una clave\\
Apuesta C&Sección 2&No alcanza criterio previo&Fusión logística específica\\
Decisión económica&Sección 4, Fig. 3&A minimiza costo interno&Costos y volumen académicos\\
Recomendación&Sección 5&Complementar con revisión&Requiere cohorte nueva\\\bottomrule\end{{tabularx}}

\vfill\textbf{{Declaración de IA.}} Se usó asistencia para código, redacción, visualización y auditoría. Los autores ejecutaron y verificaron datos, métricas, falsificaciones y conclusiones.\end{{document}}"""
    out_dir = ROOT / "entregables" / "informe" / "v5"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / "informe.tex"
    out.write_text(tex, encoding="utf-8", newline="\n")
    pdflatex = shutil.which("pdflatex")
    if pdflatex:
        for _ in range(2):
            subprocess.run([pdflatex, "-interaction=nonstopmode", "-halt-on-error", "-output-directory", str(out_dir), str(out)], cwd=out_dir, check=False, capture_output=True)
    return out


def build_slides() -> Path:
    """Delega la presentación ampliada al constructor visual V5."""

    from presentation_v5 import build_presentation

    return build_presentation(ROOT, R, export_pdf=True)


def build_ficha() -> None:
    out_dir = ROOT / "entregables" / "ficha" / "v5"
    out_dir.mkdir(parents=True, exist_ok=True)
    qr_path = ROOT / "evidencia" / "recursos" / "qr_repositorio_v5.png"
    qr_path.parent.mkdir(parents=True, exist_ok=True)
    qrcode.make(URL).save(qr_path)
    doc = Document()
    section = doc.sections[0]
    section.top_margin=Inches(.42); section.bottom_margin=Inches(.42); section.left_margin=Inches(.55); section.right_margin=Inches(.55)
    title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; run=title.add_run("PROYECTO 1 · MONITOREO TRANSACCIONAL · V5"); run.bold=True; run.font.size=Pt(19); run.font.color.rgb=RGBColor(24,78,119)
    sub=doc.add_paragraph("Universidad del Valle de Guatemala · Deep Learning y Sistemas Inteligentes · Sección 30"); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    names=doc.add_paragraph("Wilson Alejandro Calderón Argueta · 22018 | Pablo Daniel Barillas Moreno · 22193"); names.alignment=WD_ALIGN_PARAGRAPH.CENTER
    table=doc.add_table(rows=2,cols=4); table.style="Table Grid"
    values=[('AUC-PR',f(BENCH['A_V4']['auc_pr'])),('ROC-AUC',f(BENCH['A_V4']['roc_auc'])),('Precisión',pct(BENCH['A_V4']['precision'])),('Recall',pct(BENCH['A_V4']['recall'])),('F1',f(BENCH['A_V4']['f1'])),('Costo',money(BENCH['A_V4']['cost_q'])),('Orden ΔAP',f(FALS['order_auc_pr_drop'],4)),('Candidato','A · V4')]
    for cell,(label,value) in zip([c for row in table.rows for c in row.cells],values):
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run(label.upper()+"\n"); rr.bold=True; rr.font.size=Pt(8); vv=p.add_run(value); vv.bold=True; vv.font.size=Pt(12); vv.font.color.rgb=RGBColor(42,157,143)
    doc.add_heading("Resumen",level=1); doc.add_paragraph("V5 integra A LightGBM, B GRU y C stacking bajo bloques temporales comunes. A minimiza costo. Permutar la historia no reduce AUC-PR, por lo que no se atribuye valor material al orden.")
    doc.add_heading("Decisión",level=1); doc.add_paragraph("Conservar A para priorizar revisión humana. El benchmark es histórico reutilizado; se requiere una cohorte nueva antes de producción.")
    pic=doc.add_paragraph(); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER; pic.add_run().add_picture(str(qr_path),width=Inches(1.15)); link=doc.add_paragraph(URL); link.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.save(out_dir / "Ficha_Repositorio_Proyecto_1_V5.docx")

    styles=getSampleStyleSheet(); pdf=SimpleDocTemplate(str(out_dir/"Ficha_Repositorio_Proyecto_1_V5.pdf"),pagesize=letter,rightMargin=.55*inch,leftMargin=.55*inch,topMargin=.45*inch,bottomMargin=.45*inch)
    story=[Paragraph("<b>PROYECTO 1 · MONITOREO TRANSACCIONAL · V5</b>",styles['Title']),Paragraph("Universidad del Valle de Guatemala · Grupo 1 · Sección 30",styles['Heading3']),Paragraph("Wilson Alejandro Calderón Argueta · 22018 | Pablo Daniel Barillas Moreno · 22193",styles['BodyText']),Spacer(1,10)]
    data=[["AP A",f(BENCH['A_V4']['auc_pr']),"Recall",pct(BENCH['A_V4']['recall'])],["F1",f(BENCH['A_V4']['f1']),"Costo",money(BENCH['A_V4']['cost_q'])],["ΔAP orden",f(FALS['order_auc_pr_drop'],4),"Candidato","A · V4"]]
    tab=Table(data,colWidths=[1.1*inch,1.35*inch,1.1*inch,1.35*inch]); tab.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,-1),colors.HexColor('#EDF5FB')),('GRID',(0,0),(-1,-1),.5,colors.HexColor('#9FB7C8')),('FONTNAME',(0,0),(-1,-1),'Helvetica-Bold'),('ALIGN',(0,0),(-1,-1),'CENTER'),('PADDING',(0,0),(-1,-1),8)]))
    qr=RLImage(str(qr_path),width=1.12*inch,height=1.12*inch); qr.hAlign="CENTER"
    story.extend([tab,Spacer(1,14),Paragraph("A/B/C fueron comparados con la misma partición y horizonte. La permutación no muestra valor material del orden; se conserva A para revisión humana.",styles['BodyText']),Spacer(1,8),Paragraph("Benchmark histórico reutilizado. Requiere cohorte temporal nueva, identidad fiable y controles productivos.",styles['BodyText']),Spacer(1,10),qr,Spacer(1,6),Paragraph(URL,styles['Heading3'])]); pdf.build(story)


def main() -> None:
    build_readme()
    notebook = build_notebook()
    report = build_report()
    slides = build_slides()
    build_ficha()
    print("Entregables V5 generados desde", RESULT_PATH)
    print("Notebook:", notebook, "Informe:", report, "Presentación:", slides)


if __name__ == "__main__":
    main()
