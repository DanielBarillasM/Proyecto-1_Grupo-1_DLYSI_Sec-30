"""Genera todos los entregables rubricados de Proyecto 1 V6 desde resultados_v6.json."""

from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path

import nbformat as nbf
import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[2]
RESULT_PATH = ROOT / "artefactos" / "v6" / "resultados_v6.json"
R = json.loads(RESULT_PATH.read_text(encoding="utf-8"))
INTERNAL = R["evaluacion_interna"]
BENCH = R["benchmark_historico"]
CANDIDATE = R["candidato"]["modelo"]
FALS = R["falsificaciones"]
URL = "https://github.com/DanielBarillasM/Proyecto-1_Grupo-1_DLYSI_Sec-30"


def f(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def pct(value: float, digits: int = 2) -> str:
    return f"{100 * value:.{digits}f}%"


def money(value: float) -> str:
    return f"Q{value:,.0f}"


def notebook_style() -> str:
    return """<style>
:root{--navy:#102a43;--blue:#184e77;--teal:#2a9d8f;--gold:#e9c46a;--red:#e76f51;--ink:#172033;--pale:#edf5fb}
.hero{padding:38px 42px;border-radius:24px;color:#f8fbff;background:radial-gradient(circle at 92% 8%,rgba(255,255,255,.16) 0 8%,transparent 9%),linear-gradient(125deg,var(--navy),var(--blue) 55%,var(--teal));box-shadow:0 16px 38px #102a433d;font-family:Inter,'Segoe UI',sans-serif}.hero h1{font-size:38px;margin:18px 0 8px;color:white}.chips{display:flex;gap:8px;flex-wrap:wrap}.chips span{padding:6px 13px;border:1px solid #ffffff55;border-radius:999px;background:#ffffff20;font-size:11px;font-weight:800}.grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(205px,1fr));gap:13px;margin-top:24px}.card{padding:14px 16px;border:1px solid #ffffff44;border-radius:12px;background:#ffffff12}.section{margin:28px 0 14px;padding:16px 22px;border-radius:14px;background:linear-gradient(90deg,var(--navy),var(--blue));color:white}.section h2{margin:0;color:white}.call{padding:18px 22px;margin:14px 0;border:1px solid #c9d9e6;border-left:6px solid var(--teal);border-radius:13px;background:var(--pale);color:var(--ink);line-height:1.7}.warn{border-left-color:var(--red);background:#fff2ed}.proof{border-left-color:var(--gold);background:#fff9e9}.metric{display:inline-block;min-width:145px;padding:12px 14px;margin:4px;border-radius:10px;background:#e5f5f2;color:#184e77;text-align:center}.metric b{display:block;font-size:21px}table{font-size:14px}th{background:#184e77!important;color:white!important;text-align:left!important}td,th{padding:9px!important}code{background:#102a4312;padding:2px 5px;border-radius:4px}
</style>"""


def build_readme() -> str:
    center = R["economia_mensual"][CANDIDATE]["12"]
    text = f"""<div align="center">

# Proyecto 1 · Monitoreo transaccional · V6 integrada

### ¿El orden aporta información más allá de los agregados?

![Python](https://img.shields.io/badge/Python-3.13-3776AB?logo=python&logoColor=white)
![PyTorch](https://img.shields.io/badge/PyTorch-2.13-EE4C2C?logo=pytorch&logoColor=white)
![LightGBM](https://img.shields.io/badge/LightGBM-4.7-184e77)
![Estado](https://img.shields.io/badge/Candidato-A__V4-2a9d8f)

**Wilson Alejandro Calderón Argueta · 22018** · **Pablo Daniel Barillas Moreno · 22193**

Universidad del Valle de Guatemala · Deep Learning y Sistemas Inteligentes · Sección 30 · 2026

</div>

> [!IMPORTANT]
> El último 15 % de IEEE-CIS ya fue observado en iteraciones anteriores y se reporta como benchmark temporal histórico reutilizado. Todas las decisiones V6 se toman dentro de validación. Una promoción confirmatoria exige una cohorte nueva.

## Resumen ejecutivo

El proyecto estudia {R['datos']['filas']:,} transacciones IEEE-CIS, con {R['datos']['fraudes']:,} fraudes y prevalencia {pct(R['datos']['prevalencia'])}. Compara una línea tabular competitiva sin orden (A), GRU/TCN causales sobre hasta 32 eventos (B), una fusión condicionada (C) y un encoder–decoder entrenado solo con transacciones legítimas (D). A obtiene AP interna {f(INTERNAL['A']['auc_pr'],4)}, B {f(INTERNAL['B']['auc_pr'],4)}, C {f(INTERNAL['C']['auc_pr'],4)} y D {f(INTERNAL['D']['auc_pr'],4)}.

La permutación controlada no perjudica a B: su AP cambia de {f(FALS['original_internal']['auc_pr'],4)} a {f(FALS['permutation_mean_auc_pr'],4)} ± {f(FALS['permutation_std_auc_pr'],4)}. La diferencia original−permutada es {f(FALS['order_auc_pr_drop'],4)}. Con esta evidencia no se afirma que el orden aporte. C tampoco supera su criterio previo: cambio AP {f(R['hipotesis_C']['auc_pr_gain'],4)} y reducción de costo {pct(R['hipotesis_C']['cost_reduction'])}. El candidato es A.

## Datos y protocolo temporal

Se utiliza la competencia pública [IEEE-CIS Fraud Detection](https://www.kaggle.com/competitions/ieee-fraud-detection/overview) de Kaggle, con datos anonimizados suministrados por Vesta Corporation. `train_transaction.csv` y `train_identity.csv` se unen por `TransactionID`, las filas se ordenan por `TransactionDT` y ambos campos se excluyen como magnitudes predictivas. La identidad secuencial es una clave aproximada formada por `card1`, `card2`, `card3`, `card5` y `addr1`.

La separación es 70 % entrenamiento, 15 % validación y 15 % benchmark histórico. Validación se subdivide cronológicamente en early stopping, ajuste de C, calibración, umbral y evaluación interna. Imputación, escalado y vocabularios se aprenden exclusivamente con entrenamiento. Las características históricas utilizan solo eventos anteriores; no se aplican particiones aleatorias.

## EDA, correlación y PCA

La V6 entrega un cuaderno exploratorio ejecutado que examina las 434 columnas de la unión teórica, la cobertura parcial de identidad, faltantes extremos, variables constantes, cambios temporales y asociación univariada con `isFraud`. También compara varias definiciones de entidad y cuantifica qué proporción de transacciones dispone de 3, 8, 16 o 32 eventos. Esta evidencia explica por qué ampliar una secuencia sin mejorar la identidad puede añadir ruido en lugar de memoria útil.

La correlación de Spearman se usa para localizar familias redundantes, especialmente dentro de `V1–V339`, pero no para eliminar automáticamente todo par correlacionado. Los árboles pueden aprovechar umbrales e interacciones diferentes incluso entre variables similares; por ello cada reducción debe validarse temporalmente. PCA se estudia como compresión del bloque V: resume gran parte de su varianza con muchas menos componentes, pero una varianza reconstruida alta no garantiza conservar la señal de fraude minoritaria. Como la ablation V3 con PCA rindió peor, V6 conserva PCA como diagnóstico y no como transformación del candidato.

Las nuevas características priorizan significado operativo y causalidad: conteos y monto medio por entidad en 1, 6, 24 y 72 horas; tiempo desde el evento anterior; monto relativo al historial; cambios de dispositivo/dirección; cantidad de faltantes; variables `C`, `D`, `V` e identidad seleccionadas. `TransactionID` y `TransactionDT` permanecen fuera del vector predictivo.

## Modelos A/B/C y control D

| Pieza | Diseño | Resultado interno |
|---|---|---:|
| A | LightGBM V4 con expertos `ProductCD=W/NO-W` | AP {f(INTERNAL['A']['auc_pr'],4)} · costo {money(INTERNAL['A']['cost_q'])} |
| B | GRU frente a TCN causal; seleccionada `{R['seleccion']['B']['seleccionado']}`, hasta 32 eventos | AP {f(INTERNAL['B']['auc_pr'],4)} · costo {money(INTERNAL['B']['cost_q'])} |
| C | Regresión logística condicionada sobre A/B/D, monto, historia e identidad | AP {f(INTERNAL['C']['auc_pr'],4)} · costo {money(INTERNAL['C']['cost_q'])} |
| D | Encoder–decoder PyTorch entrenado solo con legítimas | AP {f(INTERNAL['D']['auc_pr'],4)} · costo {money(INTERNAL['D']['cost_q'])} |

A conserva `{R['seleccion']['A']['seleccionado']}` porque el refuerzo LightGBM no fue estable en las dos subventanas de selección. B compara dos arquitecturas con BCE ponderada, AdamW, clipping y early stopping. D minimiza MSE de reconstrucción legítima; su alta tasa de anomalías demuestra que anomalía y fraude no son equivalentes. C se entrena en un bloque independiente y no recibe el benchmark para decidir su arquitectura.

## Hipótesis y falsificaciones

**Hipótesis previa de C:** {R['hipotesis_C']['declaracion_previa']}

La hipótesis no se cumple. C pierde {abs(R['hipotesis_C']['auc_pr_gain']):.4f} de AP y aumenta el costo {abs(100*R['hipotesis_C']['cost_reduction']):.2f} % frente a A en evaluación interna.

Las dos pruebas obligatorias son:

1. Permutación de antecedentes con cinco semillas, manteniendo la transacción objetivo al final.
2. Recorte de la historia a 3, 8 y 16 eventos.

La historia de 3 eventos obtiene AP {f(FALS['historia_3']['auc_pr'],4)}, la de 8 obtiene {f(FALS['historia_8']['auc_pr'],4)} y la de 16 obtiene {f(FALS['historia_16']['auc_pr'],4)}. Ninguna evidencia justifica afirmar que el orden mejore el detector.

## Resultados y decisión económica

| Modelo | AUC-PR benchmark | Precisión | Recall | F1 | Costo |
|---|---:|---:|---:|---:|---:|
| **A** | **{f(BENCH['A']['auc_pr'])}** | **{pct(BENCH['A']['precision'])}** | {pct(BENCH['A']['recall'])} | **{f(BENCH['A']['f1'])}** | **{money(BENCH['A']['cost_q'])}** |
| B | {f(BENCH['B']['auc_pr'])} | {pct(BENCH['B']['precision'])} | {pct(BENCH['B']['recall'])} | {f(BENCH['B']['f1'])} | {money(BENCH['B']['cost_q'])} |
| C | {f(BENCH['C']['auc_pr'])} | {pct(BENCH['C']['precision'])} | **{pct(BENCH['C']['recall'])}** | {f(BENCH['C']['f1'])} | {money(BENCH['C']['cost_q'])} |
| D | {f(BENCH['D']['auc_pr'])} | {pct(BENCH['D']['precision'])} | {pct(BENCH['D']['recall'])} | {f(BENCH['D']['f1'])} | {money(BENCH['D']['cost_q'])} |

La política de umbral minimiza $4200FN+180FP$ sujeta a recall ≥ 0.75 en selección. El umbral de A es {R['umbrales']['A']:.5f}. En el escenario central de 12 transacciones por tarjeta al mes, A representa un costo mensual proyectado de {money(center['costo_mensual_q'])}. Es una extrapolación académica, no una cifra contable.

### Cómo interpretar las métricas

AP {f(BENCH['A']['auc_pr'],3)} significa que A mantiene una relación precisión–recall muy superior a la prevalencia de {pct(R['datos']['prevalencia'])}; no significa que {100*BENCH['A']['auc_pr']:.1f} % de sus alertas sea correcto. Esa proporción puntual la expresa la precisión: {pct(BENCH['A']['precision'])}. El recall {pct(BENCH['A']['recall'])} indica que el umbral recupera cerca de ocho de cada diez fraudes, mientras que ROC-AUC {f(BENCH['A']['roc_auc'],3)} describe la probabilidad de ordenar un fraude por encima de una transacción legítima elegida al azar. Debido al desbalance, un ROC alto puede coexistir con miles de falsas alarmas; por eso AP, alertas/100k y costo acompañan siempre a ROC.

D ilustra el mismo punto desde otro ángulo: alcanza recall {pct(BENCH['D']['recall'])}, pero precisión de solo {pct(BENCH['D']['precision'])}. El encoder–decoder reconoce rareza, no fraude: operaciones legítimas poco frecuentes, deriva o patrones con muchos faltantes también reconstruyen mal. C aparece descriptivamente competitivo en el benchmark, pero no se promueve porque su hipótesis se rechazó en evaluación interna. Reabrir esa decisión con el período final sería seleccionar con el test reutilizado.

## Tres decisiones técnicas importantes

1. **A tabular V4 frente al HistGradientBoosting V1.** Se consideró conservar el baseline antiguo. Se eligió LightGBM con expertos porque usa más variables causales y obtiene AP y costo claramente mejores.
2. **GRU frente a TCN causal.** Ambas se entrenaron con la misma población; TCN ganó en `model_select`, aunque la permutación mostró que su ranking no depende favorablemente del orden.
3. **Fusión y anomalías.** Se añadió un encoder–decoder legítimo como D y como entrada de C. Se conserva como ablation porque su baja precisión y alto costo no justifican promoverlo.

## Reproducción

```powershell
py -3.13 -m venv .venv
.\.venv\Scripts\Activate.ps1
python -m pip install -r configuracion/v6/requirements-v6.txt
python -m pip install -r configuracion/v6/requirements-docs-v6.txt
python codigo/compartido/download_data.py
python -u codigo/v6/proyecto1_v6_pipeline.py
python codigo/v6/build_v6_deliverables.py
jupyter nbconvert --to notebook --execute --inplace --ExecutePreprocessor.timeout=300 entregables/cuaderno/v6/proyecto1_calderon_barillas.ipynb
python codigo/v6/audit_project1_v6.py
```

La descarga requiere aceptar las reglas de IEEE-CIS en Kaggle y configurar las credenciales fuera del repositorio. Los CSV de casi 700 MB no se versionan.

## Estructura

```text
codigo/v6/              pipeline, construcción y auditoría de V6
configuracion/v6/       versiones exactas e instrucciones
datos/raw/              CSV locales ignorados por Git
artefactos/v6/          pesos A/B/C/D, calibradores, contrato y puntajes
evidencia/figuras/v6/   resultados reproducibles
entregables/cuaderno/v6/ notebook ejecutado
entregables/informe/v6/  fuente LaTeX y PDF
entregables/presentacion/v6/ HTML y PDF de ocho diapositivas
entregables/ficha/v6/    ficha del repositorio
```

## Candidato al Proyecto Final

- **Modelo conservado:** A — LightGBM V4 con expertos por `ProductCD` y calibrador V6.
- **Artefactos:** `artefactos/v4/modelo_experto_w_v4.txt`, `modelo_experto_no_w_v4.txt`, `artefactos/v6/calibradores_v6.joblib` y `contrato_entrada_salida_v6.json`.
- **Usuario:** analista de riesgo o equipo de monitoreo transaccional.
- **Decisión:** ordenar alertas y priorizar revisión; el puntaje no prueba fraude ni autoriza bloqueo autónomo.
- **Entrada preliminar:** transacción actual, variables categóricas y estadísticas históricas causales especificadas por el contrato.
- **Salida:** `risk_score` continuo en [0,1], umbral {R['umbrales']['A']:.5f} y política de revisión.
- **Pendientes:** nueva cohorte etiquetada, identidad bancaria fiable, costos reales, latencia, privacidad, equidad, seguridad, explicaciones y monitoreo.

## Limitaciones y uso responsable

IEEE-CIS está anonimizado y cubre aproximadamente 182 días. La clave proxy no equivale a un cliente real. El benchmark ya fue observado y ninguna conclusión se presenta como confirmación externa. Los costos y volúmenes mensuales son escenarios. El sistema debe apoyar revisión humana, no atribuir culpabilidad ni bloquear de manera autónoma.

## Declaración de uso de inteligencia artificial

Se utilizó asistencia de IA para estructurar y revisar código, diseñar documentación HTML/CSS/LaTeX, localizar bibliografía y automatizar auditorías. Los integrantes ejecutaron el pipeline y verificaron particiones, alineación de IDs, métricas, falsificaciones, umbrales y artefactos. La IA no se utilizó como fuente académica ni reemplaza la defensa de las decisiones.

## Referencias APA 7

Cho, K., et al. (2014). Learning phrase representations using RNN encoder–decoder for statistical machine translation. *Proceedings of EMNLP*, 1724–1734. https://doi.org/10.3115/v1/D14-1179

IEEE Computational Intelligence Society. (2019). *IEEE-CIS Fraud Detection* [Data set]. Kaggle. https://www.kaggle.com/competitions/ieee-fraud-detection/overview

Ke, G., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. *Advances in Neural Information Processing Systems, 30*.

Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. *PLOS ONE, 10*(3), e0118432. https://doi.org/10.1371/journal.pone.0118432

Zhou, C., & Paffenroth, R. C. (2017). Anomaly detection with robust deep autoencoders. *Proceedings of the 23rd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining*, 665–674. https://doi.org/10.1145/3097983.3098052
"""
    out = ROOT / "entregables" / "ficha" / "v6" / "README_V6_GENERADO.md"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(text, encoding="utf-8", newline="\n")
    (ROOT / "README.md").write_text(text, encoding="utf-8", newline="\n")
    return text


def build_report() -> Path:
    """Compila el informe V6 canónico sin volver a generar una plantilla antigua."""

    from materialize_report_v6 import main as materialize_report
    from report_v6 import compile_report

    materialize_report()
    return compile_report(ROOT)


def build_slides() -> Path:
    """Delega la presentación ampliada al constructor visual V6."""

    from presentation_v6 import build_presentation

    return build_presentation(ROOT, R, export_pdf=True)


def build_ficha() -> None:
    out_dir = ROOT / "entregables" / "ficha" / "v6"
    out_dir.mkdir(parents=True, exist_ok=True)
    qr_path = ROOT / "evidencia" / "recursos" / "v6" / "qr_repositorio_v6.png"
    qr_path.parent.mkdir(parents=True, exist_ok=True)
    qrcode.make(URL).save(qr_path)
    doc = Document()
    section = doc.sections[0]
    section.top_margin=Inches(.42); section.bottom_margin=Inches(.42); section.left_margin=Inches(.55); section.right_margin=Inches(.55)
    title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; run=title.add_run("PROYECTO 1 · MONITOREO TRANSACCIONAL · V6"); run.bold=True; run.font.size=Pt(19); run.font.color.rgb=RGBColor(24,78,119)
    sub=doc.add_paragraph("Universidad del Valle de Guatemala · Deep Learning y Sistemas Inteligentes · Sección 30"); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    names=doc.add_paragraph("Wilson Alejandro Calderón Argueta · 22018 | Pablo Daniel Barillas Moreno · 22193"); names.alignment=WD_ALIGN_PARAGRAPH.CENTER
    table=doc.add_table(rows=2,cols=4); table.style="Table Grid"
    values=[('AUC-PR',f(BENCH['A']['auc_pr'])),('ROC-AUC',f(BENCH['A']['roc_auc'])),('Precisión',pct(BENCH['A']['precision'])),('Recall',pct(BENCH['A']['recall'])),('F1',f(BENCH['A']['f1'])),('Costo',money(BENCH['A']['cost_q'])),('Orden ΔAP',f(FALS['order_auc_pr_drop'],4)),('Candidato','A · V4')]
    for cell,(label,value) in zip([c for row in table.rows for c in row.cells],values):
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run(label.upper()+"\n"); rr.bold=True; rr.font.size=Pt(8); vv=p.add_run(value); vv.bold=True; vv.font.size=Pt(12); vv.font.color.rgb=RGBColor(42,157,143)
    doc.add_heading("Fuente de datos",level=1); doc.add_paragraph("IEEE-CIS Fraud Detection · Vesta Corporation · Kaggle\nhttps://www.kaggle.com/competitions/ieee-fraud-detection/overview")
    doc.add_heading("Resumen",level=1); doc.add_paragraph("V6 compara A tabular, B GRU/TCN, C fusión y D encoder–decoder bajo bloques temporales comunes. A se conserva. Permutar la historia no reduce AP y D produce demasiadas falsas alarmas.")
    doc.add_heading("Decisión",level=1); doc.add_paragraph("Conservar A para priorizar revisión humana. El benchmark es histórico reutilizado; se requiere una cohorte nueva antes de producción.")
    pic=doc.add_paragraph(); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER; pic.add_run().add_picture(str(qr_path),width=Inches(1.15)); link=doc.add_paragraph(URL); link.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.save(out_dir / "Ficha_Repositorio_Proyecto_1_V6.docx")

    styles=getSampleStyleSheet(); pdf=SimpleDocTemplate(str(out_dir/"Ficha_Repositorio_Proyecto_1_V6.pdf"),pagesize=letter,rightMargin=.55*inch,leftMargin=.55*inch,topMargin=.45*inch,bottomMargin=.45*inch)
    story=[Paragraph("<b>PROYECTO 1 · MONITOREO TRANSACCIONAL · V6</b>",styles['Title']),Paragraph("Universidad del Valle de Guatemala · Grupo 1 · Sección 30",styles['Heading3']),Paragraph("Wilson Alejandro Calderón Argueta · 22018 | Pablo Daniel Barillas Moreno · 22193",styles['BodyText']),Spacer(1,10)]
    data=[["AP A",f(BENCH['A']['auc_pr']),"Recall",pct(BENCH['A']['recall'])],["F1",f(BENCH['A']['f1']),"Costo",money(BENCH['A']['cost_q'])],["ΔAP orden",f(FALS['order_auc_pr_drop'],4),"Candidato","A · V4"]]
    tab=Table(data,colWidths=[1.1*inch,1.35*inch,1.1*inch,1.35*inch]); tab.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,-1),colors.HexColor('#EDF5FB')),('GRID',(0,0),(-1,-1),.5,colors.HexColor('#9FB7C8')),('FONTNAME',(0,0),(-1,-1),'Helvetica-Bold'),('ALIGN',(0,0),(-1,-1),'CENTER'),('PADDING',(0,0),(-1,-1),8)]))
    qr=RLImage(str(qr_path),width=1.12*inch,height=1.12*inch); qr.hAlign="CENTER"
    story.extend([tab,Spacer(1,12),Paragraph("Datos: IEEE-CIS Fraud Detection, Vesta Corporation, Kaggle.",styles['BodyText']),Paragraph("A/B/C/D fueron comparados con la misma partición y horizonte. La permutación no muestra valor material del orden; el autoencoder genera demasiadas falsas alarmas; se conserva A.",styles['BodyText']),Spacer(1,8),Paragraph("Benchmark histórico reutilizado. Requiere cohorte temporal nueva, identidad fiable y controles productivos.",styles['BodyText']),Spacer(1,10),qr,Spacer(1,6),Paragraph(URL,styles['Heading3'])]); pdf.build(story)


def build_notebook_v6() -> Path:
    """Construye el cuaderno oficial V6 con narrativa derivada del resultado."""
    out_dir = ROOT / "entregables" / "cuaderno" / "v6"
    out_dir.mkdir(parents=True, exist_ok=True)
    candidate = R["candidato"]["modelo"]
    selected_a = R["seleccion"]["A"]["seleccionado"]
    selected_b = R["seleccion"]["B"]["seleccionado"]
    order_drop = FALS["order_auc_pr_drop"]
    order_text = "sí alcanzó" if FALS["orden_material"] else "no alcanzó"
    c_text = "cumplió" if R["hipotesis_C"]["success"] else "no cumplió"
    hero = notebook_style() + f"""<div class="hero"><div class="chips"><span>DEEP LEARNING</span><span>PROYECTO 1</span><span>V6</span><span>GRUPO 1 · SECCIÓN 30</span></div><h1>Monitoreo transaccional</h1><p style="font-size:20px">Datos, modelos A/B/C, encoder–decoder y evidencia del valor del orden</p><div class="grid"><div class="card"><b>Institución</b><br>Universidad del Valle de Guatemala</div><div class="card"><b>Curso</b><br>Deep Learning y Sistemas Inteligentes</div><div class="card"><b>Docente</b><br>Kevin Recinos</div><div class="card"><b>Integrantes</b><br>Wilson Alejandro Calderón Argueta · 22018<br>Pablo Daniel Barillas Moreno · 22193</div><div class="card"><b>Datos</b><br>IEEE-CIS · {R['datos']['filas']:,} transacciones</div><div class="card"><b>Candidato V6</b><br>{candidate}</div></div></div>"""
    cells = [nbf.v4.new_markdown_cell(hero)]
    cells.append(nbf.v4.new_code_cell("""from pathlib import Path
import json, pandas as pd, numpy as np
from sklearn.metrics import average_precision_score, roc_auc_score
from IPython.display import Image, display, HTML
ROOT=Path.cwd().resolve()
while not (ROOT/'artefactos').exists() and ROOT != ROOT.parent: ROOT=ROOT.parent
R=json.loads((ROOT/'artefactos/v6/resultados_v6.json').read_text(encoding='utf-8'))
VAL=pd.read_csv(ROOT/'artefactos/v6/predicciones_validacion_v6.csv')
BENCH=pd.read_csv(ROOT/'artefactos/v6/predicciones_benchmark_v6.csv')
print('Fuente única cargada · candidato:',R['candidato']['modelo'])"""))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>1 · Fuente, alcance e integridad</h2></div>

Los datos proceden de la competencia pública [IEEE-CIS Fraud Detection](https://www.kaggle.com/competitions/ieee-fraud-detection/overview), publicada en Kaggle con datos transaccionales anonimizados proporcionados por Vesta Corporation. Se unen `train_transaction.csv` y `train_identity.csv` mediante `TransactionID` y se ordenan por `TransactionDT`; ambos campos funcionan como llave y reloj, no como magnitudes predictivas.

<div class="call warn"><b>Alcance inferencial.</b> El 15 % final es un benchmark temporal histórico reutilizado después de V1–V5. Las decisiones V6 se toman dentro de validación; una afirmación confirmatoria exige una cohorte futura.</div>

Para toda variable histórica se exige causalidad:

$$x_t^{hist}=f(\{x_j:t_j<t\}).$$"""))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(R['datos']['particiones']).T"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>2 · Diagnóstico exploratorio y selección de variables</h2></div>

El EDA ejecutado se entrega en `EDA_IEEE_CIS_Diagnostico_Datos_V6.ipynb`. Revisa faltantes, constantes, deriva temporal, correlación de Spearman, asociación con fraude, cobertura de identidades proxy y PCA sobre el bloque `V`. Los IDs se excluyen como predictores. La correlación y PCA son herramientas de diagnóstico, no filtros automáticos: dos variables correlacionadas pueden aportar interacciones diferentes a árboles y conservar varianza no garantiza conservar señal discriminativa.

La V6 amplía variables `C`, `D`, `V`, identidad, faltantes, frecuencia y velocidad en 1/6/24/72 horas. Toda estadística de entidad excluye la transacción actual."""))
    cells.append(nbf.v4.new_code_cell("pd.Series(R['secuencias']).to_frame('valor')"))
    cells.append(nbf.v4.new_markdown_cell(f"""<div class="section"><h2>3 · Núcleo comparable A/B y controles</h2></div>

| Pieza | Diseño V6 | Papel experimental |
|---|---|---|
| **A** | Baseline tabular sin orden; seleccionado: `{selected_a}` | Establece cuánto se logra con variables actuales y agregados causales |
| **B** | GRU frente a TCN causal; seleccionado: `{selected_b}` | Puede explotar transiciones en hasta 32 eventos |
| **C** | Fusión A/B/D condicionada por identidad y longitud | Apuesta predeclarada de complementariedad |
| **D** | Encoder–decoder PyTorch entrenado solo con legítimas | Error de reconstrucción como score de anomalía |
| **Control** | Regresión logística | Verifica si la complejidad tabular es necesaria |

Todos producen puntajes continuos sobre las mismas filas, horizonte y etiqueta. Imputación, vocabularios y escalado se ajustan exclusivamente con entrenamiento."""))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(R['seleccion']).T"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>4 · Encoder–decoder para la clase minoritaria</h2></div>

D aprende una representación latente de transacciones no fraudulentas y minimiza:

$$\mathcal{L}_{AE}=\frac{1}{p}\sum_{j=1}^{p}(x_j-\hat{x}_j)^2.$$

Una transacción recibe el error $e(x)$ y este se transforma en percentil respecto de errores legítimos de entrenamiento. Un percentil alto indica que la observación se aleja del patrón normal; no implica fraude por sí solo. Esta estrategia ayuda solo si las anomalías coinciden con fraudes y no con cambios legítimos, faltantes o deriva."""))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(R['entrenamiento_D'])"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>5 · Hipótesis previa, calibración y umbral</h2></div>

<div class="call proof"><b>Hipótesis C.</b> """ + R["hipotesis_C"]["declaracion_previa"] + """</div>

Los scores se calibran en un bloque independiente y el umbral minimiza:

$$C(\tau)=Q4{,}200\,FN(\tau)+Q180\,FP(\tau),$$

sujeto al objetivo de recall declarado. AP es la métrica principal porque el fraude es minoritario; ROC-AUC se reporta como discriminación global, pero puede ocultar muchas falsas alarmas."""))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(R['calibracion']).T"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>6 · Resultados comunes</h2></div>

La precisión indica qué proporción de alertas es fraude; el recall, qué proporción de fraudes se detecta; F1 equilibra ambas sin incorporar quetzales. AP resume la curva precisión–recall a través de umbrales y no debe confundirse con precisión puntual."""))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(R['evaluacion_interna']).T[['auc_pr','roc_auc','precision','recall','f1','tp','fp','fn','cost_q','alertas_por_100k','threshold']]"))
    cells.append(nbf.v4.new_code_cell("""eval_rows=np.arange(int(len(VAL)*.80),len(VAL))
for name in ['A','B','C','D']:
    measured=average_precision_score(VAL.y.to_numpy()[eval_rows],VAL[f'score_{name}'].to_numpy()[eval_rows])
    assert abs(measured-R['evaluacion_interna'][name]['auc_pr'])<1e-10
print('Coherencia verificada: AP reproducida desde los scores guardados.')"""))
    cells.append(nbf.v4.new_code_cell("display(Image(filename=str(ROOT/'evidencia/figuras/v6/01_comparacion_abc_validacion.png')))"))
    cells.append(nbf.v4.new_markdown_cell(f"""<div class="section"><h2>7 · Dos intentos de falsificar el valor del orden</h2></div>

Se barajan únicamente los antecedentes, nunca la transacción objetivo, con cinco semillas. Además se recorta la historia a 3, 8 y 16 eventos. La caída original–permutada fue `{order_drop:.4f}` AP y {order_text} el umbral material predeclarado de 0.01. Esta prueba distingue “una red que procesa secuencias” de “un modelo cuya predicción depende realmente del orden”."""))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame({'original':R['falsificaciones']['original_internal'],'permutada':{'auc_pr':R['falsificaciones']['permutation_mean_auc_pr']},'historia_3':R['falsificaciones']['historia_3'],'historia_8':R['falsificaciones']['historia_8'],'historia_16':R['falsificaciones']['historia_16']}).T"))
    cells.append(nbf.v4.new_code_cell("display(Image(filename=str(ROOT/'evidencia/figuras/v6/03_falsificaciones_orden_v6.png')))"))
    cells.append(nbf.v4.new_markdown_cell(f"""<div class="section"><h2>8 · Decisión económica y veredicto</h2></div>

C {c_text} el criterio previo. El candidato `{candidate}` fue elegido por la regla congelada en evaluación interna; las cifras del benchmark son descriptivas y no reabren la selección. Un ROC-AUC alto puede coexistir con baja precisión cuando la clase negativa domina: por eso se muestran alertas por 100,000 y costo, no solo métricas de ranking."""))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(R['benchmark_historico']).T[['auc_pr','roc_auc','precision','recall','f1','tp','fp','fn','cost_q','alertas_por_100k']]"))
    cells.append(nbf.v4.new_code_cell("display(Image(filename=str(ROOT/'evidencia/figuras/v6/04_costos_abc_v6.png')))"))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>9 · Matriz de evidencias y conclusión</h2></div>

| Evidencia rubricada | Resultado verificable | Decisión o límite |
|---|---|---|
| Integridad | 70/15/15 temporal; transformaciones train-only | Benchmark reutilizado, no ciego |
| A/B común | Misma población, horizonte y scores continuos | AP principal; costo en umbral |
| Orden | 5 permutaciones + recortes 3/8/16 | Solo se atribuye valor con caída material |
| C | Regla +0.01 AP, −5 % costo y recall ≥0.75 | Veredicto no reescrito con benchmark |
| D | Reconstrucción de legítimas | Anomalía no equivale automáticamente a fraude |
| Economía | Q4,200 FN y Q180 FP | Escenario académico sensible al volumen |

La recomendación final debe leerse junto con limitaciones: identidad proxy, anonimización, deriva, benchmark reutilizado y ausencia de auditoría productiva de equidad, privacidad, latencia y seguridad."""))
    cells.append(nbf.v4.new_markdown_cell("""<div class="section"><h2>10 · Reproducibilidad, IA y referencias</h2></div>

Pipeline: `python -u codigo/v6/proyecto1_v6_pipeline.py`. Auditoría: `python codigo/v6/audit_project1_v6.py`. Los artefactos incluyen modelos, calibradores, umbrales, predicciones, contrato y manifiesto.

**Declaración de uso de inteligencia artificial.** Se utilizó asistencia para estructurar código, revisar consistencia, redactar y automatizar documentación. Los autores ejecutaron, verificaron y deben defender datos, decisiones, métricas y conclusiones.

**Referencias APA 7**

IEEE Computational Intelligence Society. (2019). *IEEE-CIS Fraud Detection* [Data set]. Kaggle. https://www.kaggle.com/competitions/ieee-fraud-detection/overview

Ke, G., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. *Advances in Neural Information Processing Systems, 30*.

Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. *PLOS ONE, 10*(3), e0118432. https://doi.org/10.1371/journal.pone.0118432

Zhou, C., & Paffenroth, R. C. (2017). Anomaly detection with robust deep autoencoders. *Proceedings of the 23rd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining*, 665–674. https://doi.org/10.1145/3097983.3098052"""))
    notebook = nbf.v4.new_notebook(cells=cells, metadata={"kernelspec":{"display_name":"Python 3","language":"python","name":"python3"},"language_info":{"name":"python","version":"3.13"}})
    out = out_dir / "proyecto1_calderon_barillas.ipynb"
    nbf.write(notebook, out)
    return out


def main() -> None:
    build_readme()
    notebook = build_notebook_v6()
    report = build_report()
    build_ficha()
    slides = build_slides()
    print("Entregables V6 generados desde", RESULT_PATH)
    print("Notebook:", notebook, "Informe:", report, "Presentación:", slides)


if __name__ == "__main__":
    main()
