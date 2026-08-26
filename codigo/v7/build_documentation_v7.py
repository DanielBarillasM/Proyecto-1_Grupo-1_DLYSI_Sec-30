"""Actualiza README, ficha y guion V7 desde resultados_v7.json."""

from __future__ import annotations

import json
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[2]
R = json.loads((ROOT / "artefactos" / "v7" / "resultados_v7.json").read_text(encoding="utf-8"))
URL = "https://github.com/DanielBarillasM/Proyecto-1_Grupo-1_DLYSI_Sec-30"


def f(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def pct(value: float, digits: int = 2) -> str:
    return f"{100 * value:.{digits}f}%"


def money(value: float) -> str:
    return f"Q{value:,.0f}"


def build_readme() -> str:
    i = R["evaluacion_interna"]
    b = R["benchmark_historico"]
    gate = R["promocion_V7"]
    c = R["hipotesis_C"]
    fals = R["falsificaciones"]
    walk = R["walk_forward"]
    text = f"""<div align="center">

# Proyecto 1 · Monitoreo transaccional · V7

### ¿El orden aporta señal incremental y cuánto vale económicamente?

![Python](https://img.shields.io/badge/Python-3.13-3776AB?logo=python&logoColor=white)
![PyTorch](https://img.shields.io/badge/PyTorch-2.13-EE4C2C?logo=pytorch&logoColor=white)
![LightGBM](https://img.shields.io/badge/LightGBM-4.7-184e77)
![CatBoost](https://img.shields.io/badge/CatBoost-1.2.10-e9c46a)
![Versión](https://img.shields.io/badge/Revisar-V7-2a9d8f)
![Estado](https://img.shields.io/badge/Promoción-exploratoria-f4b942)

**Wilson Alejandro Calderón Argueta · 22018** · **Pablo Daniel Barillas Moreno · 22193**

Universidad del Valle de Guatemala · Deep Learning y Sistemas Inteligentes · Sección 30 · 2026

</div>

> [!IMPORTANT]
> **La versión que debe revisarse es V7.** El último 15 % de IEEE-CIS ya fue observado en versiones anteriores y se reporta como benchmark temporal histórico reutilizado. Todas las decisiones V7 se toman dentro del período de desarrollo; una conclusión confirmatoria exige una cohorte futura etiquetada.

## Resumen ejecutivo

El proyecto analiza {R['datos']['filas']:,} transacciones IEEE-CIS, {R['datos']['fraudes']:,} fraudes y una prevalencia de {pct(R['datos']['prevalencia'])}. La pregunta no es simplemente qué arquitectura produce el número más alto, sino si el orden de eventos aporta información incremental frente a un baseline sin orden competitivo y si esa diferencia mejora una decisión de monitoreo con costos explícitos.

V7 amplía el espacio tabular a {R['datos']['columnas_union']} columnas después de ingeniería, compara regresión logística, LightGBM completo, LightGBM reducido por correlación, PCA, CatBoost y stacking. Mantiene el experimento obligatorio A/B/C: A es el baseline sin orden; B es una TCN causal congelada de V6; C fusiona predicciones fuera de tiempo; D es un encoder–decoder PyTorch entrenado solo con operaciones legítimas y funciona como control de anomalía.

El ganador interno es `{R['seleccion']['A']['seleccionado']}`. A obtiene AP {f(i['A']['auc_pr'],4)}, ROC-AUC {f(i['A']['roc_auc'],4)}, precisión {pct(i['A']['precision'])}, recall {pct(i['A']['recall'])}, F1 {f(i['A']['f1'],4)}, {i['A']['alertas_por_100k']:,.0f} alertas por 100,000 y costo {money(i['A']['cost_q'])}. Frente al control V6, aumenta AP {gate['delta_ap']:+.4f}, reduce costo {pct(gate['reduccion_costo'])} y reduce alertas {pct(-gate['crecimiento_alertas'])}; a cambio, recall baja de {pct(i['A_V6_control']['recall'])} a {pct(i['A']['recall'])}.

La mejora no se declara estable. Tres ventanas son favorables, pero una cae {min(row['delta_ap_V7_vs_V6'] for row in gate['ventanas']):+.4f} AP y el límite previo era −0.005. El gate de promoción es **{str(gate['success']).lower()}**. V7 queda como candidato exploratorio mejor equilibrado, no como reemplazo confirmatorio.

## Dónde está el código

Los notebooks ya no son resúmenes decorativos. El cuaderno oficial [`entregables/cuaderno/v7/proyecto1_calderon_barillas.ipynb`](entregables/cuaderno/v7/proyecto1_calderon_barillas.ipynb) materializa las definiciones ejecutables que sostienen el experimento: configuración temporal; unión e ingeniería causal; codificación aprendida con entrenamiento; selección por correlación; PCA; regresión logística; LightGBM; CatBoost; stacking; GRU; TCN causal; encoder–decoder; calibración; umbral económico; apuesta C; métricas y falsificaciones. Las celdas están ejecutadas y conservan sus salidas.

El interruptor `REENTRENAR_DESDE_CERO=False` evita que abrir o volver a ejecutar el informe dispare por accidente una corrida de varias horas. Con el valor predeterminado, el notebook reconstruye y verifica la evidencia congelada. Al cambiarlo a `True`, llama al pipeline íntegro. Esta separación distingue dos tareas que no deben confundirse: **leer y auditar una corrida reproducible** frente a **volver a estimar todos los modelos**.

El código mantenible se distribuye así:

| Archivo | Responsabilidad |
|---|---|
| `codigo/v7/proyecto1_v7_pipeline.py` | Orquestación integral, datos, características causales, familia A, stacking, C, calibración, umbrales, evaluación y exportación. |
| `codigo/v7/modelos_secuenciales_v7.py` | Implementaciones PyTorch de GRU, TCN causal y autoencoder, datasets, entrenamiento e inferencia. |
| `codigo/v7/build_notebooks_v7.py` | Inserta esas implementaciones reales en los notebooks; no reemplaza el entrenamiento. |
| `codigo/v7/finalize_v7.py` | Consolida predicciones, pruebas de orden, gates y resultados comunes. |
| `codigo/v7/audit_project1_v7.py` | Verifica artefactos, métricas, notebooks ejecutados, símbolos de modelos, páginas, diapositivas, rutas y secretos. |
| `entregables/cuaderno/v7/EDA_IEEE_CIS_Diagnostico_Datos_V7.ipynb` | Audita CSV originales, calidad, deriva, identidad proxy, asociación, correlación y PCA. |

Por tanto, el entrenamiento no está escondido en un JSON ni delegado a una librería sin mostrar su preparación. Los JSON y CSV son la salida verificable de una ejecución; el pipeline y las clases PyTorch son la implementación.

## Por qué aparece V6 en un experimento V7

V7 conserva **scores congelados** de la TCN y del autoencoder entrenados en V6 como controles de referencia. Esto no significa que falte código V7: las arquitecturas se encuentran ahora en `modelos_secuenciales_v7.py` y también se muestran dentro del notebook oficial. Se congelan las predicciones porque la pregunta experimental de esta iteración es si las nuevas variables, la reducción y el stacking agregan valor frente a un control previamente observado. Reentrenar simultáneamente el control cambiaría dos factores a la vez y volvería ambigua la atribución de la mejora.

La procedencia se declara en resultados, protocolo e informe. B es la TCN causal congelada; D es el encoder–decoder congelado; A5 puede incorporar su score como una característica de nivel meta; C contrasta la fusión. No se usan pesos V6 para aparentar que fueron producidos por V7, y el benchmark no vuelve a tratarse como ciego. En una siguiente cohorte confirmatoria sí deben congelarse el recipe completo, los umbrales y las versiones de dependencias antes de obtener etiquetas.

Eliminar la referencia a V6 haría el repositorio visualmente más aislado, pero científicamente perdería el comparador que permite medir el incremento. La solución correcta fue volver autónoma la implementación de V7 y mantener explícito el origen del control.

## Resultado principal

| Modelo | AP interna | ROC-AUC | Precisión | Recall | F1 | Costo | Alertas/100k |
|---|---:|---:|---:|---:|---:|---:|---:|
| **A · A5 stacking** | **{f(i['A']['auc_pr'],4)}** | **{f(i['A']['roc_auc'],4)}** | **{pct(i['A']['precision'])}** | {pct(i['A']['recall'])} | **{f(i['A']['f1'],4)}** | {money(i['A']['cost_q'])} | **{i['A']['alertas_por_100k']:,.0f}** |
| B · TCN causal | {f(i['B']['auc_pr'],4)} | {f(i['B']['roc_auc'],4)} | {pct(i['B']['precision'])} | {pct(i['B']['recall'])} | {f(i['B']['f1'],4)} | {money(i['B']['cost_q'])} | {i['B']['alertas_por_100k']:,.0f} |
| C · A+B | {f(i['C']['auc_pr'],4)} | {f(i['C']['roc_auc'],4)} | {pct(i['C']['precision'])} | **{pct(i['C']['recall'])}** | {f(i['C']['f1'],4)} | **{money(i['C']['cost_q'])}** | {i['C']['alertas_por_100k']:,.0f} |
| D · autoencoder | {f(i['D']['auc_pr'],4)} | {f(i['D']['roc_auc'],4)} | {pct(i['D']['precision'])} | {pct(i['D']['recall'])} | {f(i['D']['f1'],4)} | {money(i['D']['cost_q'])} | {i['D']['alertas_por_100k']:,.0f} |

C cuesta ligeramente menos, pero no se promueve: cambia AP {c['delta_ap']:+.4f}, reduce costo solo {pct(c['reduccion_costo'])} y no mejora ninguna de cuatro ventanas. La regla previa exigía +0.01 AP, −5 % costo, recall ≥0.75, crecimiento de alertas ≤10 % y mejora en tres ventanas. D conserva recall, pero su precisión de {pct(i['D']['precision'])} confirma que rareza y fraude no son sinónimos.

## Cómo interpretar las métricas

### AP, no PA

La abreviatura correcta es **AP**, de *Average Precision* o precisión promedio. “PA” no es una métrica diferente en este proyecto; es una inversión accidental de las siglas. AP aproxima el área bajo la curva precisión–recall mediante un promedio ponderado de la precisión conseguida cada vez que aumenta el recall. En forma discreta:

$$
AP = \\sum_n (R_n-R_{{n-1}})P_n.
$$

Su referencia ingenua es aproximadamente la prevalencia positiva: {pct(R['datos']['prevalencia'])}. Por ello, AP {f(i['A']['auc_pr'],3)} representa un ranking muy superior al azar en este conjunto desbalanceado. No significa, sin embargo, “{100*i['A']['auc_pr']:.1f} % de fraudes acertados” ni “{100*i['A']['auc_pr']:.1f} % de alertas correctas”. Para responder esas preguntas se necesitan recall y precisión en el umbral operativo.

- **AP o Average Precision** resume la curva precisión–recall a través de umbrales. AP {f(i['A']['auc_pr'],3)} no significa que {100*i['A']['auc_pr']:.1f} % de alertas sea correcta; esa pureza puntual la expresa la precisión ({pct(i['A']['precision'])}). AP es principal porque la clase positiva representa solo {pct(R['datos']['prevalencia'])}.
- **ROC-AUC** aproxima la probabilidad de ordenar un fraude por encima de una operación legítima elegida al azar. Un ROC {f(i['A']['roc_auc'],3)} muestra buena separación global, pero puede coexistir con muchas falsas alarmas cuando hay millones de negativos.
- **Precisión** responde: “de todas las alertas, ¿cuántas son fraude?”. A logra {pct(i['A']['precision'])}, aproximadamente una alerta verdadera por cada cinco.
- **Recall** responde: “de todos los fraudes, ¿cuántos se detectaron?”. A recupera {pct(i['A']['recall'])}, casi ocho de diez.
- **F1** combina precisión y recall, pero no conoce el costo en quetzales. Por eso se reporta junto con `Q4,200×FN + Q180×FP`.
- **Precision@1 %** es {pct(i['A']['precision_at_1pct'])}: si solo se revisa el 1 % de mayor riesgo, casi nueve de diez seleccionadas son fraude. **Recall@1 %** es {pct(i['A']['recall_at_1pct'])}: esa capacidad limitada captura alrededor de tres de diez fraudes.

### Lectura operativa conjunta

Las métricas describen capas distintas. ROC-AUC y AP evalúan el **ordenamiento continuo** antes de fijar una política. Precisión, recall y F1 describen una **decisión binaria** en un umbral. Alertas por 100,000 mide la **carga de trabajo**. El costo traduce FP y FN a la función académica `Q4,200×FN + Q180×FP`. Un modelo puede ganar ROC-AUC y perder costo si sus errores se concentran cerca del umbral; también puede mejorar F1 y perjudicar recall, lo cual es costoso cuando un FN vale 23.3 veces un FP.

En A, la precisión de {pct(i['A']['precision'])} implica que alrededor de cuatro de cada cinco alertas son falsas, pero eso no invalida automáticamente el modelo: con prevalencia {pct(R['datos']['prevalencia'])}, la muestra alertada está enriquecida más de cinco veces. Al mismo tiempo, {i['A']['alertas_por_100k']:,.0f} alertas por 100,000 pueden superar la capacidad de un equipo pequeño. La conclusión defendible es que el ranking es útil y captura {pct(i['A']['recall'])} de los fraudes con el umbral académico, mientras la política final necesita cupo diario y costos reales.

## Arquitecturas y preguntas que responde cada una

| Apuesta | Arquitectura | Qué prueba | Resultado V7 |
|---|---|---|---|
| A0 | Regresión logística regularizada | Si una frontera lineal bien controlada ya explica la señal. | Baseline interpretable y corroboración. |
| A1–A4 | LightGBM completo/reducido/PCA y CatBoost | Si no linealidad, variables ampliadas, reducción o categóricas elevan AP. | Aportan diversidad; ninguna aislada desplaza de forma estable al control. |
| A5 | Stacking logístico con predicciones fuera de tiempo | Si los errores complementarios mejoran ranking y costo. | Mejor candidato interno; gate estable no superado. |
| B | TCN causal PyTorch | Si la historia ordenada aporta señal incremental. | Predice, pero las falsificaciones no atribuyen la ganancia al orden. |
| C | Fusión A+B y variante A+B+D | Si combinar señal tabular, temporal y anomalía supera A bajo regla previa. | Rechazada: pierde AP y no mejora ventanas suficientes. |
| D | Encoder–decoder entrenado con legítimas | Si el error de reconstrucción detecta fraude raro. | Recall alto, precisión baja; rareza no equivale a fraude. |

La familia A responde “¿qué tan bien se ordena el riesgo sin requerir orden explícito?”. B responde “¿cambia la evidencia al destruir el orden?”. C responde “¿la combinación añade valor incremental?”. D responde “¿un modelo de normalidad ayuda bajo desbalance?”. Esta trazabilidad evita vender cuatro modelos como si persiguieran exactamente el mismo objetivo.

## Datos, orden y prevención de fuga

La fuente es [IEEE-CIS Fraud Detection](https://www.kaggle.com/competitions/ieee-fraud-detection/overview), publicada en Kaggle con datos anonimizados proporcionados por Vesta Corporation. `train_transaction.csv` y `train_identity.csv` se unen por `TransactionID`; las filas se ordenan por `TransactionDT`. Ambos campos se excluyen como magnitudes predictivas: el primero solo alinea evidencia y el segundo define el reloj.

La partición temporal es 70 % entrenamiento, 15 % validación y 15 % benchmark histórico. Validación se subdivide, en ese orden, en early stopping, `meta_fit`, `model_select`, calibración, umbral y evaluación. Toda imputación, frecuencia, asociación con fraude, correlación de Spearman, selección y PCA se ajusta exclusivamente con train. Los tres walk-forward vuelven a ajustar el preprocesamiento dentro de cada pliegue.

Las variables causales usan solo eventos anteriores: frecuencias previas de tarjeta, dirección, correo, dispositivo y producto; monto histórico; recencia; faltantes y resúmenes C/D/V/identidad. La entidad proxy se diagnostica con cuatro definiciones. Ninguna clave anonimizada equivale necesariamente a un cliente real.

## Correlación y PCA

El análisis encuentra {R['variables']['correlacion']['pares_eliminados']} relaciones con `|ρ de Spearman| ≥ {R['variables']['correlacion']['umbral']}` en la muestra train-only y conserva {R['variables']['correlacion']['variables_retenidas']} representantes. Esto reduce redundancia extrema, no “variables poco correlacionadas con fraude” de forma ciega. Una variable con baja asociación marginal todavía puede ser útil mediante interacciones.

PCA se limita al bloque `V1–V339`. Se ajustan 128 componentes con una muestra determinista contenida en train; se necesitan {R['variables']['pca']['componentes_para_90']} para 90 % y {R['variables']['pca']['componentes_para_95']} para 95 % de varianza. La mejor variante PCA alcanza AP de selección {R['seleccion']['A']['pca_auc_pr_model_select']['A3_pca_128']:.4f}, por debajo de LightGBM completo y correlación. La conclusión es predictiva: PCA comprime, pero no mejora el detector.

## Valor del orden

La prueba principal mantiene el evento objetivo al final y permuta solo antecedentes con cinco semillas. B original obtiene AP {f(fals['original_internal']['auc_pr'],4)} y la media permutada {f(fals['permutation_mean_auc_pr'],4)} ± {f(fals['permutation_std_auc_pr'],4)}. La diferencia original−permutada es {fals['order_auc_pr_drop']:+.4f}; se exigía una caída positiva mínima de 0.01. Por tanto, no se afirma que el orden aporte.

El segundo intento recorta la historia sin reentrenar: 3 eventos producen AP {f(fals['historia_3']['auc_pr'],4)}, 8 producen {f(fals['historia_8']['auc_pr'],4)}, 16 producen {f(fals['historia_16']['auc_pr'],4)} y 32 producen {f(fals['historia_32']['auc_pr'],4)}. No existe patrón monotónico. B aprende señal, pero esa señal puede provenir del evento actual o la composición histórica, no del orden.

## Estabilidad y benchmark histórico

Los walk-forward del recipe reducido alcanzan AP {' / '.join(f(row['auc_pr'],3) for row in walk)}. La dispersión evidencia deriva. En las cuatro ventanas del gate, las diferencias A5−V6 son {' / '.join(f"{row['delta_ap_V7_vs_V6']:+.4f}" for row in gate['ventanas'])}. La tercera ventana impide promoción estable.

El benchmark histórico, solo descriptivo, muestra:

| Modelo | AP | Precisión | Recall | F1 | Costo |
|---|---:|---:|---:|---:|---:|
| A | {f(b['A']['auc_pr'],4)} | {pct(b['A']['precision'])} | {pct(b['A']['recall'])} | {f(b['A']['f1'],4)} | {money(b['A']['cost_q'])} |
| B | {f(b['B']['auc_pr'],4)} | {pct(b['B']['precision'])} | {pct(b['B']['recall'])} | {f(b['B']['f1'],4)} | {money(b['B']['cost_q'])} |
| C | {f(b['C']['auc_pr'],4)} | {pct(b['C']['precision'])} | {pct(b['C']['recall'])} | {f(b['C']['f1'],4)} | {money(b['C']['cost_q'])} |
| D | {f(b['D']['auc_pr'],4)} | {pct(b['D']['precision'])} | {pct(b['D']['recall'])} | {f(b['D']['f1'],4)} | {money(b['D']['cost_q'])} |

## Reproducción rápida

```powershell
py -3.13 -m venv .venv
.\\.venv\\Scripts\\Activate.ps1
python -m pip install --upgrade pip
python -m pip install -r configuracion/v7/requirements-v7.txt
python -m pip install -r configuracion/v7/requirements-docs-v7.txt
$env:PROYECTO1_RAW=(Resolve-Path datos/raw)
python -u codigo/v7/proyecto1_v7_pipeline.py
python codigo/v7/build_notebooks_v7.py
jupyter nbconvert --to notebook --execute --inplace --ExecutePreprocessor.timeout=600 entregables/cuaderno/v7/proyecto1_calderon_barillas.ipynb
python codigo/v7/report_v7.py
python codigo/v7/presentation_v7.py
python codigo/v7/build_documentation_v7.py
python codigo/v7/audit_project1_v7.py
```

La descarga requiere aceptar las reglas de Kaggle. Los CSV, credenciales y tokens no se versionan. Consulte [`configuracion/v7/INSTRUCCIONES_V7.md`](configuracion/v7/INSTRUCCIONES_V7.md) para la corrida completa.

### Ejecución desde los notebooks

1. Coloque los cuatro CSV de la competencia en `datos/raw/` o defina `PROYECTO1_RAW`.
2. Abra primero el EDA para confirmar dimensiones, columnas, prevalencia y orden temporal.
3. Abra el notebook oficial. Mantenga `REENTRENAR_DESDE_CERO=False` para revisar la evidencia entregada.
4. Para una corrida nueva, cambie el interruptor a `True`, asegure memoria y tiempo suficientes y ejecute desde el inicio.
5. Reconstruya informe, presentación y documentación únicamente después de terminar modelos; todos leen `resultados_v7.json` como fuente única.
6. Ejecute la auditoría. Un resultado `APROBADO` verifica integridad del paquete, no validez externa futura.

### Problemas frecuentes

- **No se encuentran los CSV:** revise `PROYECTO1_RAW`; no copie credenciales al repositorio.
- **CatBoost o LightGBM no importan:** instale las versiones exactas de `requirements-v7.txt` dentro del mismo entorno de Jupyter.
- **El notebook termina demasiado rápido:** con el interruptor en `False` está auditando la corrida congelada; no está simulando un reentrenamiento.
- **Las métricas cambian:** confirme orden por `TransactionDT`, semillas, versiones y que ningún transformador fue ajustado con validación o benchmark.
- **El PDF cambia de páginas:** use el builder correspondiente y vuelva a ejecutar la auditoría; el informe debe tener como máximo siete páginas y la presentación ocho diapositivas.

## Correspondencia con la rúbrica

| Evidencia solicitada | Ubicación V7 | Veredicto |
|---|---|---|
| Integridad de datos, entidad y protocolo temporal | EDA, cuaderno oficial, informe §2 y protocolo | Cubierta; se declaran límites de identidad proxy. |
| Comparación común A/B con AP, precisión, recall y F1 | Cuaderno, `resultados_v7.json`, informe y diapositiva 6 | Cubierta sobre la misma población. |
| Permutación y segundo intento sobre historia | Cuaderno, falsificaciones JSON, informe §4 y diapositiva 7 | Cubierta; resultado negativo reportado. |
| Hipótesis previa, control y veredicto C | Protocolo, cuaderno, informe §5 | Cubierta; C no se promueve. |
| Umbral, costo y recomendación | Cuaderno, umbrales, informe §6 y diapositiva 8 | Cubierta con costos explícitos. |
| Reproducibilidad y comunicación | Código, notebooks ejecutados, artefactos, README, informe de 7 páginas y presentación de 8 | Cubierta y auditada automáticamente. |

La cobertura formal de la rúbrica no convierte el resultado en validación externa. El punto pendiente es confirmatorio: una cohorte futura no observada con identidad y costos operativos más fiables.

## Estructura y versiones

Las carpetas principales se mantienen en la raíz; cada una contiene subcarpetas por versión.

```text
codigo/v7/                 pipeline, finalización, builders y auditoría
configuracion/v7/          protocolo, instrucciones y dependencias exactas
datos/processed/v7/        asociación, correlación y auditoría
artefactos/v7/             modelos A0–A5/C, scores, umbrales y contrato
evidencia/figuras/v7/      seis figuras reproducibles
entregables/cuaderno/v7/   notebook oficial y EDA ejecutados
entregables/informe/v7/    LaTeX y PDF de siete páginas
entregables/presentacion/v7/ HTML/PDF de ocho diapositivas con notas integradas
entregables/ficha/v7/      ficha DOCX/PDF y copia del README
```

| Versión | Enfoque | Estado |
|---|---|---|
| V1 | Experimento A/B/C original | Histórica |
| V2 | LightGBM, walk-forward y calibración | Histórica |
| V3 | Regresión logística, PCA y controles | Histórica |
| V4 | LightGBM/CatBoost/XGBoost y expertos | Baseline fuerte heredado |
| V5 | Reintegración rubricada A/B/C | Histórica |
| V6 | EDA, TCN y encoder–decoder | Control experimental |
| **V7** | Variables completas, correlación/PCA, CatBoost y stacking con V6 | **Versión a revisar** |

## Candidato al Proyecto Final

- **Modelo:** A5, stacking logístico de A0–A4 y control V6.
- **Usuario previsto:** equipo de monitoreo o analista de riesgo.
- **Decisión:** priorizar transacciones para revisión bajo capacidad limitada.
- **Entrada:** transacción actual y agregados estrictamente causales descritos en `artefactos/v7/contrato_entrada_salida_v7.json`.
- **Salida:** `risk_score` calibrado en `[0,1]`; umbral {R['umbrales']['A']:.8f}; indicador binario derivado.
- **Faltantes:** medianas y frecuencias aprendidas en train; categorías nuevas reciben frecuencia cero.
- **Riesgos:** deriva, identidad proxy, falsos positivos, fraude adaptativo, costos hipotéticos y benchmark reutilizado.
- **Pendiente confirmatorio:** cohorte futura, identidad bancaria fiable, costos/capacidad reales, latencia, privacidad, equidad, explicación y monitoreo.

## Tres decisiones técnicas importantes

1. **Conservar el baseline V6 dentro de A5.** Alternativa: reemplazarlo por el mejor modelo nuevo. Evidencia: los nuevos aislados caen en evaluación; el stacking con V6 alcanza AP {f(i['A']['auc_pr'],4)} y costo {money(i['A']['cost_q'])}.
2. **Validar correlación y PCA como ablations.** Alternativa: borrar automáticamente columnas correlacionadas o comprimir todo el dataset. Evidencia: A2 gana selección frente a A1, pero PCA y la reducción aislada no sostienen superioridad.
3. **Rechazar C y el valor del orden según reglas previas.** Alternativa: promover C por el benchmark o asumir que TCN usa orden. Evidencia: C pierde AP interna y permutar antecedentes no perjudica B.

## Limitaciones y uso responsable

IEEE-CIS está anonimizado y cubre un período finito. La identidad es una aproximación, los costos son académicos y el benchmark no es ciego. El score mide riesgo estadístico, no certeza ni culpabilidad. Antes de cualquier despliegue se requieren pruebas con una cohorte futura, auditoría de privacidad/equidad, explicaciones, latencia, seguridad, monitoreo de deriva y un procedimiento de apelación.

## Declaración de uso de inteligencia artificial

Se utilizó asistencia de IA para estructurar y revisar código, diseñar documentación HTML/CSS/LaTeX, automatizar auditorías y mejorar la redacción. Los integrantes ejecutaron el pipeline, comprobaron alineación de IDs, particiones, transformaciones train-only, métricas, falsificaciones, gates y artefactos. La IA no se usa como fuente académica y no sustituye la defensa del proyecto.

## Referencias APA 7

IEEE Computational Intelligence Society. (2019). *IEEE-CIS Fraud Detection* [Data set]. Kaggle. https://www.kaggle.com/competitions/ieee-fraud-detection/overview

Ke, G., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. *Advances in Neural Information Processing Systems, 30*.

Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A. V., & Gulin, A. (2018). CatBoost: Unbiased boosting with categorical features. *Advances in Neural Information Processing Systems, 31*.

Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. *PLOS ONE, 10*(3), e0118432. https://doi.org/10.1371/journal.pone.0118432

Zhou, C., & Paffenroth, R. C. (2017). Anomaly detection with robust deep autoencoders. *Proceedings of the 23rd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining*, 665–674. https://doi.org/10.1145/3097983.3098052
"""
    out = ROOT / "entregables" / "ficha" / "v7"
    out.mkdir(parents=True, exist_ok=True)
    (out / "README_V7_GENERADO.md").write_text(text, encoding="utf-8", newline="\n")
    (ROOT / "README.md").write_text(text, encoding="utf-8", newline="\n")
    return text


def build_script() -> str:
    i, gate, fals = R["evaluacion_interna"], R["promocion_V7"], R["falsificaciones"]
    text = f"""# Guion de exposición · Proyecto 1 V7

Duración sugerida: 10–12 minutos. El guion sigue las ocho diapositivas; presione `N` en el HTML para mostrar notas del expositor.

## 1. Decisión ejecutiva · 60 segundos

“Nuestra V7 mejora el promedio y la utilidad operativa, pero todavía no demuestra estabilidad suficiente. El candidato es A5, un stacking que conserva el baseline V6 y añade modelos nuevos. Internamente alcanza AP {f(i['A']['auc_pr'],3)}, precisión {pct(i['A']['precision'],1)}, recall {pct(i['A']['recall'],1)} y costo {money(i['A']['cost_q'])}. Frente a V6 mejora AP {gate['delta_ap']:+.3f}, reduce costo {pct(gate['reduccion_costo'],1)} y reduce alertas, pero una ventana temporal cae demasiado. Por eso decimos candidato exploratorio, no ganador definitivo.”

## 2. Datos y problema · 60 segundos

“Usamos las 590,540 transacciones de IEEE-CIS Fraud Detection, publicadas en Kaggle con datos anonimizados de Vesta Corporation. Solo 3.5 % son fraude; por eso una accuracy de 96.5 % podría lograrse sin detectar un solo caso positivo. TransactionID une las tablas, TransactionDT define el reloj y ambos se excluyen como magnitudes predictivas.”

## 3. Protocolo temporal · 70 segundos

“La separación es 70 % train, 15 % validación y 15 % benchmark histórico. Validación se subdivide en early stopping, meta_fit, model_select, calibración, umbral y evaluación. Toda imputación, correlación, selección y PCA se aprende solo con pasado. El benchmark ya fue observado y no decide nada en V7.”

## 4. EDA, correlación y PCA · 70 segundos

“El EDA muestra alta dimensionalidad, faltantes, categorías y redundancia. Spearman elimina solo 34 representantes con correlación absoluta al menos 0.995. PCA se aplica únicamente a V1–V339: 105 componentes explican 95 % de varianza, pero la variante PCA no gana AP. La lección es que conservar varianza no garantiza conservar señal de fraude.”

## 5. Diseño A/B/C/D y familia A · 110 segundos

“A0 es regresión logística; A1 LightGBM completo; A2 LightGBM reducido; A3 PCA; A4 CatBoost y A5 stacking. A es el candidato sin orden; B es la TCN causal; C integra A, B y opcionalmente D; D es el encoder-decoder entrenado con operaciones legítimas. Ningún modelo nuevo aislado generaliza mejor que V6. A5 sí mejora porque combina logits fuera de tiempo y alcanza AP de selección {R['seleccion']['A']['auc_pr_model_select']['A5_ensamble_tabular']:.3f}. C debía mejorar AP, costo, recall, alertas y estabilidad, no solo una cifra.”

## 6. Resultados y significado de métricas · 95 segundos

“A gana AP, precisión y F1. AP evalúa el ranking precisión-recall y se compara con prevalencia 0.035; no es la proporción puntual de alertas correctas. Esa proporción es precisión: A logra {pct(i['A']['precision'],1)}, aproximadamente una alerta correcta de cada cinco. B queda en AP {i['B']['auc_pr']:.3f}; C pierde AP y D logra precisión de solo {pct(i['D']['precision'],1)}. ROC describe separación global, pero no la carga de falsas alarmas.”

## 7. Valor del orden · 75 segundos

“Probamos el orden en vez de asumirlo. Al barajar antecedentes cinco veces, AP cambia de {f(fals['original_internal']['auc_pr'],4)} a {f(fals['permutation_mean_auc_pr'],4)}. La diferencia es {fals['order_auc_pr_drop']:+.4f}: destruir el orden no perjudica. Recortar a 3, 8, 16 y 32 tampoco produce mejora monotónica. B aprende algo, pero no podemos atribuírselo al orden.”

## 8. Economía, estabilidad y decisión · 105 segundos

“Usamos costo Q4,200 por FN y Q180 por FP, con recall mínimo 0.75. A5 produce 1,700 FP y 124 FN internamente: Q826,800. Reduce las alertas a {i['A']['alertas_por_100k']:,.0f} por 100 mil. Sin embargo, las cuatro diferencias AP contra V6 son {'; '.join(f"{row['delta_ap_V7_vs_V6']:+.3f}" for row in gate['ventanas'])}; una cae −0.014. Conservamos A5 como candidato exploratorio, pero esa deriva bloquea la promoción confirmatoria. No promovemos B porque no demuestra orden; C porque falla la hipótesis; D porque genera demasiadas falsas alarmas. La decisión cambiaría con una cohorte futura, sin caídas por ventana, con identidad real y costos operativos.”

## Respuestas cortas ante preguntas

- **¿AP 0.55 es bajo?** No se compara con 1 de forma aislada; la base es prevalencia 0.035. Es un ranking útil, aunque no perfecto.
- **¿Por qué ROC alto y precisión moderada?** Porque hay muchísimas operaciones legítimas; una tasa pequeña de FP genera muchas alertas.
- **¿Por qué no usar accuracy?** Un clasificador “todo legítimo” tendría 96.5 % sin detectar fraude.
- **¿Por qué no hacer la TCN más grande?** La permutación muestra que el cuello de botella es identidad/representación, no capacidad.
- **¿V7 supera V6?** Mejora promedio, costo y carga, pero no cumple estabilidad; necesita cohorte futura.
- **¿Por qué C no gana si cuesta menos?** Su hipótesis exigía mejora conjunta de AP, costo y estabilidad; pierde AP y no mejora ventanas.
"""
    out = ROOT / "entregables" / "presentacion" / "v7" / "GUION_EXPOSICION_V7.md"
    out.write_text(text, encoding="utf-8", newline="\n")
    return text


def build_ficha() -> None:
    out = ROOT / "entregables" / "ficha" / "v7"
    out.mkdir(parents=True, exist_ok=True)
    resources = ROOT / "evidencia" / "recursos" / "v7"
    resources.mkdir(parents=True, exist_ok=True)
    qr_path = resources / "qr_repositorio_v7.png"
    qrcode.make(URL).save(qr_path)
    b = R["benchmark_historico"]["A"]
    gate = R["promocion_V7"]

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(.42); sec.bottom_margin = Inches(.42); sec.left_margin = Inches(.55); sec.right_margin = Inches(.55)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROYECTO 1 · MONITOREO TRANSACCIONAL · V7"); run.bold = True; run.font.size = Pt(19); run.font.color.rgb = RGBColor(24, 78, 119)
    sub = doc.add_paragraph("Universidad del Valle de Guatemala · Deep Learning y Sistemas Inteligentes · Sección 30"); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    names = doc.add_paragraph("Wilson Alejandro Calderón Argueta · 22018 | Pablo Daniel Barillas Moreno · 22193"); names.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=2, cols=4); table.style = "Table Grid"
    values = [("AP histórico", f(b["auc_pr"])), ("ROC-AUC", f(b["roc_auc"])), ("Precisión", pct(b["precision"])), ("Recall", pct(b["recall"])), ("F1", f(b["f1"])), ("Costo", money(b["cost_q"])), ("ΔAP interno", f(gate["delta_ap"], 4)), ("Candidato", "A5 exploratorio")]
    for cell, (label, value) in zip([c for row in table.rows for c in row.cells], values):
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        a = p.add_run(label.upper() + "\n"); a.bold = True; a.font.size = Pt(8)
        v = p.add_run(value); v.bold = True; v.font.size = Pt(12); v.font.color.rgb = RGBColor(42, 157, 143)
    doc.add_heading("Fuente y alcance", level=1)
    doc.add_paragraph("IEEE-CIS Fraud Detection · Vesta Corporation · Kaggle\n" + "https://www.kaggle.com/competitions/ieee-fraud-detection/overview")
    doc.add_heading("Resultado", level=1)
    doc.add_paragraph("A5 combina A0–A4 con el control V6. Mejora AP y costo promedio, pero una de cuatro ventanas temporales se degrada; la promoción confirmatoria permanece pendiente. B no demuestra valor material del orden y C/D no cumplen sus criterios.")
    doc.add_heading("Qué contiene el repositorio", level=1)
    doc.add_paragraph("Notebook oficial y EDA ejecutados; modelos, preprocesamiento, scores, umbrales y contrato; informe LaTeX/PDF de siete páginas; presentación HTML/PDF de ocho diapositivas con notas; guion y auditoría reproducible.")
    pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER; pic.add_run().add_picture(str(qr_path), width=Inches(1.15))
    link = doc.add_paragraph(URL); link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(out / "Ficha_Repositorio_Proyecto_1_V7.docx")

    styles = getSampleStyleSheet()
    pdf = SimpleDocTemplate(str(out / "Ficha_Repositorio_Proyecto_1_V7.pdf"), pagesize=letter, rightMargin=.55*inch, leftMargin=.55*inch, topMargin=.45*inch, bottomMargin=.45*inch)
    story = [Paragraph("<b>PROYECTO 1 · MONITOREO TRANSACCIONAL · V7</b>", styles["Title"]), Paragraph("Universidad del Valle de Guatemala · Grupo 1 · Sección 30", styles["Heading3"]), Paragraph("Wilson Alejandro Calderón Argueta · 22018 | Pablo Daniel Barillas Moreno · 22193", styles["BodyText"]), Spacer(1, 10)]
    data = [["AP A5", f(b["auc_pr"]), "Recall", pct(b["recall"])], ["Precisión", pct(b["precision"]), "Costo", money(b["cost_q"])], ["ΔAP interno", f(gate["delta_ap"], 4), "Estado", "Exploratorio"]]
    tab = Table(data, colWidths=[1.1*inch, 1.35*inch, 1.1*inch, 1.35*inch]); tab.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#EDF5FB")), ("GRID", (0,0), (-1,-1), .5, colors.HexColor("#9FB7C8")), ("FONTNAME", (0,0), (-1,-1), "Helvetica-Bold"), ("ALIGN", (0,0), (-1,-1), "CENTER"), ("PADDING", (0,0), (-1,-1), 8)]))
    qr = RLImage(str(qr_path), width=1.12*inch, height=1.12*inch); qr.hAlign = "CENTER"
    story.extend([tab, Spacer(1, 12), Paragraph("Datos: IEEE-CIS Fraud Detection, Vesta Corporation, Kaggle.", styles["BodyText"]), Paragraph("A5 mejora promedio y costo, pero no cumple estabilidad; B no demuestra orden y C/D no se promueven.", styles["BodyText"]), Spacer(1, 8), Paragraph("Versión a revisar: V7. El benchmark es histórico reutilizado.", styles["BodyText"]), Spacer(1, 10), qr, Spacer(1, 6), Paragraph(URL, styles["Heading3"])])
    pdf.build(story)


def main() -> None:
    build_readme()
    build_script()
    build_ficha()
    print("README, guion y ficha V7 actualizados desde resultados_v7.json")


if __name__ == "__main__":
    main()
